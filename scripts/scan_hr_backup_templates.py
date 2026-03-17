"""Scan backup templates in HR Documents/templates/ folder."""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

TEMPLATES_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "HR Documents", "templates"
)

TARGET = sys.argv[1] if len(sys.argv) > 1 else None

for fname in sorted(os.listdir(TEMPLATES_DIR)):
    if not fname.endswith(('.docx', '.DOCX')):
        continue
    if TARGET and TARGET.lower() not in fname.lower():
        continue
    path = os.path.join(TEMPLATES_DIR, fname)
    doc = Document(path)
    print(f"\n{'='*70}\nFILE: {fname}\n  {len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs")
    for ti, tbl in enumerate(doc.tables[:3]):
        print(f"  Table {ti}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
        for ri, row in enumerate(tbl.rows[:5]):
            cells_text = [repr((cell.text or "").strip().replace("\n"," ")[:40]) for cell in row.cells]
            if any(c != repr("") for c in cells_text):
                print(f"    R{ri}: {' | '.join(cells_text)}")
    for i, p in enumerate(doc.paragraphs[:10]):
        t = (p.text or "").strip()
        if t:
            print(f"  P{i}: {repr(t[:100])}")
