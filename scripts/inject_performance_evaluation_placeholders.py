"""
Inject docxtpl placeholders into Employee Performance Evaluation Form - INJAAZ.DOCX
so that UI form data is filled when downloading DOCX.

Run from project root: python scripts/inject_performance_evaluation_placeholders.py
"""
import os
import sys
import shutil

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

def inject(doc_path, backup=True):
    if not os.path.isfile(doc_path):
        raise FileNotFoundError(f"Template not found: {doc_path}")
    if backup:
        backup_dir = os.path.join(os.path.dirname(doc_path), "templates")
        os.makedirs(backup_dir, exist_ok=True)
        backup_path = os.path.join(backup_dir, "Employee Performance Evaluation Form - INJAAZ (before placeholders).docx")
        shutil.copy2(doc_path, backup_path)
        print(f"Backed up to: {backup_path}")

    doc = Document(doc_path)
    # Table 1: Employee info (4 rows x 4 cols). Value cells are 1 and 3.
    t1 = doc.tables[1]
    t1.rows[0].cells[1].text = "{{ employee_name }}"
    t1.rows[0].cells[3].text = "{{ date_of_evaluation }}"
    t1.rows[1].cells[1].text = "{{ employee_id }}"
    t1.rows[1].cells[3].text = "{{ date_of_joining }}"
    t1.rows[2].cells[1].text = "{{ department }}"
    t1.rows[2].cells[3].text = "{{ designation }}"
    t1.rows[3].cells[1].text = "{{ evaluation_done_by }}"

    # Table 2: Scores. Col 3 is SCORE for rows 3-12, row 13 col 3 is overall.
    t2 = doc.tables[2]
    for i in range(1, 11):
        t2.rows[2 + i].cells[3].text = "{{ score_%02d }}" % i
    t2.rows[13].cells[3].text = "{{ overall_score }}"

    # Table 3: Evaluation details
    t3 = doc.tables[3]
    t3.rows[1].cells[1].text = "{{ evaluator_name }}"
    t3.rows[1].cells[4].text = "{{ evaluator_designation }}"
    t3.rows[2].cells[1].text = "{{ evaluator_observation }}"
    t3.rows[3].cells[1].text = "{{ area_of_concern }}"
    t3.rows[4].cells[1].text = "{{ training_required }}"
    t3.rows[5].cells[1].text = "{{ employee_comments }}"
    t3.rows[6].cells[1].text = "{{ employee_signature }}"
    t3.rows[6].cells[4].text = "{{ employee_sign_date }}"
    t3.rows[7].cells[1].text = "{{ evaluator_signature }}"
    t3.rows[7].cells[4].text = "{{ evaluator_sign_date }}"

    # Table 4: Manager / GM / HR
    t4 = doc.tables[4]
    t4.rows[0].cells[1].text = "{{ concern_incharge_name }}"
    t4.rows[1].cells[1].text = "{{ incharge_comments }}"
    t4.rows[2].cells[1].text = "{{ gm_remarks }}"
    t4.rows[4].cells[1].text = "{{ hr_remarks }}"

    doc.save(doc_path)
    print("Placeholders injected. Saved:", doc_path)

if __name__ == "__main__":
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(base, "HR Documents", "Employee Performance Evaluation Form - INJAAZ.DOCX")
    inject(path, backup=True)
