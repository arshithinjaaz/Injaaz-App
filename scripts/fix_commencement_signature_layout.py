"""
Fix Commencement Form signature layout:
1. Signature paragraphs (29, 39): label on left, signature + date centered in the box (via center tab).
2. Clear duplicate date paragraphs (31, 41).

Run from project root: python scripts/fix_commencement_signature_layout.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches


def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(base, "HR Documents", "Commencement Form - INJAAZ.DOCX")
    if not os.path.isfile(path):
        print("Template not found:", path)
        return
    doc = Document(path)

    # Signature paragraphs: label left, then center tab (indices after header table)
    sig_rewrites = [
        (22, "Employee's Signature:", "{{ employee_signature }}    Date: {{ employee_sign_date }}"),
        (32, "Signature:", "{{ reporting_to_signature }}    Date: {{ reporting_sign_date }}"),
    ]
    for idx, label, placeholders in sig_rewrites:
        if idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        para.clear()
        # Center tab: content after \t will be centered at this position
        para.paragraph_format.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
        para.add_run(label)
        para.add_run("\t")
        para.add_run(placeholders)
        print(f"Rewrote paragraph {idx}: label left, signature+date centered")

    # Remove duplicate date lines (indices after header table)
    for idx in (24, 34):
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].clear()
            print(f"Cleared paragraph {idx}")

    doc.save(path)
    print("Saved. Signature and date are now centered in the box.")


if __name__ == "__main__":
    main()
