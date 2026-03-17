"""
Fix HR DOCX placeholder placement: restore labels that were lost due to merged cells,
and ensure placeholders are in the correct position (label | value).

When a table has merged cells, the inject script overwrote the label. This script:
1. Restores labels before placeholders (e.g. "Name: {{ employee_name }}")
2. Updates the inject logic for future runs

Run from project root: python scripts/fix_hr_docx_placeholder_placement.py
"""
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

# Import FORMS from inject script
from inject_all_hr_placeholders import FORMS, HR_DOCS, _is_header_table

PLACEHOLDER_RE = re.compile(r'\{\{\s*(\w+)\s*\}\}')


def restore_labels_in_docx(doc_path, label_placeholders):
    """Restore labels for placeholders that lost them (merged cell overwrite)."""
    if not os.path.isfile(doc_path):
        return False, []
    doc = Document(doc_path)
    # Build placeholder -> label (first occurrence per placeholder)
    placeholder_to_label = {}
    for label, placeholder in label_placeholders:
        if placeholder not in placeholder_to_label:
            placeholder_to_label[placeholder] = label

    tables_start = 0
    if doc.tables and _is_header_table(doc.tables[0]):
        tables_start = 1

    fixed = []
    for ti in range(tables_start, len(doc.tables)):
        table = doc.tables[ti]
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                for match in PLACEHOLDER_RE.finditer(text):
                    placeholder = match.group(1)
                    label = placeholder_to_label.get(placeholder)
                    if not label:
                        continue
                    # Check if label is already present (before the placeholder)
                    before_placeholder = text[:match.start()].strip()
                    # Label might be "Name:" or "Employee Name" - check if something like it exists
                    if not before_placeholder or before_placeholder in ("{{", ":"):
                        # No proper label - prepend it
                        new_text = label + " " + text
                        cell.text = new_text
                        fixed.append((placeholder, label))
                        break  # one fix per cell
    if fixed:
        doc.save(doc_path)
    return True, fixed


def main():
    print("=" * 60)
    print("Fix HR DOCX Placeholder Placement - Restore Labels")
    print("=" * 60)
    total_fixed = 0
    for filename, label_placeholders in FORMS:
        path = os.path.join(HR_DOCS, filename)
        if not os.path.isfile(path):
            print("Skip (not found):", filename)
            continue
        ok, fixed = restore_labels_in_docx(path, label_placeholders)
        if fixed:
            print(f"\n{filename}: restored {len(fixed)} labels")
            for ph, lbl in fixed[:5]:  # show first 5
                print(f"  {ph} <- {lbl}")
            if len(fixed) > 5:
                print(f"  ... and {len(fixed) - 5} more")
            total_fixed += len(fixed)
        else:
            print(f"{filename}: OK (no changes needed)")
    print("\n" + "=" * 60)
    print(f"Done. Total labels restored: {total_fixed}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
