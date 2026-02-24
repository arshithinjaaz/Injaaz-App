"""
Inject docxtpl placeholders into all HR form DOCX templates by matching label text
and setting the adjacent cell to {{ placeholder_name }}.

Run from project root: python scripts/inject_all_hr_placeholders.py
"""
import os
import sys
import shutil

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
HR_DOCS = os.path.join(BASE, "HR Documents")

# (filename, list of (label_substring, placeholder_name))
# Label is matched case-insensitive; the NEXT cell in the row gets the placeholder.
FORMS = [
    ("Leave Application Form - INJAAZ.DOCX", [
        ("Name:", "employee_name"), ("Job Title:", "job_title"), ("Employee ID:", "employee_id"),
        ("Date of Joining:", "date_of_joining"), ("Last Leave Date:", "last_leave_date"),
        ("Today's Date:", "today_date"), ("Department:", "department"), ("Mobile No.", "mobile_no"),
        ("Total No. of Days", "total_days_requested"), ("First Day of leave", "first_day_of_leave"),
        ("Last Day of Leave", "last_day_of_leave"), ("Date returning to work", "date_returning_to_work"),
        ("Telephone (reachable)", "telephone_reachable"), ("Replacement Name", "replacement_name"),
        ("leave_type", "leave_type"), ("number_of_days", "number_of_days"),
        ("Employee Signature", "employee_signature"), ("GM Signature", "gm_signature"),
        ("Checked by HR", "hr_checked"), ("HR Comments", "hr_comments"), ("Balance C/F", "hr_balance_cf"),
        ("HR Signature", "hr_signature"),
        ("Contract Year", "hr_contract_year"), ("Paid", "hr_paid"), ("Unpaid", "hr_unpaid"),
    ]),
    ("Duty Resumption Form - INJAAZ.DOCX", [
        ("Requester", "requester"), ("Employee Name", "employee_name"), ("Employee ID", "employee_id"),
        ("Job Title", "job_title"), ("Company", "company"), ("Leave Started", "leave_started"),
        ("Leave Ended", "leave_ended"), ("Planned Resumption", "planned_resumption_date"),
        ("Actual Resumption", "actual_resumption_date"), ("Note", "note"),
        ("Employee Signature", "employee_signature"), ("GM Signature", "gm_signature"),
        ("Line Manager Remarks", "line_manager_remarks"), ("sign_date", "sign_date"),
        ("HR Signature", "hr_signature"),
    ]),
    ("Passport Release & Submission Form - INJAAZ.DOCX", [
        ("Requester", "requester"), ("Employee Name", "employee_name"), ("Employee ID", "employee_id"),
        ("Job Title", "job_title"), ("Project", "project"), ("Date", "form_date"),
        ("Purpose of Release", "purpose_of_release"), ("Release Date", "release_date"),
        ("Employee Signature", "employee_signature"), ("GM Signature", "gm_signature"),
        ("HR Signature", "hr_signature"),
    ]),
    ("Employee grievance disciplinary action-form.docx", [
        ("Employee name:", "complainant_name"), ("Employee ID #:", "complainant_id"),
        ("Designation:", "complainant_designation"), ("Date of incident", "date_of_incident"),
        ("Shift / time:", "shift_time"), ("Employee contact #:", "complainant_contact"),
        ("Employee name:", "second_party_name"), ("Staff ID number:", "second_party_id"),
        ("Department:", "second_party_department"), ("Place of Incident:", "place_of_incident"),
        ("Employee contact #:", "second_party_contact"),
        ("complaint_description", "complaint_description"), ("Witnesses", "witnesses"),
        ("Who was informed", "who_informed"), ("Attachment", "attachment"),
        ("Signature of complainant", "complainant_signature"),
        ("HR signature", "hr_signature"), ("GM signature", "gm_signature"),
        ("hr_remarks", "hr_remarks"), ("gm_remarks", "gm_remarks"),
    ]),
    ("Interview Assessment Form - INJAAZ.DOCX", [
        ("Candidate Name", "candidate_name"), ("Position Title", "position_title"),
        ("Academic Qualification", "academic_qualification"), ("Age", "age"),
        ("Marital Status", "marital_status"), ("No. of Dependents", "dependents"),
        ("Nationality", "nationality"), ("Gender", "gender"), ("Current Job Title", "current_job_title"),
        ("Years of Experience", "years_experience"), ("Current Salary", "current_salary"),
        ("Expected Salary", "expected_salary"), ("Interview Date", "interview_date"),
        ("Interview by", "interview_by"), ("rating_turnout", "rating_turnout"),
        ("rating_confidence", "rating_confidence"), ("rating_mental_alertness", "rating_mental_alertness"),
        ("rating_maturity", "rating_maturity"), ("rating_communication", "rating_communication"),
        ("rating_technical", "rating_technical"), ("rating_training", "rating_training"),
        ("rating_experience", "rating_experience"), ("rating_overall", "rating_overall"),
        ("overall_assessment", "overall_assessment"), ("eligibility", "eligibility"),
        ("HR Signature", "hr_signature"), ("GM Signature", "gm_signature"),
        ("HR Comments", "hr_comments"), ("GM Comments", "gm_comments"),
        ("HR Remarks", "hr_remarks"), ("GM Remarks", "gm_remarks"),
    ]),
    ("Staff Appraisal Form - INJAAZ.DOCX", [
        ("Name:", "employee_name"), ("Employee ID:", "employee_id"), ("Department:", "department"),
        ("Position:", "position"), ("Appraisal Period", "appraisal_period"), ("Reviewer:", "reviewer"),
        ("rating_punctuality", "rating_punctuality"), ("comments_punctuality", "comments_punctuality"),
        ("rating_job_knowledge", "rating_job_knowledge"), ("comments_job_knowledge", "comments_job_knowledge"),
        ("rating_quality", "rating_quality"), ("comments_quality", "comments_quality"),
        ("rating_productivity", "rating_productivity"), ("comments_productivity", "comments_productivity"),
        ("rating_communication", "rating_communication"), ("comments_communication", "comments_communication"),
        ("rating_teamwork", "rating_teamwork"), ("comments_teamwork", "comments_teamwork"),
        ("rating_problem_solving", "rating_problem_solving"), ("comments_problem_solving", "comments_problem_solving"),
        ("rating_adaptability", "rating_adaptability"), ("comments_adaptability", "comments_adaptability"),
        ("rating_leadership", "rating_leadership"), ("comments_leadership", "comments_leadership"),
        ("Total Score:", "total_score"), ("employee_strengths", "employee_strengths"),
        ("Employee Signature:", "employee_signature"),
        ("HR Signature", "hr_signature"), ("GM Signature", "gm_signature"),
        ("HR Comments", "hr_comments"), ("GM Comments", "gm_comments"),
        ("HR Remarks", "hr_remarks"), ("GM Remarks", "gm_remarks"),
    ]),
    ("Station Clearance Form - INJAAZ.DOCX", [
        ("Employee Name", "employee_name"), ("Employee ID", "employee_id"),
        ("Employment Date", "employment_date"), ("Position", "position"),
        ("Department", "department"), ("Section", "section"),
        ("Type of Departure", "type_of_departure"), ("Last Working Date", "last_working_date"),
        ("tasks_handed_over", "tasks_handed_over"), ("documents_handed_over", "documents_handed_over"),
        ("files_handed_over", "files_handed_over"), ("keys_returned", "keys_returned"),
        ("toolbox_returned", "toolbox_returned"), ("access_card_returned", "access_card_returned"),
        ("dept_others", "dept_others"), ("email_cancelled", "email_cancelled"),
        ("software_hardware_returned", "software_hardware_returned"), ("laptop_returned", "laptop_returned"),
        ("it_others", "it_others"), ("file_shifted", "file_shifted"), ("dues_paid", "dues_paid"),
        ("medical_card_returned", "medical_card_returned"), ("hr_others", "hr_others"),
        ("eos_transfer", "eos_transfer"), ("finance_others", "finance_others"),
        ("Remarks", "remarks"), ("Employee Signature", "employee_signature"),
        ("HR Signature", "hr_signature"), ("Human Resources Manager", "hr_signature"),
    ]),
    ("Visa Renewal Form - INJAAZ.DOCX", [
        ("form_date", "form_date"), ("employee_name", "employee_name"), ("employee_id", "employee_id"),
        ("employer", "employer"), ("position", "position"), ("years_completed", "years_completed"),
        ("decision", "decision"), ("Signature of Employee", "employee_signature"),
        ("HR Signature", "hr_signature"), ("GM Signature", "gm_signature"),
        ("HR Comments", "hr_comments"), ("GM Comments", "gm_comments"),
        ("HR Remarks", "hr_remarks"), ("GM Remarks", "gm_remarks"),
    ]),
    ("Employee Contract Renewal Assessment Form Word.docx", [
        ("Employee Name:", "employee_name"), ("Date Of Evaluation:", "date_of_evaluation"),
        ("Employee Id:", "employee_id"), ("Date Of Joining:", "date_of_joining"),
        ("Department/Section:", "department"), ("Contract End Date:", "contract_end_date"),
        ("Evaluation By:", "evaluation_by"), ("Designation:", "designation"),
        ("rating_01", "rating_01"), ("rating_02", "rating_02"), ("rating_03", "rating_03"),
        ("rating_04", "rating_04"), ("strength", "strength"), ("areas_for_improvement", "areas_for_improvement"),
        ("OVERALL SCORE:", "overall_score"), ("recommendation", "recommendation"),
        ("Evaluator's Signature", "evaluator_signature"), ("evaluator_date", "evaluator_date"),
        ("HR Signature", "hr_signature"), ("GM Signature", "gm_signature"),
        ("HR Comments", "hr_comments"), ("GM Comments", "gm_comments"),
        ("HR Remarks", "hr_remarks"), ("GM Remarks", "gm_remarks"),
    ]),
]

# First table with 1 row and 2 cols is our standard header (headline + logo) - skip when injecting
def _is_header_table(table):
    return len(table.rows) == 1 and len(table.rows[0].cells) == 2


def inject_by_labels(doc_path, label_placeholders, backup=True):
    """Set placeholder in cell following a cell that contains label (substring match).
    Same label can map to different placeholders in order (e.g. first 'Employee name:' -> complainant_name, second -> second_party_name).
    """
    if not os.path.isfile(doc_path):
        return False
    if backup:
        backup_dir = os.path.join(os.path.dirname(doc_path), "templates")
        os.makedirs(backup_dir, exist_ok=True)
        name = os.path.basename(doc_path)
        backup_path = os.path.join(backup_dir, name.replace(".docx", " (before placeholders).docx").replace(".DOCX", " (before placeholders).docx"))
        shutil.copy2(doc_path, backup_path)
    doc = Document(doc_path)
    # Skip first table if it is our standard header (1 row x 2 cols) so we don't put placeholders in header
    tables_start = 0
    if doc.tables and _is_header_table(doc.tables[0]):
        tables_start = 1
    # Build list of (ti, ri, ci, label) for each cell that matches a label, in document order
    label_cells = []
    for ti in range(tables_start, len(doc.tables)):
        table = doc.tables[ti]
        for ri, row in enumerate(table.rows):
            cells = row.cells
            for ci in range(len(cells) - 1):
                text = (cells[ci].text or "").strip().lower()
                for label, _ in label_placeholders:
                    if label.lower() in text:
                        label_cells.append((ti, ri, ci, label))
                        break
    # For each (label, placeholder) find the next unused cell that matches this label
    used = set()
    for label, placeholder in label_placeholders:
        for ti, ri, ci, cell_label in label_cells:
            key = (ti, ri, ci)
            if key in used or cell_label != label:
                continue
            label_cell = doc.tables[ti].rows[ri].cells[ci]
            value_cell = doc.tables[ti].rows[ri].cells[ci + 1]
            placeholder_text = "{{ %s }}" % placeholder
            # Merged cells: value_cell may be same object as label_cell - preserve label
            if value_cell is label_cell:
                current = (label_cell.text or "").strip().replace(placeholder_text, "").strip()
                if placeholder_text not in (label_cell.text or ""):
                    label_cell.text = (current or label) + " " + placeholder_text
            else:
                value_cell.text = placeholder_text
            used.add(key)
            break
    doc.save(doc_path)
    return True


def main():
    for filename, label_placeholders in FORMS:
        path = os.path.join(HR_DOCS, filename)
        if not os.path.isfile(path):
            print("Skip (not found):", filename)
            continue
        print("Injecting:", filename)
        inject_by_labels(path, label_placeholders, backup=True)
    print("Done. Headers should already be applied (run apply_hr_header_all_forms.py first if you edited raw DOCX).")


if __name__ == "__main__":
    main()
