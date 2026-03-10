"""
Rebuild HR DOCX templates with proper Label | Value cell structure.
Restores from 'before header' originals, keeps the header table (form name + logo),
and rebuilds content tables with clear alignment.

Run from project root:  python scripts/rebuild_hr_templates.py
"""
import os
import sys
import shutil
import copy

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
HR_DOCS = os.path.join(BASE, "HR Documents")
TMPL_DIR = os.path.join(HR_DOCS, "templates")

FONT_NAME = "Calibri"
LABEL_SIZE = Pt(10)
VALUE_SIZE = Pt(10)
HEADER_SIZE = Pt(9)
SECTION_SIZE = Pt(10)


def _set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with values like {'sz': '4', 'color': '000000', 'val': 'single'}."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    for side, attrs in kwargs.items():
        el = OxmlElement(f'w:{side}')
        for k, v in attrs.items():
            el.set(qn(f'w:{k}'), v)
        tcBorders.append(el)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, color):
    """Set cell background colour."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 1cm = 567 twips
    tcW.set(qn('w:type'), 'dxa')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)


def _add_table_borders(tbl):
    """Add borders to table via XML."""
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def _label_cell(cell, text):
    """Format a cell as a label (bold, dark background)."""
    for p in cell.paragraphs:
        p.clear()
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = LABEL_SIZE
    _set_cell_shading(cell, 'F2F2F2')


def _value_cell(cell, placeholder):
    """Format a cell as a value placeholder."""
    for p in cell.paragraphs:
        p.clear()
    para = cell.paragraphs[0]
    run = para.add_run("{{ %s }}" % placeholder)
    run.font.name = FONT_NAME
    run.font.size = VALUE_SIZE


def _text_cell(cell, text, bold=False, size=None):
    for p in cell.paragraphs:
        p.clear()
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or VALUE_SIZE


def _new_table(doc, ncols, section_title=None):
    """Create a new table with optional section header row. Returns (table, start_adding_after_header)."""
    if section_title:
        tbl = doc.add_table(rows=1, cols=ncols)
        _add_table_borders(tbl)
        _text_cell(tbl.rows[0].cells[0], section_title, bold=True, size=SECTION_SIZE)
        _set_cell_shading(tbl.rows[0].cells[0], 'D9E2F3')
        if ncols > 1:
            tbl.rows[0].cells[0].merge(tbl.rows[0].cells[ncols - 1])
        return tbl
    tbl = doc.add_table(rows=0, cols=ncols)
    _add_table_borders(tbl)
    return tbl


def _section_row(tbl, text, ncols):
    """Add a section header row spanning all columns to an existing table."""
    row = tbl.add_row()
    _text_cell(row.cells[0], text, bold=True, size=SECTION_SIZE)
    _set_cell_shading(row.cells[0], 'D9E2F3')
    if ncols > 1:
        row.cells[0].merge(row.cells[ncols - 1])
    return row


def _label_value_row(tbl, pairs):
    """Add a row with Label | Value | Label | Value pattern.
    pairs = [(label, placeholder), ...]
    """
    row = tbl.add_row()
    ci = 0
    for label, placeholder in pairs:
        if ci < len(row.cells):
            _label_cell(row.cells[ci], label)
        if ci + 1 < len(row.cells):
            _value_cell(row.cells[ci + 1], placeholder)
        ci += 2
    return row


def _merge_row(tbl, text, ncols, bold=False, placeholder=None, shading=None):
    """Add a row where all cells are merged into one."""
    row = tbl.add_row()
    if placeholder:
        _value_cell(row.cells[0], placeholder)
    else:
        _text_cell(row.cells[0], text, bold=bold)
    if shading:
        _set_cell_shading(row.cells[0], shading)
    # Merge cells via XML
    if ncols > 1:
        first_tc = row.cells[0]._tc
        for ci in range(1, ncols):
            # Mark cells as part of horizontal merge
            pass
        # Use python-docx merge
        row.cells[0].merge(row.cells[ncols - 1])
    return row


def _get_header_table_xml(template_path):
    """Extract the header table (table 0) XML from an existing template."""
    doc = Document(template_path)
    if doc.tables:
        return copy.deepcopy(doc.tables[0]._tbl)
    return None


def _dest_path(filename):
    return os.path.join(HR_DOCS, filename)


def _backup(filename):
    dest = _dest_path(filename)
    if os.path.isfile(dest):
        bk = os.path.join(TMPL_DIR, filename.replace(".docx", " (before rebuild).docx").replace(".DOCX", " (before rebuild).docx"))
        os.makedirs(TMPL_DIR, exist_ok=True)
        shutil.copy2(dest, bk)


def _create_doc_with_header(filename):
    """Create a new doc, copying the header table from the current template."""
    current = _dest_path(filename)
    if os.path.isfile(current):
        doc = Document(current)
        body = doc.element.body
        header_tbl = doc.tables[0]._tbl if doc.tables else None
        sectPr = body.find(qn('w:sectPr'))
        keep = set()
        if header_tbl is not None:
            keep.add(id(header_tbl))
        if sectPr is not None:
            keep.add(id(sectPr))
        for el in list(body):
            if id(el) not in keep:
                body.remove(el)
        return doc
    doc = Document()
    return doc


def _add_footer_row(doc, doc_no, date_str, issue="00", revision="00"):
    """Add a standard INJAAZ footer table."""
    tbl = doc.add_table(rows=1, cols=4)
    _add_table_borders(tbl)
    tbl.rows[0].cells[0].text = f"DOC NO : {doc_no}"
    tbl.rows[0].cells[1].text = f"DATE : {date_str}"
    tbl.rows[0].cells[2].text = f"ISSUE : {issue}"
    tbl.rows[0].cells[3].text = f"REVISION : {revision}"
    for cell in tbl.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.name = FONT_NAME


# ─────────────────────────────────────────────────────────────────────────────
# Leave Application Form
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_leave():
    filename = "Leave Application Form - INJAAZ.DOCX"
    _backup(filename)
    doc = _create_doc_with_header(filename)
    doc.add_paragraph()  # spacer

    tbl = _new_table(doc, 4, "Employee Information")
    _label_value_row(tbl, [("Name:", "employee_name"), ("Today's Date:", "today_date")])
    _label_value_row(tbl, [("Job Title:", "job_title"), ("Department:", "department")])
    _label_value_row(tbl, [("Employee ID:", "employee_id"), ("Mobile No.:", "mobile_no")])
    _label_value_row(tbl, [("Date of Joining:", "date_of_joining"), ("Last Leave Date:", "last_leave_date")])

    doc.add_paragraph()

    tbl2 = _new_table(doc, 4, "Details of Leave")
    _label_value_row(tbl2, [("Leave Type:", "leave_type_display"), ("No. of Days:", "total_days_requested")])
    _label_value_row(tbl2, [("First Day of Leave:", "first_day_of_leave"), ("Last Day of Leave:", "last_day_of_leave")])
    _label_value_row(tbl2, [("Date Returning:", "date_returning_to_work"), ("Salary Advance:", "salary_advance")])
    _label_value_row(tbl2, [("Telephone (reachable):", "telephone_reachable"), ("Replacement Name:", "replacement_name")])

    doc.add_paragraph()

    tbl3 = _new_table(doc, 4, "Signatures")
    _label_value_row(tbl3, [("Employee Signature:", "employee_signature"), ("Manager Signature:", "gm_signature")])
    _label_value_row(tbl3, [("Replacement Signature:", "replacement_signature"), ("Date:", "today_date")])

    doc.add_paragraph()

    tbl4 = _new_table(doc, 4, "For Human Resources Use Only")
    _label_value_row(tbl4, [("Checked by HR:", "hr_checked"), ("HR Comments:", "hr_comments")])
    _label_value_row(tbl4, [("Balance C/F:", "hr_balance_cf"), ("Contract Year:", "hr_contract_year")])
    _label_value_row(tbl4, [("Paid:", "hr_paid"), ("Unpaid:", "hr_unpaid")])
    _label_value_row(tbl4, [("HR Signature:", "hr_signature"), ("Date:", "hr_date")])

    doc.add_paragraph()
    _add_footer_row(doc, "HR-FRM-007", "23/10/2024")

    doc.save(_dest_path(filename))
    print(f"  Rebuilt: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Duty Resumption Form
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_duty_resumption():
    filename = "Duty Resumption Form - INJAAZ.DOCX"
    _backup(filename)
    doc = _create_doc_with_header(filename)
    doc.add_paragraph()

    tbl = _new_table(doc, 4, "Employee Details")
    _label_value_row(tbl, [("Requester:", "requester"), ("Company:", "company")])
    _label_value_row(tbl, [("Employee Name:", "employee_name"), ("Employee ID:", "employee_id")])
    _label_value_row(tbl, [("Job Title:", "job_title"), ("", "placeholder_empty")])

    doc.add_paragraph()

    tbl2 = _new_table(doc, 4, "Leave & Resumption Information")
    _label_value_row(tbl2, [("Leave Started:", "leave_started"), ("Leave Ended:", "leave_ended")])
    _label_value_row(tbl2, [("Planned Resumption:", "planned_resumption_date"), ("Actual Resumption:", "actual_resumption_date")])

    # Note field (full width)
    row = tbl2.add_row()
    _label_cell(row.cells[0], "Note:")
    _value_cell(row.cells[1], "note")
    row.cells[1].merge(row.cells[3])

    doc.add_paragraph()

    tbl3 = _new_table(doc, 2)
    row = tbl3.add_row()
    _label_cell(row.cells[0], "Line Manager Remarks:")
    _value_cell(row.cells[1], "line_manager_remarks")

    doc.add_paragraph()

    tbl4 = _new_table(doc, 4, "Signatures")
    _label_value_row(tbl4, [("Employee Signature:", "employee_signature"), ("Date:", "sign_date")])
    _label_value_row(tbl4, [("Line Manager:", "gm_signature"), ("HR Signature:", "hr_signature")])

    doc.save(_dest_path(filename))
    print(f"  Rebuilt: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Passport Release & Submission Form
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_passport_release():
    filename = "Passport Release & Submission Form - INJAAZ.DOCX"
    _backup(filename)
    doc = _create_doc_with_header(filename)
    doc.add_paragraph()

    tbl = _new_table(doc, 4, "Request Details")
    _label_value_row(tbl, [("Requester:", "requester"), ("Date:", "form_date")])
    _label_value_row(tbl, [("Employee Name:", "employee_name"), ("Employee ID:", "employee_id")])
    _label_value_row(tbl, [("Job Title:", "job_title"), ("Project:", "project")])

    # Purpose and release date
    row = tbl.add_row()
    _label_cell(row.cells[0], "Purpose of Release:")
    _value_cell(row.cells[1], "purpose_of_release")
    row.cells[1].merge(row.cells[3])

    row2 = tbl.add_row()
    _label_cell(row2.cells[0], "Release Date:")
    _value_cell(row2.cells[1], "release_date")
    row2.cells[1].merge(row2.cells[3])

    doc.add_paragraph()

    tbl2 = _new_table(doc, 4, "Signatures")
    _label_value_row(tbl2, [("Employee Signature:", "employee_signature"), ("Date:", "form_date")])
    _label_value_row(tbl2, [("Line Manager:", "gm_signature"), ("HR Signature:", "hr_signature")])

    doc.add_paragraph()

    tbl3 = _new_table(doc, 4, "Requisition for Safekeeping")
    _label_value_row(tbl3, [("Employee Name:", "employee_name"), ("Employee ID:", "employee_id")])
    _label_value_row(tbl3, [("Job Title:", "job_title"), ("Project:", "project")])
    _label_value_row(tbl3, [("Date:", "form_date"), ("HR Signature:", "hr_signature")])

    doc.add_paragraph()
    _add_footer_row(doc, "HR-FRM-003", "09/05/2025", "01", "02")

    doc.save(_dest_path(filename))
    print(f"  Rebuilt: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Station Clearance Form
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_station_clearance():
    filename = "Station Clearance Form - INJAAZ.DOCX"
    _backup(filename)
    doc = _create_doc_with_header(filename)
    doc.add_paragraph()

    tbl = _new_table(doc, 4, "Employee Information")
    _label_value_row(tbl, [("Employee Name:", "employee_name"), ("Employee ID:", "employee_id")])
    _label_value_row(tbl, [("Employment Date:", "employment_date"), ("Position:", "position")])
    _label_value_row(tbl, [("Department:", "department"), ("Section:", "section")])
    _label_value_row(tbl, [("Type of Departure:", "type_of_departure"), ("Last Working Date:", "last_working_date")])

    doc.add_paragraph()

    tbl2 = _new_table(doc, 3)

    # Header row
    hdr = tbl2.add_row()
    _text_cell(hdr.cells[0], "Clearance Item", bold=True, size=Pt(9))
    _text_cell(hdr.cells[1], "Status", bold=True, size=Pt(9))
    _text_cell(hdr.cells[2], "Signature/Date", bold=True, size=Pt(9))
    for c in hdr.cells:
        _set_cell_shading(c, 'D9E2F3')

    # Employee Department
    _merge_row(tbl2, "EMPLOYEE'S DEPARTMENT / SECTION", 3, bold=True, shading='F2F2F2')
    dept_items = [
        ("Completed/handed over all tasks", "tasks_handed_over"),
        ("Handed over all original documents", "documents_handed_over"),
        ("Handed over all normal & electronic files", "files_handed_over"),
        ("Keys Returned", "keys_returned"),
        ("Toolbox Returned", "toolbox_returned"),
        ("Access Card Returned", "access_card_returned"),
        ("Others", "dept_others"),
    ]
    for label, ph in dept_items:
        row = tbl2.add_row()
        _text_cell(row.cells[0], label)
        _value_cell(row.cells[1], ph)
        row.cells[2].text = ""

    # IT Department
    _merge_row(tbl2, "INFORMATION TECHNOLOGY DEPARTMENT", 3, bold=True, shading='F2F2F2')
    it_items = [
        ("E-mail Account Cancelled", "email_cancelled"),
        ("Returned all software/hardware", "software_hardware_returned"),
        ("Laptop Returned", "laptop_returned"),
        ("Others", "it_others"),
    ]
    for label, ph in it_items:
        row = tbl2.add_row()
        _text_cell(row.cells[0], label)
        _value_cell(row.cells[1], ph)
        row.cells[2].text = ""

    # HR Department
    _merge_row(tbl2, "HUMAN RESOURCES & ADMINISTRATION", 3, bold=True, shading='F2F2F2')
    hr_items = [
        ("Employee file shifted to Exit folder", "file_shifted"),
        ("Payment of outstanding dues (Salary)", "dues_paid"),
        ("Medical Card Returned", "medical_card_returned"),
        ("Others", "hr_others"),
    ]
    for label, ph in hr_items:
        row = tbl2.add_row()
        _text_cell(row.cells[0], label)
        _value_cell(row.cells[1], ph)
        row.cells[2].text = ""

    # Finance Department
    _merge_row(tbl2, "FINANCE DEPARTMENT", 3, bold=True, shading='F2F2F2')
    fin_items = [
        ("EOS Benefits Transfer", "eos_transfer"),
        ("Others", "finance_others"),
    ]
    for label, ph in fin_items:
        row = tbl2.add_row()
        _text_cell(row.cells[0], label)
        _value_cell(row.cells[1], ph)
        row.cells[2].text = ""

    doc.add_paragraph()

    tbl3 = _new_table(doc, 2)
    row = tbl3.add_row()
    _label_cell(row.cells[0], "Remarks:")
    _value_cell(row.cells[1], "remarks")

    doc.add_paragraph()

    tbl4 = _new_table(doc, 4, "Signatures")
    _label_value_row(tbl4, [("Employee Signature:", "employee_signature"), ("HR Manager:", "hr_signature")])

    doc.save(_dest_path(filename))
    print(f"  Rebuilt: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Visa Renewal Form
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_visa_renewal():
    filename = "Visa Renewal Form - INJAAZ.DOCX"
    _backup(filename)
    doc = _create_doc_with_header(filename)
    doc.add_paragraph()

    tbl = _new_table(doc, 4, "Employee Information")
    _label_value_row(tbl, [("Employee Name:", "employee_name"), ("Employee ID:", "employee_id")])
    _label_value_row(tbl, [("Employer:", "employer"), ("Position:", "position")])
    _label_value_row(tbl, [("Years Completed:", "years_completed"), ("Date:", "form_date")])

    doc.add_paragraph()

    # Declaration
    para = doc.add_paragraph()
    run = para.add_run("To: HR Department,")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = VALUE_SIZE

    doc.add_paragraph()

    para2 = doc.add_paragraph()
    r = para2.add_run("I would like to confirm my decision to:")
    r.font.name = FONT_NAME
    r.font.size = VALUE_SIZE

    doc.add_paragraph()

    para3 = doc.add_paragraph()
    r = para3.add_run("1. Continue my employment with the company for the next 2 years and willing to have my visa renewed.")
    r.font.name = FONT_NAME
    r.font.size = VALUE_SIZE

    doc.add_paragraph()

    para4 = doc.add_paragraph()
    r = para4.add_run("2. Discontinue my service and require visa cancellation.")
    r.font.name = FONT_NAME
    r.font.size = VALUE_SIZE

    doc.add_paragraph()

    tbl2 = _new_table(doc, 2)
    row = tbl2.add_row()
    _label_cell(row.cells[0], "Decision:")
    _value_cell(row.cells[1], "decision_display")

    doc.add_paragraph()

    para5 = doc.add_paragraph()
    r = para5.add_run("Note: In case I do not continue my service, I will be liable to cover any additional charges covered by the company.")
    r.font.name = FONT_NAME
    r.font.size = Pt(9)
    r.italic = True

    doc.add_paragraph()

    tbl3 = _new_table(doc, 4, "Signatures")
    _label_value_row(tbl3, [("Employee Signature:", "employee_signature"), ("Date:", "form_date")])
    _label_value_row(tbl3, [("HR Signature:", "hr_signature"), ("GM Signature:", "gm_signature")])

    doc.add_paragraph()

    para6 = doc.add_paragraph()
    r = para6.add_run("INJAAZ LLC")
    r.bold = True
    r.font.name = FONT_NAME

    doc.save(_dest_path(filename))
    print(f"  Rebuilt: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Interview Assessment Form - fix the duplicated/misplaced rows
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_interview_assessment():
    filename = "Interview Assessment Form - INJAAZ.DOCX"
    _backup(filename)
    doc = _create_doc_with_header(filename)
    doc.add_paragraph()

    tbl = _new_table(doc, 4, "Candidate Information")
    _label_value_row(tbl, [("Candidate Name:", "candidate_name"), ("Interview Date:", "interview_date")])
    _label_value_row(tbl, [("Position Applied:", "position_title"), ("Interviewer:", "interviewer_name")])
    _label_value_row(tbl, [("Academic Qualification:", "academic_qualification"), ("Years of Experience:", "years_experience")])
    _label_value_row(tbl, [("Current Job Title:", "current_job_title"), ("Nationality:", "nationality")])

    doc.add_paragraph()

    tbl2 = _new_table(doc, 3)

    hdr = tbl2.add_row()
    _text_cell(hdr.cells[0], "Factor", bold=True, size=Pt(9))
    _text_cell(hdr.cells[1], "Description", bold=True, size=Pt(9))
    _text_cell(hdr.cells[2], "Rating", bold=True, size=Pt(9))
    for c in hdr.cells:
        _set_cell_shading(c, 'D9E2F3')

    ratings = [
        ("Turn-out & Appearance", "Appearance appropriate to the position", "rating_turnout"),
        ("Confidence", "Professional competence and self-confidence", "rating_confidence"),
        ("Mental Alertness", "Comprehends and responds coherently", "rating_mental_alertness"),
        ("Maturity & Stability", "Composure and balanced behaviour", "rating_maturity"),
        ("Communication Skills", "Listens well, expresses thoughts clearly", "rating_communication"),
        ("Technical Knowledge", "Awareness of duties and responsibilities", "rating_technical"),
        ("Relevant Training", "Professional/technical training for the role", "rating_training"),
        ("Relevant Experience", "Previous work experience meets requirement", "rating_experience"),
        ("Overall Rating", "Overall suitability for the position", "rating_overall"),
    ]
    for factor, desc, ph in ratings:
        row = tbl2.add_row()
        _text_cell(row.cells[0], factor, bold=True, size=Pt(9))
        _text_cell(row.cells[1], desc, size=Pt(9))
        _value_cell(row.cells[2], ph)

    doc.add_paragraph()

    tbl3 = _new_table(doc, 2, "Overall Assessment")
    row = tbl3.add_row()
    _label_cell(row.cells[0], "Assessment / Comments:")
    _value_cell(row.cells[1], "overall_assessment")
    row2 = tbl3.add_row()
    _label_cell(row2.cells[0], "Eligible for Hire?")
    _value_cell(row2.cells[1], "eligibility")

    doc.add_paragraph()

    tbl4 = _new_table(doc, 4, "Signatures")
    _label_value_row(tbl4, [("Interviewer:", "interviewer_signature"), ("HR Signature:", "hr_signature")])
    _label_value_row(tbl4, [("GM Signature:", "gm_signature"), ("HR Remarks:", "hr_remarks")])

    doc.save(_dest_path(filename))
    print(f"  Rebuilt: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Grievance Form
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_grievance():
    filename = "Employee grievance disciplinary action-form.docx"
    _backup(filename)
    doc = _create_doc_with_header(filename)
    doc.add_paragraph()

    tbl = _new_table(doc, 4, "First Party (Complainant)")
    _label_value_row(tbl, [("Employee Name:", "complainant_name"), ("Employee ID:", "complainant_id")])
    _label_value_row(tbl, [("Designation:", "complainant_designation"), ("Date of Incident:", "date_of_incident")])
    _label_value_row(tbl, [("Shift / Time:", "shift_time"), ("Contact No.:", "complainant_contact")])

    doc.add_paragraph()

    tbl2 = _new_table(doc, 4, "Second Party")
    _label_value_row(tbl2, [("Employee Name:", "second_party_name"), ("Staff ID:", "second_party_id")])
    _label_value_row(tbl2, [("Department:", "second_party_department"), ("Place of Incident:", "place_of_incident")])
    row = tbl2.add_row()
    _label_cell(row.cells[0], "Contact No.:")
    _value_cell(row.cells[1], "second_party_contact")
    row.cells[1].merge(row.cells[3])

    doc.add_paragraph()

    tbl3 = _new_table(doc, 2, "Complaint / Issue Details")
    row = tbl3.add_row()
    _label_cell(row.cells[0], "Description:")
    _value_cell(row.cells[1], "complaint_description")

    row = tbl3.add_row()
    _label_cell(row.cells[0], "Witnesses:")
    _value_cell(row.cells[1], "witnesses")

    row = tbl3.add_row()
    _label_cell(row.cells[0], "Who was informed?")
    _value_cell(row.cells[1], "who_informed")

    row = tbl3.add_row()
    _label_cell(row.cells[0], "Attachment:")
    _value_cell(row.cells[1], "attachment")

    doc.add_paragraph()

    tbl4 = _new_table(doc, 4, "HR & Management Review")
    _label_value_row(tbl4, [("HR Remarks:", "hr_remarks"), ("GM Remarks:", "gm_remarks")])

    doc.add_paragraph()

    tbl5 = _new_table(doc, 4, "Signatures")
    _label_value_row(tbl5, [("Complainant:", "complainant_signature"), ("HR Signature:", "hr_signature")])
    _label_value_row(tbl5, [("GM Signature:", "gm_signature"), ("Date:", "form_date")])

    doc.save(_dest_path(filename))
    print(f"  Rebuilt: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Contract Renewal - fix the R26-R32 issues
# ─────────────────────────────────────────────────────────────────────────────
def rebuild_contract_renewal():
    filename = "Employee Contract Renewal Assessment Form Word.docx"
    _backup(filename)

    # Keep existing doc (tables 0 and 1 are good), fix table 2
    current = _dest_path(filename)
    doc = Document(current)

    # Table 2 (rating table) - rows 26-32 have issues
    # Fix R26: strength label
    t2 = doc.tables[2]
    try:
        r26 = t2.rows[26]
        # Set cell 1 properly
        r26.cells[1].paragraphs[0].clear()
        run = r26.cells[1].paragraphs[0].add_run("Strength:")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        # Cell 2 should have the value
        for ci in [2, 3]:
            r26.cells[ci].paragraphs[0].clear()
            r26.cells[ci].paragraphs[0].add_run("{{ strength }}").font.size = Pt(9)

        # R27: Areas for improvement
        r27 = t2.rows[27]
        r27.cells[1].paragraphs[0].clear()
        run = r27.cells[1].paragraphs[0].add_run("Areas for Improvement:")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        for ci in [2, 3, 4]:
            if ci < len(r27.cells):
                r27.cells[ci].paragraphs[0].clear()
                r27.cells[ci].paragraphs[0].add_run("{{ areas_for_improvement }}").font.size = Pt(9)

        # R28: Overall score - fix label
        r28 = t2.rows[28]
        r28.cells[0].paragraphs[0].clear()
        run = r28.cells[0].paragraphs[0].add_run("OVERALL SCORE:")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(9)

        # R29-31: Recommendation - fix
        for ri in [29, 30, 31]:
            row = t2.rows[ri]
            row.cells[0].paragraphs[0].clear()
            if ri == 29:
                row.cells[0].paragraphs[0].add_run("Recommendation: {{ recommendation }}").font.size = Pt(9)
            else:
                row.cells[0].paragraphs[0].clear()

        # R32: Evaluator signature
        r32 = t2.rows[32]
        r32.cells[0].paragraphs[0].clear()
        run = r32.cells[0].paragraphs[0].add_run("Evaluator Signature: {{ evaluator_signature }}   Date: {{ evaluator_date }}")
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
    except (IndexError, AttributeError) as e:
        print(f"  Warning: Contract renewal fix incomplete: {e}")

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────
def main():
    os.makedirs(TMPL_DIR, exist_ok=True)
    print("Rebuilding HR DOCX templates with proper alignment...")
    rebuild_leave()
    rebuild_duty_resumption()
    rebuild_passport_release()
    rebuild_grievance()
    rebuild_interview_assessment()
    rebuild_station_clearance()
    rebuild_visa_renewal()
    rebuild_contract_renewal()
    print("\nDone. Run: python scripts/auto_test_hr_forms.py")


if __name__ == "__main__":
    main()
