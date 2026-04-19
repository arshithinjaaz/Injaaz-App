"""Print paragraph structure of HR DOCX templates."""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

HR_DOCS = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "HR Documents")

TARGET = sys.argv[1] if len(sys.argv) > 1 else None

FORMS = [
    "Staff Appraisal Form - INJAAZ.DOCX",
    "Visa Renewal Form - INJAAZ.DOCX",
]

for fname in FORMS:
    if TARGET and TARGET.lower() not in fname.lower():
        continue
    path = os.path.join(HR_DOCS, fname)
    if not os.path.exists(path):
        print(f"MISSING: {fname}"); continue
    doc = Document(path)
    print(f"\n{'='*70}\nFILE: {fname}  ({len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs)\n{'='*70}")
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t:
            print(f"  P{i}: {repr(t[:120])}")
