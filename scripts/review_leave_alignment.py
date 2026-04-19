#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Deep alignment review: MAIN vs generated leave DOCX, leave-type block rows."""
import sys, io, glob, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn

MAIN = r'd:\Work\Injaaz-App\HR Documents - Main\Leave Application Form - INJAAZ.DOCX'
OUT_DIR = r'd:\Work\Injaaz-App\test_output\leave_bold_fix'
gen_files = glob.glob(os.path.join(OUT_DIR, 'leave_application_*.docx'))
GEN = gen_files[0] if gen_files else None


def cell_margin(cell, side):
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        return None
    el = tcMar.find(qn(f'w:{side}'))
    return el.get(qn('w:w')) if el is not None else None


def inspect_leave_block(label, path):
    d = Document(path)
    t = d.tables[0]
    print(f'\n{"="*60}')
    print(f'  {label}  ({path})')
    print(f'{"="*60}')
    print(f'  {"Row":<5} {"sb":>7} {"ind":>7} {"bold":>5} {"sz":>5} {"cL":>5} {"cR":>5}  runs')
    print(f'  {"-"*110}')
    for ri in range(5, 22):
        row = t.rows[ri]
        cell = row.cells[0]
        para = cell.paragraphs[0] if cell.paragraphs else None
        if not para:
            continue
        pf = para.paragraph_format
        runs = para.runs
        sb  = pf.space_before.pt if pf.space_before else None
        ind = round(pf.left_indent.pt, 1) if pf.left_indent else None
        bold = runs[0].bold if runs else None
        sz   = runs[0].font.size.pt if (runs and runs[0].font.size) else None
        cL = cell_margin(cell, 'start') or cell_margin(cell, 'left')
        cR = cell_margin(cell, 'end')   or cell_margin(cell, 'right')
        run_texts = [r.text.replace('\u2713', 'X').replace('\u2705', 'OK') for r in runs]
        print(f'  r{ri:<4} {str(sb):>7} {str(ind):>7} {str(bold):>5} {str(sz):>5} {str(cL):>5} {str(cR):>5}  {run_texts}')


inspect_leave_block('MAIN', MAIN)
if GEN:
    inspect_leave_block('GEN', GEN)
    # Side-by-side diff for leave-type rows (r7–r15)
    dm = Document(MAIN); dg = Document(GEN)
    tm = dm.tables[0]; tg = dg.tables[0]
    print(f'\n{"="*60}')
    print('  DIFF (leave-type rows r7–r15)')
    print(f'{"="*60}')
    for ri in range(7, 16):
        pm = tm.rows[ri].cells[0].paragraphs[0]
        pg = tg.rows[ri].cells[0].paragraphs[0]
        sb_m = pm.paragraph_format.space_before.pt if pm.paragraph_format.space_before else None
        sb_g = pg.paragraph_format.space_before.pt if pg.paragraph_format.space_before else None
        runs_m = [r.text.replace('\u2713','X') for r in pm.runs]
        runs_g = [r.text.replace('\u2713','X') for r in pg.runs]
        sb_ok  = '  OK' if sb_m == sb_g else f'  DIFF sb: MAIN={sb_m} GEN={sb_g}'
        run_ok = '  runs OK' if runs_m == runs_g else f'  runs DIFF: MAIN={runs_m}  GEN={runs_g}'
        print(f'  r{ri}: {sb_ok}{run_ok}')
else:
    print('ERROR: no generated leave_application file found in', OUT_DIR)
