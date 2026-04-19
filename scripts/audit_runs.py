"""Check paragraph runs inside cells to see if labels exist as separate runs."""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

TMPL = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "HR Documents", "templates")
TARGET = sys.argv[1] if len(sys.argv) > 1 else None

for fname in sorted(os.listdir(TMPL)):
    if "before header" not in fname:
        continue
    if TARGET and TARGET.lower() not in fname.lower():
        continue
    path = os.path.join(TMPL, fname)
    doc = Document(path)
    print(f"\n{'='*80}\n  {fname}\n{'='*80}")
    for ti, tbl in enumerate(doc.tables):
        print(f"\n  TABLE {ti}: {len(tbl.rows)}R x {len(tbl.columns)}C")
        for ri, row in enumerate(tbl.rows[:10]):
            cells = row.cells
            seen_tc = set()
            for ci, cell in enumerate(cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                for pi, para in enumerate(cell.paragraphs):
                    runs = [(r.text, r.bold, r.font.size) for r in para.runs if r.text.strip()]
                    if runs:
                        print(f"    R{ri} C{ci} P{pi}: {runs}")
