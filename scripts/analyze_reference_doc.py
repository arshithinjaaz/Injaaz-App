"""Analyze the reference Commencement Form document to understand its format."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

def analyze_document(doc_path):
    """Analyze document structure, spacing, and formatting."""
    doc = Document(doc_path)
    
    print("=" * 60)
    print(f"Analyzing: {os.path.basename(doc_path)}")
    print("=" * 60)
    
    # Check tables
    print(f"\nTables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
        if i == 0:  # First table (likely header)
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    print(f"\n    Row {j}, Cell {k}:")
                    # Get all text from cell
                    cell_text = cell.text.strip()
                    print(f"      Text: '{cell_text[:100]}'")
                    
                    # Check for images
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run._element.findall('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                                print(f"      Contains IMAGE/LOGO")
                    
                    # Check cell margins
                    tc = cell._tc
                    tcPr = tc.tcPr
                    if tcPr is not None:
                        tcMar = tcPr.find(qn('w:tcMar'))
                        if tcMar is not None:
                            margins = {}
                            for m in ['top', 'bottom', 'start', 'end']:
                                elem = tcMar.find(qn(f'w:{m}'))
                                if elem is not None:
                                    margins[m] = elem.get(qn('w:w'))
                            if margins:
                                print(f"      Margins: {margins}")
                    
                    # Check all paragraphs in cell
                    print(f"      Paragraphs in cell: {len(cell.paragraphs)}")
                    for p_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()[:60]
                        pf = para.paragraph_format
                        spacing_info = f"before={pf.space_before}, after={pf.space_after}"
                        if pf.line_spacing:
                            spacing_info += f", line={pf.line_spacing}"
                        print(f"        Para {p_idx}: '{text}'")
                        print(f"          Spacing: {spacing_info}")
                        if para.runs:
                            first_run = para.runs[0]
                            if first_run.font.size:
                                print(f"          Font size: {first_run.font.size}")
                            if first_run.bold:
                                print(f"          Bold: {first_run.bold}")
                    
                    # Check cell width
                    try:
                        width = cell.width
                        print(f"      Cell width: {width}")
                    except:
                        pass
    
    # Check first 5 paragraphs after table
    print(f"\nFirst 5 paragraphs after table:")
    for i, para in enumerate(doc.paragraphs[:5]):
        text = para.text.strip()[:60]
        pf = para.paragraph_format
        spacing_info = f"before={pf.space_before}, after={pf.space_after}"
        print(f"  Para {i}: '{text}'")
        print(f"    Spacing: {spacing_info}")

if __name__ == "__main__":
    ref_path = r"c:\Users\Lenovo\Downloads\Commencement_Form_HR-COMMENCEMENT-A05D1664 (11).docx"
    if os.path.exists(ref_path):
        analyze_document(ref_path)
    else:
        print(f"File not found: {ref_path}")
