#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Compare leave template rows: MAIN vs generated."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn

MAIN = r'd:\Work\Injaaz-App\HR Documents - Main\Leave Application Form - INJAAZ.DOCX'
GEN  = r'd:\Work\Injaaz-App\test_output\leave_leftpad_fix\leave_application_HR-LEAVE-APPLICATION-39B9EB1D.docx'


def inspect_doc(label, path):
    d = Document(path)
    t = d.tables[0]
    print(f'\n=== {label} === rows={len(t.rows)}')
    for ri in range(5, 20):
        row = t.rows[ri]
        cell = row.cells[0]
        para = cell.paragraphs[0] if cell.paragraphs else None
        pf = para.paragraph_format if para else None
        runs = [r.text for r in para.runs] if para else []
        tcPr = cell._tc.find(qn('w:tcPr'))
        tcMar = tcPr.find(qn('w:tcMar')) if tcPr is not None else None
        margins = {}
        if tcMar is not None:
            for side in ('top', 'start', 'bottom', 'end', 'left', 'right'):
                el = tcMar.find(qn(f'w:{side}'))
                if el is not None:
                    margins[side] = el.get(qn('w:w'))
        sb = str(pf.space_before) if (pf and pf.space_before) else 'None'
        ind = str(pf.left_indent) if (pf and pf.left_indent) else 'None'
        safe_runs = [r.replace('\u2713', 'X').replace('\u2705', 'OK') for r in runs]
        print(f'  r{ri}: sb={sb}  ind={ind}  cellMar={margins}  runs={safe_runs}')


inspect_doc('MAIN', MAIN)
inspect_doc('GEN',  GEN)
