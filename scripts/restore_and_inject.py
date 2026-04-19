"""
Restore HR DOCX templates from 'before placeholders' backups (original design intact)
and inject {{ placeholder }} tags into VALUE cells only — preserving all original
formatting, headers, logos, merged cells, and labels exactly as they are.

Run from project root:  python scripts/restore_and_inject.py
"""
import os
import sys
import shutil
import re

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
HR_DOCS = os.path.join(BASE, "HR Documents")
TMPL_DIR = os.path.join(HR_DOCS, "templates")


def _dest(filename):
    return os.path.join(HR_DOCS, filename)


def _backup_src(filename):
    """Return the 'before placeholders' backup path (has header + original layout)."""
    name = (filename
            .replace(".docx", " (before placeholders).docx")
            .replace(".DOCX", " (before placeholders).docx"))
    return os.path.join(TMPL_DIR, name)


def _restore(filename):
    """Copy the backup over the current file, restoring original layout."""
    src = _backup_src(filename)
    if not os.path.isfile(src):
        print(f"  WARNING: No backup for {filename}")
        return False
    shutil.copy2(src, _dest(filename))
    return True


def _set_cell_value(cell, text):
    """Set cell text content while preserving formatting of the first run."""
    para = cell.paragraphs[0]
    if para.runs:
        # Clear all runs, set text in first
        for run in para.runs:
            run.text = ""
        para.runs[0].text = text
    else:
        para.text = text


def _inject_into_value_cell(table, row_idx, col_idx, placeholder):
    """Put {{ placeholder }} into a specific cell position — only if the cell exists and is empty or has old placeholder."""
    try:
        cell = table.rows[row_idx].cells[col_idx]
        txt = (cell.text or "").strip()
        ph = "{{ %s }}" % placeholder
        if ph in txt:
            return  # already present
        if not txt or txt.startswith("{{"):
            _set_cell_value(cell, ph)
    except IndexError:
        pass


def _inject_after_label(table, row_idx, col_idx, placeholder):
    """Append {{ placeholder }} to a cell that contains a label (for merged cells like 'Label: {{ value }}')."""
    try:
        cell = table.rows[row_idx].cells[col_idx]
        txt = (cell.text or "").strip()
        ph = "{{ %s }}" % placeholder
        if ph in txt:
            return
        if txt and not txt.startswith("{{"):
            _set_cell_value(cell, txt + " " + ph)
        elif not txt:
            _set_cell_value(cell, ph)
    except IndexError:
        pass


def _inject_into_paragraph(doc, para_idx, replacements):
    """Replace blanks in a paragraph with placeholders. replacements = [(search, replacement), ...]."""
    try:
        para = doc.paragraphs[para_idx]
        txt = para.text or ""
        new_txt = txt
        for search, repl in replacements:
            new_txt = new_txt.replace(search, repl, 1)
        if new_txt != txt:
            # Set via first run to preserve formatting
            if para.runs:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = new_txt
            else:
                para.text = new_txt
    except IndexError:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# Leave Application Form — original has merged cells, inject into correct positions
# ─────────────────────────────────────────────────────────────────────────────
def fix_leave():
    filename = "Leave Application Form - INJAAZ.DOCX"
    _restore(filename)
    doc = Document(_dest(filename))
    t = doc.tables[1]  # main content (29 rows x 4 cols)

    # R0: merged C0-C3 has '{{ employee_name }}' — already has placeholder, OK
    # R1: C0-C2 merged '{{ job_title }}', C3 'Today's Date:' → inject today_date after label
    _inject_after_label(t, 1, 3, "today_date")
    # R2: C0-C2 '{{ employee_id }}', C3 'Department:' → inject department
    _inject_after_label(t, 2, 3, "department")
    # R3: C0-C2 '{{ date_of_joining }}', C3 'Mobile No.:' → inject mobile_no
    _inject_after_label(t, 3, 3, "mobile_no")

    # R17: C0 'First Day of leave:', C1-C3 merged → fix: change placeholder to first_day_of_leave
    try:
        cell17 = t.rows[17].cells[1]
        txt17 = (cell17.text or "").strip()
        if "last_day_of_leave" in txt17 or "first_day_of_leave" in txt17:
            _set_cell_value(cell17, "{{ first_day_of_leave }}")
    except IndexError:
        pass

    # R20: 'Telephone Number...' merged → inject telephone_reachable
    _inject_after_label(t, 20, 0, "telephone_reachable")
    # R21: C0-C1 'Replacement Name:' merged, C2-C3 'Signature:' → inject names
    _inject_after_label(t, 21, 0, "replacement_name")
    _inject_after_label(t, 21, 2, "replacement_signature")
    # R22: C0-C1 'Employee Signature:', C2-C3 'Manager Signature:' → inject sigs
    _inject_after_label(t, 22, 0, "employee_signature")
    _inject_after_label(t, 22, 2, "gm_signature")
    # R28: C0-C2 'HR Signature:', C3 'Date:' → inject
    _inject_after_label(t, 28, 0, "hr_signature")

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Duty Resumption Form
# ─────────────────────────────────────────────────────────────────────────────
def fix_duty_resumption():
    filename = "Duty Resumption Form - INJAAZ.DOCX"
    _restore(filename)
    doc = Document(_dest(filename))

    t1 = doc.tables[1]  # 4R x 2C
    # R0: merged '{{ requester }}' → keep (label "Requester" is part of Word visual design)
    # R1: merged '{{ employee_name }}' → keep
    # R2: C0 'Employee ID:', C1 '{{ job_title }}' → FIX: should be employee_id
    _inject_into_value_cell(t1, 2, 1, "employee_id")
    # R3: C0 'Job Title:', C1 '{{ job_title }}' → keep

    t2 = doc.tables[2]  # 6R x 2C
    # R2: merged '{{ planned_resumption_date }}' → keep
    # R3: merged '{{ actual_resumption_date }}' → keep
    # R4: merged '{{ note }}' → keep
    # R5: C0 'Employee Signature:', C1 'Date:' → inject
    _inject_after_label(t2, 5, 0, "employee_signature")
    _inject_after_label(t2, 5, 1, "sign_date")

    t3 = doc.tables[3]  # 2R x 1C (Line Manager Remarks)
    # R1: blank line → inject
    _inject_into_value_cell(t3, 1, 0, "line_manager_remarks")

    t4 = doc.tables[4]  # 2R x 2C
    # R0: 'Approved by Line Manager:' | 'Date:' → inject gm_signature
    _inject_after_label(t4, 0, 0, "gm_signature")
    # R1: 'HR Signature:' | 'Date:' → inject hr_signature
    _inject_after_label(t4, 1, 0, "hr_signature")
    _inject_after_label(t4, 1, 1, "hr_date")

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Passport Release & Submission Form
# ─────────────────────────────────────────────────────────────────────────────
def fix_passport_release():
    filename = "Passport Release & Submission Form - INJAAZ.DOCX"
    _restore(filename)
    doc = Document(_dest(filename))

    t2 = doc.tables[2]  # 4R x 1C (purpose, release, signature, note)
    _inject_after_label(t2, 0, 0, "purpose_of_release")  # Purpose of Release:
    _inject_after_label(t2, 1, 0, "release_date")         # Release Date:
    _inject_after_label(t2, 2, 0, "employee_signature")   # Employee Signature:

    t3 = doc.tables[3]  # 2R x 2C (approvals)
    _inject_after_label(t3, 0, 0, "gm_signature")  # Approved by Line Manager:
    _inject_after_label(t3, 1, 0, "hr_signature")  # HR Signature:

    t4 = doc.tables[4]  # 6R x 4C (safekeeping section)
    # R1: 'Employee Name:' merged → inject
    _inject_after_label(t4, 1, 0, "employee_name")
    # R2: C0 'Employee ID:', C1 '{{ project }}' → fix project placement, add employee_id
    try:
        # project was incorrectly placed in Employee ID row
        r2c1 = t4.rows[2].cells[1]
        if "project" in (r2c1.text or ""):
            _set_cell_value(r2c1, "{{ employee_id }}")
    except IndexError:
        pass
    # R3: C0 'Job Title:', C1 '{{ form_date }}' → fix
    try:
        r3c1 = t4.rows[3].cells[1]
        if "form_date" in (r3c1.text or ""):
            _set_cell_value(r3c1, "{{ job_title }}")
    except IndexError:
        pass
    # R4: 'Approved by Line Manager:' → inject
    _inject_after_label(t4, 4, 1, "gm_signature")
    _inject_after_label(t4, 4, 2, "form_date")
    # R5: 'HR Signature:' → inject
    _inject_after_label(t4, 5, 1, "hr_signature")

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Employee Grievance Form
# ─────────────────────────────────────────────────────────────────────────────
def fix_grievance():
    filename = "Employee grievance disciplinary action-form.docx"
    _restore(filename)
    doc = Document(_dest(filename))

    # Tables 1 & 2 already have correct label | value structure from original
    # Table 3: complaint description (2R x 1C)
    t3 = doc.tables[3]
    try:
        txt = (t3.rows[0].cells[0].text or "").strip()
        if "complaint_description" not in txt:
            _inject_after_label(t3, 0, 0, "complaint_description")
    except IndexError:
        pass

    # Table 4: witnesses etc. - R0 witnesses label is merged
    t4 = doc.tables[4]
    # Check if witnesses placeholder exists
    try:
        r0txt = (t4.rows[0].cells[0].text or "").strip()
        if "witnesses" not in r0txt:
            _inject_after_label(t4, 0, 0, "witnesses")
    except IndexError:
        pass

    # Table 5: Employee Acknowledgment
    t5 = doc.tables[5]
    try:
        txt = (t5.rows[0].cells[0].text or "").strip()
        if "complainant_signature" not in txt:
            _inject_after_label(t5, 0, 0, "complainant_signature")
    except IndexError:
        pass

    # Inject HR/GM signatures into paragraphs at the end
    hr_sig_found = any(
        "{{ hr_signature }}" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not hr_sig_found:
        # Find "HR signature:" paragraph and inject after it
        for i, p in enumerate(doc.paragraphs):
            txt = (p.text or "").strip()
            if txt.lower().startswith("hr signature"):
                _inject_into_paragraph(doc, i, [("HR signature:", "HR signature: {{ hr_signature }}")])
            elif txt.lower().startswith("gm signature"):
                _inject_into_paragraph(doc, i, [("GM Signature:", "GM Signature: {{ gm_signature }}")])
            elif "HR remarks" in txt.lower() and "{{ hr_remarks }}" not in txt:
                # Add hr_remarks after the remarks heading
                pass

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Interview Assessment Form
# ─────────────────────────────────────────────────────────────────────────────
def fix_interview():
    filename = "Interview Assessment Form - INJAAZ.DOCX"
    _restore(filename)
    doc = Document(_dest(filename))

    t1 = doc.tables[1]  # 16R x 7C (rating table)

    # Rating rows: R3-R11 have Factor | Description | 5 rating columns (Outstanding to Low)
    # Inject rating placeholders into the appropriate column
    # We'll put the rating value in column 6 (last col = "Low")
    rating_map = {
        3: "rating_turnout", 4: "rating_confidence", 5: "rating_mental_alertness",
        6: "rating_maturity", 7: "rating_communication", 8: "rating_technical",
        9: "rating_training", 10: "rating_experience", 11: "rating_overall",
    }
    for ri, ph in rating_map.items():
        try:
            cell = t1.rows[ri].cells[6]
            if ("{{ %s }}" % ph) not in (cell.text or ""):
                _set_cell_value(cell, "{{ %s }}" % ph)
        except IndexError:
            pass

    # R13: 'Overall Assessment / Comments' merged → inject overall_assessment
    try:
        r13 = t1.rows[13]
        txt13 = (r13.cells[0].text or "").strip()
        if "overall_assessment" not in txt13 and "Overall Assessment" in txt13:
            _inject_after_label(t1, 13, 0, "overall_assessment")
    except IndexError:
        pass

    # R14: was academic_qualification (wrong position) → clear and use for candidate info
    # R15: was eligibility → keep for eligibility
    try:
        r14 = t1.rows[14]
        txt14 = (r14.cells[0].text or "").strip()
        if "academic_qualification" in txt14:
            _set_cell_value(r14.cells[0], "Academic Qualification: {{ academic_qualification }}")
        r15 = t1.rows[15]
        txt15 = (r15.cells[0].text or "").strip()
        if "eligibility" in txt15:
            _set_cell_value(r15.cells[0], "Eligible for hire: {{ eligibility }}")
    except IndexError:
        pass

    # We need candidate info fields - add them before the rating table
    # Check if candidate_name exists anywhere
    has_candidate = any(
        "candidate_name" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not has_candidate:
        from docx.oxml import OxmlElement
        tbl = doc.add_table(rows=4, cols=4)
        # Add borders
        tbl_xml = tbl._tbl
        tblPr = tbl_xml.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl_xml.insert(0, tblPr)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        tblPr.append(borders)

        rows_data = [
            ("Candidate Name:", "candidate_name", "Interview Date:", "interview_date"),
            ("Position Applied:", "position_title", "Interviewer:", "interviewer_name"),
            ("Qualification:", "academic_qualification", "Experience:", "years_experience"),
            ("Current Title:", "current_job_title", "Nationality:", "nationality"),
        ]
        for ri, (l1, ph1, l2, ph2) in enumerate(rows_data):
            tbl.rows[ri].cells[0].text = l1
            tbl.rows[ri].cells[1].text = "{{ %s }}" % ph1
            tbl.rows[ri].cells[2].text = l2
            tbl.rows[ri].cells[3].text = "{{ %s }}" % ph2

        # Move after header table (table 0)
        body = doc.element.body
        header_tbl = doc.tables[0]._tbl
        header_tbl.addnext(tbl._tbl)

    # Add signature table at end if missing
    has_hr = any("{{ hr_signature }}" in (c.text or "") for tbl in doc.tables for row in tbl.rows for c in row.cells)
    if not has_hr:
        sig = doc.add_table(rows=2, cols=4)
        sig_xml = sig._tbl
        sigPr = sig_xml.find(qn('w:tblPr'))
        if sigPr is None:
            sigPr = OxmlElement('w:tblPr')
            sig_xml.insert(0, sigPr)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        sigPr.append(borders)
        sig.rows[0].cells[0].text = "Interviewer Signature:"
        sig.rows[0].cells[1].text = "{{ interviewer_signature }}"
        sig.rows[0].cells[2].text = "HR Signature:"
        sig.rows[0].cells[3].text = "{{ hr_signature }}"
        sig.rows[1].cells[0].text = "GM Signature:"
        sig.rows[1].cells[1].text = "{{ gm_signature }}"
        sig.rows[1].cells[2].text = "HR Remarks:"
        sig.rows[1].cells[3].text = "{{ hr_remarks }}"

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Station Clearance Form
# ─────────────────────────────────────────────────────────────────────────────
def fix_station_clearance():
    filename = "Station Clearance Form - INJAAZ.DOCX"
    _restore(filename)
    doc = Document(_dest(filename))

    t = doc.tables[1]  # 28R x 6C

    # R0-R3: left side merged cells have {{ placeholder }}, right side C4=label, C5=value
    # R3 C5: 'Last Working Date' value is empty → inject
    _inject_into_value_cell(t, 3, 5, "last_working_date")

    # Checklist items: inject status into C4 (DATE column) for each item row
    # Department rows R5-R11
    checklist = {
        5: "tasks_handed_over", 6: "documents_handed_over", 7: "files_handed_over",
        8: "keys_returned", 9: "toolbox_returned", 10: "access_card_returned",
        11: "dept_others",
        13: "email_cancelled", 14: "software_hardware_returned", 15: "laptop_returned",
        17: "it_others",
        19: "file_shifted", 20: "dues_paid", 21: "medical_card_returned", 22: "hr_others",
        24: "eos_transfer", 25: "finance_others",
    }
    for ri, ph in checklist.items():
        _inject_into_value_cell(t, ri, 4, ph)

    # R26: merged '{{ remarks }}' → already there
    # R27: C0-C2 'Employee Signature:', C3-C5 'Human Resources Manager:' → inject
    _inject_after_label(t, 27, 0, "employee_signature")
    _inject_after_label(t, 27, 3, "hr_signature")

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Visa Renewal Form — paragraph-based
# ─────────────────────────────────────────────────────────────────────────────
def fix_visa_renewal():
    filename = "Visa Renewal Form - INJAAZ.DOCX"
    _restore(filename)
    doc = Document(_dest(filename))

    LONG_BLANK = "___________________________"
    SHORT_BLANK = "_________"

    for i, para in enumerate(doc.paragraphs):
        txt = para.text or ""
        if "I, " in txt and LONG_BLANK in txt and "ID number" in txt:
            new = txt.replace(LONG_BLANK, "{{ employee_name }}", 1)
            remaining = new[new.index("{{ employee_name }}") + len("{{ employee_name }}"):]
            if SHORT_BLANK in remaining:
                new = new.replace(SHORT_BLANK, "{{ employee_id }}", 1)
            if LONG_BLANK in new:
                new = new.replace(LONG_BLANK, "{{ employer }}", 1)
            if para.runs:
                for r in para.runs:
                    r.text = ""
                para.runs[0].text = new
        elif "in the position " in txt and "have completed" in txt:
            new = txt
            if "_________________________________" in new:
                new = new.replace("_________________________________", "{{ position }}", 1)
            blanks = re.findall(r'_{5,}', new)
            if blanks:
                new = new.replace(blanks[0], "{{ years_completed }}", 1)
            if para.runs:
                for r in para.runs:
                    r.text = ""
                para.runs[0].text = new
        elif "Signature of Employee" in txt and "{{ employee_signature }}" not in txt:
            new = txt + " {{ employee_signature }}"
            if para.runs:
                for r in para.runs:
                    r.text = ""
                para.runs[0].text = new

    # Add HR/GM signature table
    has_hr = any("{{ hr_signature }}" in (c.text or "") for tbl in doc.tables for row in tbl.rows for c in row.cells)
    if not has_hr:
        sig = doc.add_table(rows=2, cols=4)
        sig_xml = sig._tbl
        sigPr = sig_xml.find(qn('w:tblPr'))
        if sigPr is None:
            sigPr = OxmlElement('w:tblPr')
            sig_xml.insert(0, sigPr)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        sigPr.append(borders)
        sig.rows[0].cells[0].text = "HR Signature:"
        sig.rows[0].cells[1].text = "{{ hr_signature }}"
        sig.rows[0].cells[2].text = "GM Signature:"
        sig.rows[0].cells[3].text = "{{ gm_signature }}"
        sig.rows[1].cells[0].text = "Decision:"
        sig.rows[1].cells[1].text = "{{ decision_display }}"
        sig.rows[1].cells[2].text = "Date:"
        sig.rows[1].cells[3].text = "{{ form_date }}"

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Staff Appraisal — original is empty (just header), keep rebuilt version
# ─────────────────────────────────────────────────────────────────────────────
def fix_staff_appraisal():
    # The original Staff Appraisal is empty (only header table, no content)
    # Keep the rebuilt version as-is since there's no original content to restore
    print(f"  Staff Appraisal: keeping rebuilt version (original was empty)")


# ─────────────────────────────────────────────────────────────────────────────
# Contract Renewal — original structure is fine, restore and inject ratings
# ─────────────────────────────────────────────────────────────────────────────
def fix_contract_renewal():
    filename = "Employee Contract Renewal Assessment Form Word.docx"
    _restore(filename)
    doc = Document(_dest(filename))

    t2 = doc.tables[2]  # 33R x 7C rating table
    # Rating col = col 6 (last column)
    rating_map = {
        4: "rating_01a", 5: "rating_01b", 6: "rating_01c", 7: "rating_01d", 8: "rating_01e",
        10: "rating_02a", 11: "rating_02b", 12: "rating_02c", 13: "rating_02d", 14: "rating_02e",
        16: "rating_03a", 17: "rating_03b", 18: "rating_03c", 19: "rating_03d", 20: "rating_03e",
        22: "rating_04a", 23: "rating_04b", 24: "rating_04c",
    }
    for ri, ph in rating_map.items():
        _inject_into_value_cell(t2, ri, 6, ph)

    # R26: strength row — inject into cells 2-3 (value area)
    try:
        r26 = t2.rows[26]
        for ci in [2, 3]:
            txt = (r26.cells[ci].text or "").strip()
            if not txt:
                _set_cell_value(r26.cells[ci], "{{ strength }}")
                break
    except IndexError:
        pass

    # R27: areas for improvement — inject into cells 4-5
    try:
        r27 = t2.rows[27]
        for ci in [4, 5]:
            if ci < len(r27.cells):
                txt = (r27.cells[ci].text or "").strip()
                if not txt:
                    _set_cell_value(r27.cells[ci], "{{ areas_for_improvement }}")
                    break
    except IndexError:
        pass

    # R28: overall score — inject into col 6
    _inject_into_value_cell(t2, 28, 6, "overall_score")

    # R29: recommendation — inject into col 0 (merged)
    try:
        r29 = t2.rows[29]
        txt = (r29.cells[0].text or "").strip()
        if "{{ recommendation }}" not in txt:
            _inject_after_label(t2, 29, 0, "recommendation")
    except IndexError:
        pass

    # R32: evaluator signature — inject into the merged row
    try:
        r32 = t2.rows[32]
        txt = (r32.cells[0].text or "").strip()
        if "evaluator_signature" not in txt and "Evaluator" in txt:
            new = txt.replace("_____________________", "{{ evaluator_signature }}", 1)
            new = new.replace("_________", "{{ evaluator_date }}", 1)
            _set_cell_value(r32.cells[0], new)
    except IndexError:
        pass

    # Add HR/GM signature table at end
    has_hr = any("{{ hr_signature }}" in (c.text or "") for tbl in doc.tables for row in tbl.rows for c in row.cells)
    if not has_hr:
        sig = doc.add_table(rows=2, cols=4)
        sig_xml = sig._tbl
        sigPr = sig_xml.find(qn('w:tblPr'))
        if sigPr is None:
            sigPr = OxmlElement('w:tblPr')
            sig_xml.insert(0, sigPr)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        sigPr.append(borders)
        sig.rows[0].cells[0].text = "HR Signature:"
        sig.rows[0].cells[1].text = "{{ hr_signature }}"
        sig.rows[0].cells[2].text = "GM Signature:"
        sig.rows[0].cells[3].text = "{{ gm_signature }}"
        sig.rows[1].cells[0].text = "HR Remarks:"
        sig.rows[1].cells[1].text = "{{ hr_remarks }}"
        sig.rows[1].cells[2].text = "GM Remarks:"
        sig.rows[1].cells[3].text = "{{ gm_remarks }}"

    doc.save(_dest(filename))
    print(f"  Restored + injected: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Performance Evaluation & Commencement — already good, restore from backup
# ─────────────────────────────────────────────────────────────────────────────
def fix_performance():
    filename = "Employee Performance Evaluation Form - INJAAZ.DOCX"
    _restore(filename)
    print(f"  Restored: {filename} (already had proper placeholders)")


def fix_commencement():
    filename = "Commencement Form - INJAAZ.DOCX"
    # Commencement backup is in templates/ subfolder
    src = os.path.join(TMPL_DIR, "Commencement Form - INJAAZ.DOCX")
    if os.path.isfile(src):
        shutil.copy2(src, _dest(filename))
        print(f"  Restored: {filename}")
    else:
        print(f"  Commencement: keeping current (no backup in templates/)")


def main():
    print("Restoring original HR DOCX templates and injecting placeholders...")
    fix_leave()
    fix_duty_resumption()
    fix_passport_release()
    fix_grievance()
    fix_interview()
    fix_station_clearance()
    fix_visa_renewal()
    fix_staff_appraisal()
    fix_contract_renewal()
    fix_performance()
    fix_commencement()
    print("\nDone. Run: python scripts/auto_test_hr_forms.py")


if __name__ == "__main__":
    main()
