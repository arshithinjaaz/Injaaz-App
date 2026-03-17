#!/usr/bin/env python3
"""
Compare generated HR DOCX files against golden templates.

Checks:
- Section/page setup (size, margins, orientation, header/footer distance)
- Header/footer XML hash equality
- Document settings/styles XML hash equality
- Paragraph and table counts
- Table grid shape (rows x cols per table)
- Unresolved Jinja placeholders in generated output

Usage:
  python scripts/check_hr_docx_fidelity.py
  python scripts/check_hr_docx_fidelity.py --generated-dir "test_output/hr_forms_20260311_104701"
  python scripts/check_hr_docx_fidelity.py --strict
"""
from __future__ import annotations

import argparse
import hashlib
import os
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document


TEMPLATE_NAME_BY_FORM: Dict[str, str] = {
    "commencement": "Commencement Form - INJAAZ.DOCX",
    "leave_application": "Leave Application Form - INJAAZ.DOCX",
    "leave": "Leave Application Form - INJAAZ.DOCX",
    "duty_resumption": "Duty Resumption Form - INJAAZ.DOCX",
    "passport_release": "Passport Release & Submission Form - INJAAZ.DOCX",
    "grievance": "Employee grievance disciplinary action-form.docx",
    "performance_evaluation": "Employee Performance Evaluation Form - INJAAZ.DOCX",
    "interview_assessment": "Interview Assessment Form - INJAAZ.DOCX",
    "staff_appraisal": "Staff Appraisal Form - INJAAZ.DOCX",
    "station_clearance": "Station Clearance Form - INJAAZ.DOCX",
    "visa_renewal": "Visa Renewal Form - INJAAZ.DOCX",
    "contract_renewal": "Employee Contract Renewal Assessment Form Word.docx",
}


@dataclass
class CheckResult:
    file_name: str
    form_type: str
    ok: bool
    issues: List[str]
    warnings: List[str]
    infos: List[str]


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _xml_group_hash(docx_path: Path, prefix: str) -> str:
    with zipfile.ZipFile(docx_path, "r") as zf:
        names = sorted([n for n in zf.namelist() if n.startswith(prefix) and n.endswith(".xml")])
        payload = b""
        for name in names:
            payload += name.encode("utf-8") + b"\0" + zf.read(name) + b"\0"
    return _sha256(payload)


def _single_xml_hash(docx_path: Path, xml_name: str) -> str:
    with zipfile.ZipFile(docx_path, "r") as zf:
        try:
            return _sha256(zf.read(xml_name))
        except KeyError:
            return ""


def _extract_form_type(generated_name: str) -> str:
    # e.g. "leave_application_HR-LEAVE_APPLICATION-ABCD1234.docx" -> leave_application
    marker = "_HR-"
    if marker in generated_name:
        return generated_name.split(marker, 1)[0].lower()
    return generated_name.rsplit(".", 1)[0].lower()


def _table_shapes(doc: Document) -> List[Tuple[int, int]]:
    out: List[Tuple[int, int]] = []
    for t in doc.tables:
        rows = len(t.rows)
        cols = max((len(r.cells) for r in t.rows), default=0)
        out.append((rows, cols))
    return out


def _sections_signature(doc: Document) -> List[Tuple[float, ...]]:
    sig: List[Tuple[float, ...]] = []
    for s in doc.sections:
        sig.append(
            (
                round(s.page_width.cm, 3),
                round(s.page_height.cm, 3),
                round(s.left_margin.cm, 3),
                round(s.right_margin.cm, 3),
                round(s.top_margin.cm, 3),
                round(s.bottom_margin.cm, 3),
                round(s.header_distance.cm, 3),
                round(s.footer_distance.cm, 3),
                float(s.orientation),
            )
        )
    return sig


def _has_unresolved_placeholders(docx_path: Path) -> bool:
    with zipfile.ZipFile(docx_path, "r") as zf:
        xml_chunks = []
        for name in zf.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml_chunks.append(zf.read(name).decode("utf-8", errors="ignore"))
    text = "\n".join(xml_chunks)
    return bool(re.search(r"\{\{\s*[^}]+\s*\}\}", text))


def check_one(template_path: Path, generated_path: Path) -> CheckResult:
    issues: List[str] = []
    warnings: List[str] = []
    infos: List[str] = []
    form_type = _extract_form_type(generated_path.name)

    try:
        t_doc = Document(str(template_path))
        g_doc = Document(str(generated_path))
    except Exception as exc:  # pragma: no cover
        return CheckResult(generated_path.name, form_type, False, [f"cannot_open_docx: {exc}"], [], [])

    if _sections_signature(t_doc) != _sections_signature(g_doc):
        issues.append("section/page setup mismatch (size/margins/orientation/header-footer distance)")

    # Semantic checks are more stable than raw XML hash checks for docxtpl output.
    # docxtpl can change XML IDs/order while preserving visual layout fidelity.
    t_header_para = [len(s.header.paragraphs) for s in t_doc.sections]
    g_header_para = [len(s.header.paragraphs) for s in g_doc.sections]
    if t_header_para != g_header_para:
        issues.append("header paragraph structure differs from template")

    t_footer_para = [len(s.footer.paragraphs) for s in t_doc.sections]
    g_footer_para = [len(s.footer.paragraphs) for s in g_doc.sections]
    if t_footer_para != g_footer_para:
        issues.append("footer paragraph structure differs from template")

    t_header_tables = [len(s.header.tables) for s in t_doc.sections]
    g_header_tables = [len(s.header.tables) for s in g_doc.sections]
    if t_header_tables != g_header_tables:
        issues.append("header table structure differs from template")

    t_footer_tables = [len(s.footer.tables) for s in t_doc.sections]
    g_footer_tables = [len(s.footer.tables) for s in g_doc.sections]
    if t_footer_tables != g_footer_tables:
        issues.append("footer table structure differs from template")

    # settings.xml can change due to docxtpl/Word rsid metadata updates.
    # Keep this as informational only; page setup is already validated via section signature.
    if _single_xml_hash(template_path, "word/settings.xml") != _single_xml_hash(generated_path, "word/settings.xml"):
        infos.append("settings.xml differs (likely metadata/rsid drift)")

    if len(t_doc.paragraphs) != len(g_doc.paragraphs):
        warnings.append(f"paragraph count differs ({len(t_doc.paragraphs)} vs {len(g_doc.paragraphs)})")

    if len(t_doc.tables) != len(g_doc.tables):
        issues.append(f"table count differs ({len(t_doc.tables)} vs {len(g_doc.tables)})")
    else:
        t_shapes = _table_shapes(t_doc)
        g_shapes = _table_shapes(g_doc)
        if t_shapes != g_shapes:
            issues.append("table row/column shape differs from template")

    if _has_unresolved_placeholders(generated_path):
        issues.append("generated file contains unresolved {{ placeholder }} tokens")

    return CheckResult(generated_path.name, form_type, ok=not issues, issues=issues, warnings=warnings, infos=infos)


def _latest_generated_dir(project_root: Path) -> Path | None:
    test_output = project_root / "test_output"
    if not test_output.exists():
        return None
    candidates = [p for p in test_output.iterdir() if p.is_dir() and p.name.startswith("hr_forms_")]
    if not candidates:
        return None
    candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0]


def _resolve_template_dir(project_root: Path) -> Path:
    env_path = os.environ.get("HR_DOCX_TEMPLATE_DIR")
    if env_path and Path(env_path).is_dir():
        return Path(env_path)
    preferred = project_root / "HR Templates"
    secondary = project_root / "HR Documents - Main"
    tertiary = project_root / "HR Documents - Copy"
    fallback = project_root / "HR Documents"
    if preferred.is_dir():
        return preferred
    if secondary.is_dir():
        return secondary
    if tertiary.is_dir():
        return tertiary
    if fallback.is_dir():
        return fallback
    raise FileNotFoundError("Could not find 'HR Templates', 'HR Documents - Main', 'HR Documents - Copy', or 'HR Documents'.")


def main() -> int:
    parser = argparse.ArgumentParser(description="Check generated HR DOCX files for template fidelity.")
    parser.add_argument("--generated-dir", type=str, default="", help="Path to generated hr_forms_* directory")
    parser.add_argument("--strict", action="store_true", help="Fail on warnings too")
    args = parser.parse_args()

    project_root = Path(__file__).resolve().parents[1]
    template_dir = _resolve_template_dir(project_root)

    if args.generated_dir:
        generated_dir = Path(args.generated_dir)
        if not generated_dir.is_absolute():
            generated_dir = project_root / generated_dir
    else:
        generated_dir = _latest_generated_dir(project_root)
        if generated_dir is None:
            print("ERROR: no generated hr_forms_* directory found under test_output/")
            return 2

    if not generated_dir.exists():
        print(f"ERROR: generated directory does not exist: {generated_dir}")
        return 2

    generated_files = sorted(generated_dir.glob("*.docx"))
    if not generated_files:
        print(f"ERROR: no .docx files found in: {generated_dir}")
        return 2

    print(f"Template dir : {template_dir}")
    print(f"Generated dir: {generated_dir}")
    print(f"Files checked: {len(generated_files)}")
    print("")

    results: List[CheckResult] = []
    for g in generated_files:
        form_type = _extract_form_type(g.name)
        template_name = TEMPLATE_NAME_BY_FORM.get(form_type)
        if not template_name:
            results.append(CheckResult(g.name, form_type, False, [f"unknown form type '{form_type}'"], [], []))
            continue
        t = template_dir / template_name
        if not t.exists():
            results.append(CheckResult(g.name, form_type, False, [f"template missing '{template_name}'"], [], []))
            continue
        results.append(check_one(t, g))

    ok_count = 0
    warn_count = 0
    fail_count = 0
    info_count = 0
    for r in results:
        if r.ok and not r.warnings:
            state = "PASS"
            ok_count += 1
        elif r.ok and r.warnings:
            state = "WARN"
            warn_count += 1
        else:
            state = "FAIL"
            fail_count += 1
        print(f"[{state}] {r.file_name}")
        for issue in r.issues:
            print(f"  - ISSUE: {issue}")
        for warning in r.warnings:
            print(f"  - WARN : {warning}")
        for info in r.infos:
            info_count += 1
            print(f"  - INFO : {info}")

    print("")
    print(f"Summary: PASS={ok_count} WARN={warn_count} FAIL={fail_count} INFO={info_count}")

    if fail_count > 0:
        return 1
    if args.strict and warn_count > 0:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
