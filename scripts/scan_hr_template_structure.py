"""Print table structure of HR DOCX templates to see label positions for placeholder injection."""
import os, sys, re
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

HR_DOCS = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "HR Documents")
PH = re.compile(r'\{\{\s*(\w+)\s*\}\}')

TARGET = sys.argv[1] if len(sys.argv) > 1 else None

FORMS = [
    "Staff Appraisal Form - INJAAZ.DOCX",
    "Visa Renewal Form - INJAAZ.DOCX",
    "Leave Application Form - INJAAZ.DOCX",
    "Duty Resumption Form - INJAAZ.DOCX",
    "Passport Release & Submission Form - INJAAZ.DOCX",
    "Employee grievance disciplinary action-form.docx",
    "Interview Assessment Form - INJAAZ.DOCX",
    "Station Clearance Form - INJAAZ.DOCX",
    "Employee Contract Renewal Assessment Form Word.docx",
]

for fname in FORMS:
    if TARGET and TARGET.lower() not in fname.lower():
        continue
    path = os.path.join(HR_DOCS, fname)
    if not os.path.exists(path):
        print(f"MISSING: {fname}"); continue
    doc = Document(path)
    print(f"\n{'='*70}")
    print(f"FILE: {fname}  ({len(doc.tables)} tables)")
    print('='*70)
    for ti, tbl in enumerate(doc.tables):
        print(f"\n  Table {ti}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
        for ri, row in enumerate(tbl.rows):
            cells_text = []
            for cell in row.cells:
                t = (cell.text or "").strip().replace("\n", " ")[:60]
                cells_text.append(repr(t))
            # Only print rows with text
            if any(c != repr("") for c in cells_text):
                print(f"    R{ri}: {' | '.join(cells_text)}")
