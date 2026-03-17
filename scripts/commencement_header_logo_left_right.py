"""
Apply reversed header to Commencement Form: heading on the left, logo on the right.
Inserts a 1x2 table at the top and removes the old centered header paragraphs.

Run from project root: python scripts/commencement_header_logo_left_right.py
"""
import os
import sys
import shutil

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_margins(cell, **kwargs):
    """Set table cell margins. Values in twentieths of a point (dxa). Use 0 for minimal gap."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m in ("top", "start", "bottom", "end"):
        if m in kwargs:
            node = OxmlElement(f"w:{m}")
            node.set(qn("w:w"), str(kwargs[m]))
            node.set(qn("w:type"), "dxa")
            tcMar.append(node)
    tcPr.append(tcMar)


def add_header_table_to_commencement(doc_path, logo_path=None, backup=True):
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if logo_path is None:
        logo_path = os.path.join(base, "static", "logo.png")
    if not os.path.isfile(doc_path):
        raise FileNotFoundError(f"Template not found: {doc_path}")
    if not os.path.isfile(logo_path):
        raise FileNotFoundError(f"Logo not found: {logo_path}")

    if backup:
        backup_dir = os.path.join(os.path.dirname(doc_path), "templates")
        os.makedirs(backup_dir, exist_ok=True)
        backup_path = os.path.join(backup_dir, "Commencement Form - INJAAZ (before header).docx")
        shutil.copy2(doc_path, backup_path)
        print(f"Backed up to: {backup_path}")

    doc = Document(doc_path)
    body = doc._element.body

    # Find intro text BEFORE modifying document structure
    intro_text = None
    intro_para_element = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("To complete the administrative aspect"):
            intro_text = text
            intro_para_element = para._element
            break
    
    # If not found, use default text
    if not intro_text:
        intro_text = "To complete the administrative aspect of your Employment please complete this form within 5 days of joining and fax it back to AH or email it to joana@ajmanholding.ae"

    # If there's already a header table (e.g. from a previous run with broken image), remove it
    had_existing_table = False
    first = body[0] if len(body) else None
    if first is not None:
        tag = first.tag.split("}")[-1] if "}" in first.tag else first.tag
        if tag == "tbl":
            body.remove(first)
            had_existing_table = True

    # Build header table IN THE SAME DOC so the logo is embedded in this docx package
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Cell widths matching reference: left cell wider (~13.5cm), right cell narrower (~3cm for logo)
    try:
        table.rows[0].cells[0].width = Cm(13.5)  # Wider for headline + intro text
        table.rows[0].cells[1].width = Cm(3.5)  # Narrower for logo
    except Exception:
        pass
    FORM_FONT = "Calibri"
    # Left cell: "Commencement Form" + intro text (matching reference format)
    left_cell = table.rows[0].cells[0]
    left_cell.text = ""
    
    # First paragraph: "Commencement Form" (bold, larger headline)
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.paragraph_format.keep_together = True
    p_left.paragraph_format.space_after = Pt(0)
    run = p_left.add_run("Commencement Form")
    run.bold = True
    run.font.size = Pt(20)  # Larger headline
    run.font.name = FORM_FONT
    
    # Second paragraph: intro text (smaller subheading)
    p_intro = left_cell.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_intro.paragraph_format.space_before = Pt(0)
    p_intro.paragraph_format.space_after = Pt(12)
    run_intro = p_intro.add_run(intro_text)
    run_intro.font.name = FORM_FONT
    run_intro.font.size = Pt(9)  # Smaller subheading
    # Right cell: logo flush right (paragraph right-aligned, wide cell)
    right_cell = table.rows[0].cells[1]
    for p in right_cell.paragraphs:
        p.clear()
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    run_right = p_right.add_run()
    run_right.add_picture(logo_path, width=Cm(1.8))

    # Minimise vertical gap: small top/bottom margins on header cells
    set_cell_margins(left_cell, top=60, bottom=0)
    set_cell_margins(right_cell, top=60, bottom=0)
    
    # Minimize table row height - set row height to "at least" a small value
    try:
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "240")  # ~12pt minimum height
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)
    except Exception:
        pass

    tbl_element = table._tbl

    # Table was appended at end; move it to the beginning
    body.remove(tbl_element)
    body.insert(0, tbl_element)

    # On first run (no existing table): remove old header paragraphs and intro text paragraph
    if not had_existing_table:
        to_remove = []
        # Remove old header paragraphs (empty/INJAAZ/Commencement/empty)
        for i in range(1, min(8, len(body))):
            child = body[i]
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                to_remove.append(child)
        # Also remove the intro text paragraph if we found it (it's now in the table)
        if intro_para_element is not None and intro_para_element in body:
            to_remove.append(intro_para_element)
        for el in to_remove:
            if el in body:
                body.remove(el)

    # Ensure first paragraph after table has minimal spacing (matching reference format)
    if doc.paragraphs:
        # First paragraph after table should have space_after matching reference (~12pt)
        doc.paragraphs[0].paragraph_format.space_before = Pt(0)
        # Keep default spacing for content paragraphs

    doc.save(doc_path)
    print("Header updated: heading left, logo right. Saved:", doc_path)


def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    doc_path = os.path.join(base, "HR Documents", "Commencement Form - INJAAZ.DOCX")
    add_header_table_to_commencement(doc_path, backup=True)


if __name__ == "__main__":
    main()
