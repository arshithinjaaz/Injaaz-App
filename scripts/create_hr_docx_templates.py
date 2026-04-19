"""
Create HR DOCX templates with placeholders for docxtpl.
Run: python scripts/create_hr_docx_templates.py

This creates template versions in HR Documents/templates/ that mirror the
structure of your HR forms. The placeholders (e.g. {{ employee_name }}) will
be replaced with form data when users download filled documents.
"""
import os
import sys

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_commencement_template():
    """Create Commencement Form template with docxtpl placeholders."""
    doc = Document()
    doc.add_paragraph('COMMENCEMENT FORM', style='Heading 1')
    doc.add_paragraph()
    doc.add_paragraph('Personal Details:')
    doc.add_paragraph('Name: {{ employee_name }}')
    doc.add_paragraph('Position: {{ position }}')
    doc.add_paragraph('Contacts: {{ contacts }}')
    doc.add_paragraph('Department: {{ department }}')
    doc.add_paragraph('Organization: {{ organization }}')
    doc.add_paragraph('Date of Joining: {{ date_of_joining }}')
    doc.add_paragraph()
    doc.add_paragraph('Bank Account Details:')
    doc.add_paragraph('Bank Name: {{ bank_name }}')
    doc.add_paragraph('Branch: {{ bank_branch }}')
    doc.add_paragraph('Account Number: {{ account_number }}')
    doc.add_paragraph()
    doc.add_paragraph("Employee's Signature:")
    doc.add_paragraph('{{ employee_signature }}')
    doc.add_paragraph('Date: {{ employee_sign_date }}')
    doc.add_paragraph()
    doc.add_paragraph('Reporting To:')
    doc.add_paragraph('Name: {{ reporting_to_name }}')
    doc.add_paragraph('Designation / Title: {{ reporting_to_designation }}')
    doc.add_paragraph('Contact No: {{ reporting_to_contact }}')
    doc.add_paragraph('Signature: {{ reporting_to_signature }}')
    doc.add_paragraph('Date: {{ reporting_sign_date }}')
    return doc


def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    hr_docs = os.path.join(base, 'HR Documents')
    templates_dir = os.path.join(hr_docs, 'templates')
    os.makedirs(templates_dir, exist_ok=True)

    # Commencement
    doc = create_commencement_template()
    out = os.path.join(templates_dir, 'Commencement Form - INJAAZ.DOCX')
    doc.save(out)
    print(f'Created: {out}')


if __name__ == '__main__':
    main()
