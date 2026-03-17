"""Scan all HR Word templates and print every placeholder found."""
import os, sys, re
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

HR_DOCS = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "HR Documents")
PH = re.compile(r'\{\{\s*(\w+)\s*\}\}')

FORMS = [
    "Leave Application Form - INJAAZ.DOCX",
    "Commencement Form - INJAAZ.DOCX",
    "Duty Resumption Form - INJAAZ.DOCX",
    "Passport Release & Submission Form - INJAAZ.DOCX",
    "Employee grievance disciplinary action-form.docx",
    "Interview Assessment Form - INJAAZ.DOCX",
    "Staff Appraisal Form - INJAAZ.DOCX",
    "Station Clearance Form - INJAAZ.DOCX",
    "Visa Renewal Form - INJAAZ.DOCX",
    "Employee Performance Evaluation Form - INJAAZ.DOCX",
    "Employee Contract Renewal Assessment Form Word.docx",
]

for fname in FORMS:
    path = os.path.join(HR_DOCS, fname)
    if not os.path.exists(path):
        print(f"  MISSING: {fname}")
        continue
    doc = Document(path)
    phs = set()
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for m in PH.findall(cell.text or ""):
                    phs.add(m)
    for para in doc.paragraphs:
        for m in PH.findall(para.text or ""):
            phs.add(m)
    label = "OK" if phs else "NO PLACEHOLDERS"
    print(f"\n[{label}] {fname}")
    for p in sorted(phs):
        print(f"    {{ {p} }}")
