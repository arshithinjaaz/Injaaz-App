#!/usr/bin/env python3
"""
Prepare placeholder-enabled HR templates from exact Main templates.

Flow:
1) Copy .docx files from "HR Documents - Main" to "HR Documents - Main/templates"
2) Inject docxtpl placeholders into copied files only
3) Keep original Main files untouched

Run:
  python scripts/prepare_hr_main_templates.py
"""
import os
import shutil
import sys

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MAIN_DIR = os.path.join(BASE, "HR Documents - Main")
TARGET_DIR = os.path.join(MAIN_DIR, "templates")

sys.path.insert(0, BASE)
from scripts.inject_all_hr_placeholders import FORMS, inject_by_labels  # noqa: E402
from scripts.inject_commencement_placeholders import inject_commencement_placeholders  # noqa: E402
from docx import Document  # noqa: E402


def _copy_leave_option_row_text_and_style(source_cell, target_cell, from_word, to_word):
    """Copy leave option row preserving run style, replacing label text."""
    src_para = source_cell.paragraphs[0] if source_cell.paragraphs else None
    if not src_para or not target_cell.paragraphs:
        target_cell.text = f"[ ] {to_word} Leave"
        return
    tgt_para = target_cell.paragraphs[0]
    tgt_para.clear()
    for run in src_para.runs:
        txt = (run.text or "").replace(from_word, to_word)
        r = tgt_para.add_run(txt)
        if run.font.size:
            r.font.size = run.font.size
        r.bold = run.bold


def inject_visa_placeholders(doc_path):
    """Inject key placeholders for Visa Renewal paragraph-based template."""
    doc = Document(doc_path)
    try:
        if len(doc.paragraphs) > 5 and "{{ form_date }}" not in doc.paragraphs[5].text:
            doc.paragraphs[5].add_run(" {{ form_date }}")

        if len(doc.paragraphs) > 9:
            p9 = doc.paragraphs[9]
            if "{{ employee_name }}" not in p9.text:
                p9.text = "I, {{ employee_name }} with ID number {{ employee_id }} employed by {{ employer }}"

        if len(doc.paragraphs) > 11:
            p11 = doc.paragraphs[11]
            if "{{ position }}" not in p11.text:
                p11.text = "in the position {{ position }} have completed {{ years_completed }} years."

        if len(doc.paragraphs) > 12 and "{{ decision }}" not in doc.paragraphs[12].text:
            doc.paragraphs[12].add_run(" {{ decision }}")

        if len(doc.paragraphs) > 21 and "{{ employee_signature }}" not in doc.paragraphs[21].text:
            doc.paragraphs[21].add_run(" {{ employee_signature }}")
    finally:
        doc.save(doc_path)


def inject_performance_placeholders_safe(doc_path):
    """Inject placeholders in Performance Evaluation template (Main structure)."""
    doc = Document(doc_path)
    # Table 0: employee information
    t0 = doc.tables[0]
    t0.rows[0].cells[1].text = "{{ employee_name }}"
    t0.rows[0].cells[3].text = "{{ date_of_evaluation }}"
    t0.rows[1].cells[1].text = "{{ employee_id }}"
    t0.rows[1].cells[3].text = "{{ date_of_joining }}"
    t0.rows[2].cells[1].text = "{{ department }}"
    t0.rows[2].cells[3].text = "{{ designation }}"
    t0.rows[3].cells[1].text = "{{ evaluation_done_by }}"

    # Table 1: score rows + overall score
    t1 = doc.tables[1]
    for i in range(1, 11):
        t1.rows[2 + i].cells[3].text = "{{ score_%02d }}" % i
    t1.rows[13].cells[3].text = "{{ overall_score }}"

    # Table 2: evaluator and signatures
    t2 = doc.tables[2]
    t2.rows[1].cells[1].text = "{{ evaluator_name }}"
    t2.rows[1].cells[4].text = "{{ evaluator_designation }}"
    t2.rows[2].cells[1].text = "{{ evaluator_observation }}"
    t2.rows[3].cells[1].text = "{{ area_of_concern }}"
    t2.rows[4].cells[1].text = "{{ training_required }}"
    t2.rows[5].cells[1].text = "{{ employee_comments }}"
    t2.rows[6].cells[1].text = "{{ employee_signature }}"
    t2.rows[6].cells[5].text = "{{ employee_sign_date }}"
    t2.rows[7].cells[1].text = "{{ evaluator_signature }}"
    t2.rows[7].cells[5].text = "{{ evaluator_sign_date }}"

    # Table 3: manager / GM / HR
    t3 = doc.tables[3]
    t3.rows[0].cells[1].text = "{{ concern_incharge_name }}"
    t3.rows[1].cells[1].text = "{{ incharge_comments }}"
    t3.rows[2].cells[1].text = "{{ gm_remarks }}\n\n{{ gm_signature }}"
    t3.rows[4].cells[1].text = "{{ hr_remarks }}\n\n{{ hr_signature }}"
    doc.save(doc_path)


def inject_leave_placeholders_safe(doc_path):
    """Fix leave form placeholder positions for Main template structure."""
    doc = Document(doc_path)
    t = doc.tables[0]

    # Personal info rows
    t.rows[0].cells[0].text = "Name: {{ employee_name }}"
    t.rows[1].cells[0].text = "Job Title: {{ job_title }}"
    t.rows[1].cells[3].text = "Today's Date: {{ today_date }}"
    t.rows[2].cells[0].text = "Employee ID: {{ employee_id }}"
    t.rows[2].cells[3].text = "Department: {{ department }}"
    t.rows[3].cells[0].text = "Date of Joining: {{ date_of_joining }}"
    t.rows[3].cells[3].text = "Mobile No.: {{ mobile_no }}"
    t.rows[4].cells[0].text = "Last Leave Date: {{ last_leave_date }}"

    # Leave details rows
    # Keep leave-type option rows exactly as in Main template for visual consistency
    # (Annual/Sick/Compassionate/Unpaid/Hajj/Other checkbox layout).
    # Re-assert Unpaid row (it can be blanked by generic label matching on some variants).
    try:
        for ci in (0, 1, 2):
            _copy_leave_option_row_text_and_style(
                t.rows[14].cells[ci],  # Hajj row (reference style)
                t.rows[13].cells[ci],  # Unpaid row
                "Hajj",
                "Unpaid",
            )
    except Exception:
        # Fallback text-only restore
        t.rows[13].cells[0].text = "[ ] Unpaid Leave"
        for ci in (1, 2):
            t.rows[13].cells[ci].text = ""
    t.rows[16].cells[0].text = "Total No. of Days Requested: {{ total_days_requested }}"
    t.rows[17].cells[0].text = "First Day of leave: {{ first_day_of_leave }}"
    t.rows[17].cells[2].text = "Last Day of Leave: {{ last_day_of_leave }}"
    t.rows[18].cells[0].text = "Date returning to work: {{ date_returning_to_work }}"
    # Salary row: both options in same format as leave-type checkboxes; ticks applied in post-render.
    t.rows[19].cells[0].text = "Leave Salary Advance Requested: [ ] YES  [ ] NO"
    for ci in (1, 2, 3):
        t.rows[19].cells[ci].text = ""
    t.rows[20].cells[0].text = "Telephone Number where you can be reached: {{ telephone_reachable }}"

    # Signatures and replacement
    t.rows[21].cells[0].text = "Replacement Name: {{ replacement_name }}"
    t.rows[21].cells[2].text = "Signature: {{ replacement_signature }}"
    t.rows[22].cells[0].text = "Employee Signature: {{ employee_signature }}"
    t.rows[22].cells[2].text = "Manager Signature: {{ gm_signature }}"

    # HR section
    t.rows[23].cells[0].text = "Checked by HR: {{ hr_checked }}"
    t.rows[24].cells[0].text = "HR Comments: {{ hr_comments }}"
    t.rows[26].cells[0].text = "Balance C/F: {{ hr_balance_cf }}"
    t.rows[26].cells[3].text = "Contract Year: {{ hr_contract_year }}"
    t.rows[27].cells[0].text = "Paid: {{ hr_paid }}"
    t.rows[27].cells[3].text = "Unpaid: {{ hr_unpaid }}"
    t.rows[28].cells[0].text = "HR Signature: {{ hr_signature }}"
    t.rows[28].cells[3].text = "Date: {{ hr_date }}"
    doc.save(doc_path)


def main():
    if not os.path.isdir(MAIN_DIR):
        raise FileNotFoundError(f"Missing folder: {MAIN_DIR}")

    os.makedirs(TARGET_DIR, exist_ok=True)

    copied = 0
    for name in os.listdir(MAIN_DIR):
        src = os.path.join(MAIN_DIR, name)
        if os.path.isfile(src) and name.lower().endswith(".docx"):
            dst = os.path.join(TARGET_DIR, name)
            shutil.copy2(src, dst)
            copied += 1

    print(f"Copied {copied} DOCX files into: {TARGET_DIR}")

    injected = 0
    for filename, label_placeholders in FORMS:
        dst = os.path.join(TARGET_DIR, filename)
        if not os.path.isfile(dst):
            print(f"Skip (missing in target): {filename}")
            continue
        ok = inject_by_labels(dst, label_placeholders, backup=False)
        if ok:
            injected += 1
            print(f"Injected placeholders: {filename}")

    # Dedicated injectors for forms where label-based table scanning is insufficient.
    leave_path = os.path.join(TARGET_DIR, "Leave Application Form - INJAAZ.DOCX")
    if os.path.isfile(leave_path):
        try:
            inject_leave_placeholders_safe(leave_path)
            print("Injected placeholders: Leave Application Form - INJAAZ.DOCX (safe override)")
        except Exception as exc:
            print(f"Warn: leave injection failed: {exc}")

    commencement_path = os.path.join(TARGET_DIR, "Commencement Form - INJAAZ.DOCX")
    if os.path.isfile(commencement_path):
        try:
            inject_commencement_placeholders(commencement_path, backup=False)
            print("Injected placeholders: Commencement Form - INJAAZ.DOCX")
        except Exception as exc:
            print(f"Warn: commencement injection failed: {exc}")

    perf_path = os.path.join(TARGET_DIR, "Employee Performance Evaluation Form - INJAAZ.DOCX")
    if os.path.isfile(perf_path):
        try:
            inject_performance_placeholders_safe(perf_path)
            print("Injected placeholders: Employee Performance Evaluation Form - INJAAZ.DOCX")
        except Exception as exc:
            print(f"Warn: performance injection failed: {exc}")

    visa_path = os.path.join(TARGET_DIR, "Visa Renewal Form - INJAAZ.DOCX")
    if os.path.isfile(visa_path):
        try:
            inject_visa_placeholders(visa_path)
            print("Injected placeholders: Visa Renewal Form - INJAAZ.DOCX")
        except Exception as exc:
            print(f"Warn: visa injection failed: {exc}")

    print(f"Done. Placeholder injection completed for {injected} files.")


if __name__ == "__main__":
    main()
