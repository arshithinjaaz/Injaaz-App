"""
Comprehensive HR DOCX template placeholder injection.

Restores each HR Word template from the cleanest available backup and injects
{{ placeholder }} tags at the exact cell positions mapped from the template
structure scan.  Handles paragraph-based forms (Visa Renewal) and re-builds
empty templates (Staff Appraisal).

Run from project root:
    python scripts/fix_all_hr_templates.py
"""
import os
import re
import shutil
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
HR_DOCS = os.path.join(BASE, "HR Documents")
TMPL_DIR = os.path.join(HR_DOCS, "templates")


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_text(cell, text):
    """Set cell text, preserving the paragraph object (keeps style)."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        cell.text = text


def _set_run_in_cell(cell, new_text):
    """Replace the text content of a cell, keeping the first paragraph/run style."""
    # Clear all runs in all paragraphs
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
    elif cell.paragraphs:
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run(new_text)
    else:
        cell.text = new_text


def _inject_cell(table, row_idx, col_idx, placeholder):
    """Inject a placeholder into the cell at table[row_idx][col_idx]."""
    ph = "{{ %s }}" % placeholder
    try:
        cell = table.rows[row_idx].cells[col_idx]
        existing = (cell.text or "").strip()
        # Don't re-inject if already present
        if ("{{ %s }}" % placeholder) in existing:
            return
        # If cell is empty, just set the placeholder
        if not existing:
            _set_cell_text(cell, ph)
        else:
            # Append placeholder after existing text (handles merged label cells)
            _set_cell_text(cell, existing + " " + ph)
    except IndexError:
        pass


def _inject_col_value_cell(table, row_idx, col_idx, placeholder):
    """Set cell to ONLY the placeholder (pure value cell, no label)."""
    ph = "{{ %s }}" % placeholder
    try:
        cell = table.rows[row_idx].cells[col_idx]
        if ("{{ %s }}" % placeholder) in (cell.text or ""):
            return
        _set_cell_text(cell, ph)
    except IndexError:
        pass


def _add_table(doc, rows, cols):
    """Add a plain table with borders, cross-document safe (no named style needed)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    # Add simple borders via XML without relying on a named style
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    return tbl


def _add_row_to_table(table, label, placeholder, label_bold=False):
    """Append a new 2-column row to a table with label | {{ placeholder }}."""
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = "{{ %s }}" % placeholder
    if label_bold:
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
    return row


def _source_path(filename):
    """Return the best available source for the template:
    prefer 'before placeholders' backup (cleanest), fall back to current file."""
    bk = os.path.join(
        TMPL_DIR,
        filename.replace(".docx", " (before placeholders).docx")
               .replace(".DOCX", " (before placeholders).docx"),
    )
    if os.path.isfile(bk):
        return bk
    return os.path.join(HR_DOCS, filename)


def _dest_path(filename):
    return os.path.join(HR_DOCS, filename)


def _backup_current(filename):
    dest = _dest_path(filename)
    if os.path.isfile(dest):
        bk_name = filename.replace(".docx", " (before fix).docx").replace(".DOCX", " (before fix).docx")
        bk = os.path.join(TMPL_DIR, bk_name)
        os.makedirs(TMPL_DIR, exist_ok=True)
        shutil.copy2(dest, bk)


# ─────────────────────────────────────────────────────────────────────────────
# 1. Leave Application Form
# ─────────────────────────────────────────────────────────────────────────────

def fix_leave_application():
    filename = "Leave Application Form - INJAAZ.DOCX"
    _backup_current(filename)
    doc = Document(_source_path(filename))
    # Table 1 is the main content table (table 0 = header)
    t = doc.tables[1]
    # Row 0: Name merged - inject employee_name into the row
    _inject_cell(t, 0, 0, "employee_name")

    # Row 1: col0='Job Title', col3='Today's Date:' → add today_date
    _inject_col_value_cell(t, 1, 3, "today_date")

    # Row 2: col0='Employee ID', col3='Department:' → add department
    _inject_col_value_cell(t, 2, 3, "department")

    # Row 3: col0='Date of Joining', col3='Mobile No.:' → add mobile_no
    _inject_col_value_cell(t, 3, 3, "mobile_no")

    # Row 6: col3 = 'Number of Days' - add leave_type_display to a better cell
    # leave_type_display spans the leave checkboxes - we add it to a new row if needed
    # For now, inject leave_type into the leave rows as a summary
    # Row 16: 'Total No. of Days' → cols 0-3 merged, already has placeholder
    # Row 17: R17 col0='First Day of leave:' label cell, col1-3='First Day of leave {{ first_day_of_leave }}'
    # We need to add last_day_of_leave - col3 is merged with col1/2, so add a new col if not possible
    # Actually R17 has 4 columns. col0=label, col1(merged with 2,3)=value
    # Since merging means only 2 unique cells, we can only inject in col0 or col1
    # The label "First Day of leave:" is col0. last_day_of_leave has no row - check if a row exists
    # Looking at original form: may not have "Last Day of Leave" row - skip if not present

    # Row 18: 'Date returning to work' - already has placeholder

    # Row 20: 'Telephone Number...' value cell - inject telephone_reachable
    # Check col3 or merged cols
    try:
        r20 = t.rows[20]
        # Find the value cell (not the label)
        cells_seen = {}
        for ci, cell in enumerate(r20.cells):
            txt = (cell.text or "").strip()
            if txt not in cells_seen:
                cells_seen[txt] = ci
        # The last unique cell (rightmost) gets the placeholder
        for ci in range(len(r20.cells) - 1, -1, -1):
            cell = r20.cells[ci]
            txt = (cell.text or "").strip()
            if "Telephone" in txt and "{{ telephone_reachable }}" not in txt:
                _set_cell_text(cell, txt + " {{ telephone_reachable }}")
                break
    except (IndexError, AttributeError):
        pass

    # Row 21: 'Name:' (replacement) col2+3 = 'Signature:' → add replacement_name and replacement_signature
    try:
        r21 = t.rows[21]
        # col2 or col3 has 'Signature:' - find first empty-ish cell
        for ci in [2, 3]:
            cell = r21.cells[ci]
            txt = (cell.text or "").strip()
            if "Signature" in txt and "{{ " not in txt:
                _set_cell_text(cell, txt + " {{ replacement_signature }}")
                break
        # col0/col1 has 'Name:' - inject replacement_name
        for ci in [0, 1]:
            cell = r21.cells[ci]
            txt = (cell.text or "").strip()
            if "Name" in txt and "{{ " not in txt:
                _set_cell_text(cell, txt + " {{ replacement_name }}")
                break
    except (IndexError, AttributeError):
        pass

    # Row 22: 'Employee Signature' col0/1 - already has placeholder
    # col2/3 = 'Manager Signature:' → add gm_signature
    try:
        r22 = t.rows[22]
        for ci in [2, 3]:
            cell = r22.cells[ci]
            txt = (cell.text or "").strip()
            if "Manager" in txt and "{{ " not in txt:
                _set_cell_text(cell, txt + " {{ gm_signature }}")
                break
    except (IndexError, AttributeError):
        pass

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 2. Duty Resumption Form
# ─────────────────────────────────────────────────────────────────────────────

def fix_duty_resumption():
    filename = "Duty Resumption Form - INJAAZ.DOCX"
    _backup_current(filename)
    doc = Document(_source_path(filename))
    # Table 1: 4 rows x 2 cols
    t1 = doc.tables[1]
    # R2 col0='Employee ID:' col1='Job Title {{ job_title }}' - wrong! fix employee_id
    r2c0 = (t1.rows[2].cells[0].text or "").strip()
    r2c1 = (t1.rows[2].cells[1].text or "").strip()
    if "Employee ID" in r2c0 and "employee_id" not in r2c1:
        _inject_col_value_cell(t1, 2, 1, "employee_id")
    # R3 col0='Job Title:' col1='Job Title {{ job_title }}' - fix: add company
    # Actually company is missing - inject into correct row
    # Add company row after row 2 or find 'Company' label
    company_found = any("company" in (r.cells[0].text or "").lower() for r in t1.rows)
    if not company_found:
        row = t1.add_row()
        row.cells[0].text = "Company:"
        row.cells[1].text = "{{ company }}"

    # Table 2: leave dates table
    t2 = doc.tables[2]
    # R1: 'Leave Started:' col1='Leave Started {{ leave_started }}' - OK
    # 'Leave Ended:' is missing
    leave_ended_found = any("leave ended" in (r.cells[0].text or "").lower() for r in t2.rows)
    if not leave_ended_found:
        # Insert after row 1 (Leave Started)
        row = t2.add_row()
        row.cells[0].text = "Leave Ended:"
        row.cells[1].text = "{{ leave_ended }}"

    # R5: 'Employee Signature:' col1='Employee Signature {{ employee_signature }}' - OK
    # Add gm_signature to table 4
    t4 = doc.tables[4]
    # R0: 'Approved by Line Manager:' | 'Date:' → add gm_signature
    r0c0 = (t4.rows[0].cells[0].text or "").strip()
    if "Approved" in r0c0 and "{{ gm_signature }}" not in r0c0:
        _set_cell_text(t4.rows[0].cells[0], r0c0 + " {{ gm_signature }}")

    # Add line_manager_remarks to table 3
    t3 = doc.tables[3]
    try:
        r0 = t3.rows[0].cells[0]
        r1 = t3.rows[1].cells[0]
        lmr_txt = (r0.text or "").strip()
        if "Line Manager" in lmr_txt and "{{ line_manager_remarks }}" not in lmr_txt:
            _set_cell_text(r1, "{{ line_manager_remarks }}")
    except (IndexError, AttributeError):
        pass

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 3. Passport Release Form
# ─────────────────────────────────────────────────────────────────────────────

def fix_passport_release():
    filename = "Passport Release & Submission Form - INJAAZ.DOCX"
    _backup_current(filename)
    doc = Document(_source_path(filename))
    # Table 2: purpose/release date/signature (4 rows x 1 col)
    t2 = doc.tables[2]
    for ri, row in enumerate(t2.rows):
        txt = (row.cells[0].text or "").strip()
        if "Purpose of Release" in txt and "{{ purpose_of_release }}" not in txt:
            _set_cell_text(row.cells[0], txt + " {{ purpose_of_release }}")
        elif "Release Date" in txt and "{{ release_date }}" not in txt:
            _set_cell_text(row.cells[0], txt + " {{ release_date }}")
        elif "Employee Signature" in txt and "{{ employee_signature }}" not in txt:
            _set_cell_text(row.cells[0], txt + " {{ employee_signature }}")

    # Table 3: HR/GM approval
    t3 = doc.tables[3]
    # R0: 'Approved by Line Manager:' | 'Date:' → add gm_signature
    for ri, row in enumerate(t3.rows):
        txt = (row.cells[0].text or "").strip()
        if "Approved" in txt and "{{ gm_signature }}" not in txt:
            _set_cell_text(row.cells[0], txt + " {{ gm_signature }}")
            break

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 4. Grievance Form
# ─────────────────────────────────────────────────────────────────────────────

def fix_grievance():
    filename = "Employee grievance disciplinary action-form.docx"
    _backup_current(filename)
    doc = Document(_source_path(filename))

    # Table 3: complaint description (1 row x 1 col)
    t3 = doc.tables[3]
    try:
        cell = t3.rows[0].cells[0]
        txt = (cell.text or "").strip()
        if "{{ complaint_description }}" not in txt:
            _set_cell_text(cell, txt + "\n{{ complaint_description }}")
    except (IndexError, AttributeError):
        pass

    # Table 5: Employee Acknowledgment
    t5 = doc.tables[5]
    try:
        cell = t5.rows[0].cells[0]
        txt = (cell.text or "").strip()
        if "{{ complainant_signature }}" not in txt:
            _set_cell_text(cell, txt + " {{ complainant_signature }}")
    except (IndexError, AttributeError):
        pass

    # Add HR/GM signatures table if not present
    has_hr_sig = any(
        "{{ hr_signature }}" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not has_hr_sig:
        table = _add_table(doc, 2, 4)
        table.rows[0].cells[0].text = "HR Signature:"
        table.rows[0].cells[1].text = "{{ hr_signature }}"
        table.rows[0].cells[2].text = "GM Signature:"
        table.rows[0].cells[3].text = "{{ gm_signature }}"
        table.rows[1].cells[0].text = "HR Remarks:"
        table.rows[1].cells[1].text = "{{ hr_remarks }}"
        table.rows[1].cells[2].text = "GM Remarks:"
        table.rows[1].cells[3].text = "{{ gm_remarks }}"

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 5. Interview Assessment Form
# ─────────────────────────────────────────────────────────────────────────────
# Rating rows map: row index (in table 1) → placeholder name
_INTERVIEW_RATINGS = {
    3: "rating_turnout",
    4: "rating_confidence",
    5: "rating_mental_alertness",
    6: "rating_maturity",
    7: "rating_communication",
    8: "rating_technical",
    9: "rating_training",
    10: "rating_experience",
    11: "rating_overall",
}


def fix_interview_assessment():
    filename = "Interview Assessment Form - INJAAZ.DOCX"
    _backup_current(filename)
    doc = Document(_source_path(filename))
    t1 = doc.tables[1]

    # Add candidate info rows before the instruction row (row 0)
    # We'll prepend a new table with candidate info
    # Check if a candidate info table already exists
    has_candidate = any(
        "candidate_name" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not has_candidate:
        # Insert a candidate info table BEFORE the existing table
        # python-docx doesn't support "insert before", so we add after header and move
        # Strategy: add a new table after the header table by appending to body
        from docx.oxml.ns import qn as _qn
        from copy import deepcopy

        # Build the candidate info table
        tbl = _add_table(doc, 4, 4)
        rows_data = [
            ("Candidate Name:", "candidate_name", "Interview Date:", "interview_date"),
            ("Position Applied:", "position_title", "Interviewer Name:", "interviewer_name"),
            ("Academic Qualification:", "academic_qualification", "Years of Experience:", "years_experience"),
            ("Current Job Title:", "current_job_title", "Nationality:", "nationality"),
        ]
        for ri, (l1, ph1, l2, ph2) in enumerate(rows_data):
            tbl.rows[ri].cells[0].text = l1
            tbl.rows[ri].cells[1].text = "{{ %s }}" % ph1
            tbl.rows[ri].cells[2].text = l2
            tbl.rows[ri].cells[3].text = "{{ %s }}" % ph2

        # Move the new table element to right after the header table (index 0)
        body = doc.element.body
        new_tbl_elem = tbl._tbl
        header_tbl_elem = doc.tables[0]._tbl
        # Insert after header table
        header_tbl_elem.addnext(new_tbl_elem)

    # Now inject rating placeholders into the rating rows of t1
    # t1 has 7 cols: Factor | Indicators | Outstanding | V.Good | Good | Fair | Low
    # For each factor row, col 6 (Low) is last - we inject into col 2 (Outstanding) 
    # by setting it to the placeholder which will render as the selected rating value
    # Actually better: collapse the 5 rating cols into 1 value that shows the chosen rating
    # We put {{ rating_X }} into col 2 (Outstanding) and clear cols 3-6
    t1 = doc.tables[1]  # re-fetch after potential table addition
    # But wait - adding a table may shift indices. Re-fetch.
    # Find the rating assessment table by looking for "Turn-out" in first col
    rating_tbl = None
    for tbl in doc.tables:
        for row in tbl.rows:
            if "Turn-out" in (row.cells[0].text or ""):
                rating_tbl = tbl
                break
        if rating_tbl:
            break

    if rating_tbl:
        for ri, row in enumerate(rating_tbl.rows):
            factor = (row.cells[0].text or "").strip()
            ph_map = {
                "Turn-out": "rating_turnout",
                "Confidence": "rating_confidence",
                "Mental Alertness": "rating_mental_alertness",
                "Maturity": "rating_maturity",
                "Communication": "rating_communication",
                "Technical": "rating_technical",
                "Relevant Training": "rating_training",
                "Relevant Experience": "rating_experience",
                "Overall Rating": "rating_overall",
            }
            for key, ph in ph_map.items():
                if key in factor:
                    # Put placeholder in the last col (index 6)
                    cell = row.cells[6]
                    if "{{ %s }}" % ph not in (cell.text or ""):
                        _set_cell_text(cell, "{{ %s }}" % ph)
                    break

    # Add overall assessment / recommendation row if missing
    has_assess = any(
        "overall_assessment" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not has_assess and rating_tbl:
        row = rating_tbl.add_row()
        row.cells[0].text = "Overall Assessment / Comments:"
        row.cells[1].text = "{{ overall_assessment }}"
        row2 = rating_tbl.add_row()
        row2.cells[0].text = "Recommendation (Eligible for hire?):"
        row2.cells[1].text = "{{ eligibility }}"

    has_hr = any(
        "{{ hr_signature }}" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not has_hr:
        sig_tbl = _add_table(doc, 2, 4)
        sig_tbl.rows[0].cells[0].text = "Interviewer Signature:"
        sig_tbl.rows[0].cells[1].text = "{{ interviewer_signature }}"
        sig_tbl.rows[0].cells[2].text = "HR Signature:"
        sig_tbl.rows[0].cells[3].text = "{{ hr_signature }}"
        sig_tbl.rows[1].cells[0].text = "GM Signature:"
        sig_tbl.rows[1].cells[1].text = "{{ gm_signature }}"
        sig_tbl.rows[1].cells[2].text = "HR Remarks:"
        sig_tbl.rows[1].cells[3].text = "{{ hr_remarks }}"

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 6. Staff Appraisal Form — REBUILD (template is essentially empty)
# ─────────────────────────────────────────────────────────────────────────────

def fix_staff_appraisal():
    filename = "Staff Appraisal Form - INJAAZ.DOCX"
    _backup_current(filename)
    doc = Document(_source_path(filename))

    # Keep the header table (table 0), add all content tables
    # ── Employee Info ──
    info_tbl = _add_table(doc, 4, 4)
    info_rows = [
        ("Name:", "employee_name", "Employee ID:", "employee_id"),
        ("Department:", "department", "Position:", "position"),
        ("Appraisal Period:", "appraisal_period", "Reviewer:", "reviewer"),
        ("Date:", "review_date", "Designation:", "designation"),
    ]
    for ri, (l1, ph1, l2, ph2) in enumerate(info_rows):
        info_tbl.rows[ri].cells[0].text = l1
        info_tbl.rows[ri].cells[1].text = "{{ %s }}" % ph1
        info_tbl.rows[ri].cells[2].text = l2
        info_tbl.rows[ri].cells[3].text = "{{ %s }}" % ph2

    # ── Rating Table ──
    criteria = [
        ("Punctuality", "rating_punctuality", "comments_punctuality"),
        ("Job Knowledge", "rating_job_knowledge", "comments_job_knowledge"),
        ("Quality of Work", "rating_quality", "comments_quality"),
        ("Productivity", "rating_productivity", "comments_productivity"),
        ("Communication Skills", "rating_communication", "comments_communication"),
        ("Team Work", "rating_teamwork", "comments_teamwork"),
        ("Problem Solving", "rating_problem_solving", "comments_problem_solving"),
        ("Adaptability", "rating_adaptability", "comments_adaptability"),
        ("Leadership", "rating_leadership", "comments_leadership"),
    ]
    rating_tbl = _add_table(doc, 1 + len(criteria) + 1, 3)
    # Header
    rating_tbl.rows[0].cells[0].text = "Criterion"
    rating_tbl.rows[0].cells[1].text = "Rating (1-5)"
    rating_tbl.rows[0].cells[2].text = "Comments"
    for pi in range(3):
        for run in rating_tbl.rows[0].cells[pi].paragraphs[0].runs:
            run.bold = True
    # Criteria rows
    for ri, (lbl, ph_r, ph_c) in enumerate(criteria, start=1):
        rating_tbl.rows[ri].cells[0].text = lbl
        rating_tbl.rows[ri].cells[1].text = "{{ %s }}" % ph_r
        rating_tbl.rows[ri].cells[2].text = "{{ %s }}" % ph_c
    # Total Score row
    total_ri = len(criteria) + 1
    rating_tbl.rows[total_ri].cells[0].text = "Total Score:"
    rating_tbl.rows[total_ri].cells[1].text = "{{ total_score }}"

    # ── Strengths / Improvement ──
    summary_tbl = _add_table(doc, 2, 2)
    summary_tbl.rows[0].cells[0].text = "Employee Strengths:"
    summary_tbl.rows[0].cells[1].text = "{{ employee_strengths }}"
    summary_tbl.rows[1].cells[0].text = "Areas for Improvement:"
    summary_tbl.rows[1].cells[1].text = "{{ areas_for_improvement }}"

    # ── Signatures ──
    sig_tbl = _add_table(doc, 2, 4)
    sig_tbl.rows[0].cells[0].text = "Employee Signature:"
    sig_tbl.rows[0].cells[1].text = "{{ employee_signature }}"
    sig_tbl.rows[0].cells[2].text = "Evaluator Signature:"
    sig_tbl.rows[0].cells[3].text = "{{ evaluator_signature }}"
    sig_tbl.rows[1].cells[0].text = "HR Signature:"
    sig_tbl.rows[1].cells[1].text = "{{ hr_signature }}"
    sig_tbl.rows[1].cells[2].text = "GM Signature:"
    sig_tbl.rows[1].cells[3].text = "{{ gm_signature }}"

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 7. Station Clearance Form
# ─────────────────────────────────────────────────────────────────────────────

def fix_station_clearance():
    filename = "Station Clearance Form - INJAAZ.DOCX"
    _backup_current(filename)
    doc = Document(_source_path(filename))
    t = doc.tables[1]  # main content table

    # Add last_working_date if missing
    lwd_found = any(
        "last_working" in (c.text or "").lower()
        for row in t.rows for c in row.cells
    )
    if not lwd_found:
        # Add to row 3 (after type_of_departure)
        # R3: 'Type of Departure...' - find it and append date after
        for ri, row in enumerate(t.rows):
            txt = (row.cells[0].text or "").strip()
            if "Type of Departure" in txt and "{{ type_of_departure }}" not in txt:
                _set_cell_text(row.cells[0], txt + " {{ type_of_departure }}")
            if "Last Working" in txt and "{{ last_working_date }}" not in txt:
                _set_cell_text(row.cells[0], txt + " {{ last_working_date }}")

    # The checklist items should be in the table
    # Map of label substring → placeholder
    checklist_map = {
        "tasks handed over": "tasks_handed_over",
        "documents handed over": "documents_handed_over",
        "files handed over": "files_handed_over",
        "keys returned": "keys_returned",
        "toolbox returned": "toolbox_returned",
        "access card returned": "access_card_returned",
        "e-mail cancelled": "email_cancelled",
        "email cancelled": "email_cancelled",
        "software": "software_hardware_returned",
        "laptop returned": "laptop_returned",
        "file shifted": "file_shifted",
        "dues paid": "dues_paid",
        "medical card": "medical_card_returned",
        "eos": "eos_transfer",
    }
    for row in t.rows:
        for ci, cell in enumerate(row.cells):
            txt = (cell.text or "").lower().strip()
            for label, ph in checklist_map.items():
                if label in txt and "{{ %s }}" % ph not in (cell.text or ""):
                    # Inject into the NEXT cell (value/status cell)
                    try:
                        next_cell = row.cells[ci + 1]
                        if next_cell is not cell and "{{ " not in (next_cell.text or ""):
                            _set_cell_text(next_cell, "{{ %s }}" % ph)
                        break
                    except IndexError:
                        pass

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 8. Visa Renewal Form — paragraph-based injection
# ─────────────────────────────────────────────────────────────────────────────

def fix_visa_renewal():
    filename = "Visa Renewal Form - INJAAZ.DOCX"
    _backup_current(filename)
    doc = Document(_source_path(filename))

    # Paragraph text replacements: (search_substring, replacement_text)
    BLANKS = "___________________________"
    SHORT_BLANK = "_________"

    for para in doc.paragraphs:
        txt = para.text or ""
        if "I, " in txt and BLANKS in txt and "with ID number" in txt:
            new = re.sub(r'_{5,}', lambda m: (
                "{{ employee_name }}" if m.start() == txt.index(BLANKS)
                else "{{ employee_id }}" if "ID number" in txt[:m.start()]
                else "{{ employer }}"
            ), txt)
            # More reliable replacement
            new = txt
            new = new.replace(BLANKS, "{{ employee_name }}", 1)
            new = new.replace(SHORT_BLANK, "{{ employee_id }}", 1)
            rest = txt[txt.index(BLANKS) + len(BLANKS):]
            new = txt.replace(BLANKS, "{{ employee_name }}", 1)
            if BLANKS in new:
                new = new.replace(BLANKS, "{{ employer }}", 1)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new
            else:
                para.add_run(new)

        elif "in the position " in txt and "have completed" in txt:
            new = txt
            new = new.replace(BLANKS, "{{ position }}", 1)
            new = re.sub(r'_{5,}', "{{ years_completed }}", new, count=1)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new
            else:
                para.add_run(new)

        elif "Signature of Employee" in txt and "{{ employee_signature }}" not in txt:
            new = txt + " {{ employee_signature }}"
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new
            else:
                para.add_run(new)

    # Add HR/GM signature table
    has_hr = any(
        "{{ hr_signature }}" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not has_hr:
        sig_tbl = _add_table(doc, 2, 4)
        sig_tbl.rows[0].cells[0].text = "HR Signature:"
        sig_tbl.rows[0].cells[1].text = "{{ hr_signature }}"
        sig_tbl.rows[0].cells[2].text = "GM Signature:"
        sig_tbl.rows[0].cells[3].text = "{{ gm_signature }}"
        sig_tbl.rows[1].cells[0].text = "Decision:"
        sig_tbl.rows[1].cells[1].text = "{{ decision_display }}"
        sig_tbl.rows[1].cells[2].text = "HR Remarks:"
        sig_tbl.rows[1].cells[3].text = "{{ hr_remarks }}"

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# 9. Contract Renewal Form
# ─────────────────────────────────────────────────────────────────────────────

def fix_contract_renewal():
    filename = "Employee Contract Renewal Assessment Form Word.docx"
    _backup_current(filename)
    doc = Document(_source_path(filename))
    t2 = doc.tables[2]  # Rating table (33 rows x 7 cols)

    # Rating rows: section 01 = R4-R8, section 02 = R10-R14, section 03 = R16-R20, section 04 = R22-R24
    # Col 6 is the Rating column (last col)
    rating_map = {
        # (row_index, sub_letter): placeholder
        (4, "a"): "rating_01a", (5, "b"): "rating_01b", (6, "c"): "rating_01c",
        (7, "d"): "rating_01d", (8, "e"): "rating_01e",
        (10, "a"): "rating_02a", (11, "b"): "rating_02b", (12, "c"): "rating_02c",
        (13, "d"): "rating_02d", (14, "e"): "rating_02e",
        (16, "a"): "rating_03a", (17, "b"): "rating_03b", (18, "c"): "rating_03c",
        (19, "d"): "rating_03d", (20, "e"): "rating_03e",
        (22, "a"): "rating_04a", (23, "b"): "rating_04b", (24, "c"): "rating_04c",
    }
    for (ri, _), ph in rating_map.items():
        try:
            cell = t2.rows[ri].cells[6]
            if "{{ %s }}" % ph not in (cell.text or ""):
                _set_cell_text(cell, "{{ %s }}" % ph)
        except IndexError:
            pass

    # R27: Areas for Improvement row - inject into value cell
    try:
        r27 = t2.rows[27]
        for ci in range(1, len(r27.cells)):
            cell = r27.cells[ci]
            if not (cell.text or "").strip() or "{{ " not in (cell.text or ""):
                # Find the first empty cell after the label
                first_empty = None
                for ci2 in range(1, len(r27.cells)):
                    c = r27.cells[ci2]
                    if not (c.text or "").strip():
                        first_empty = ci2
                        break
                if first_empty:
                    _set_cell_text(r27.cells[first_empty], "{{ areas_for_improvement }}")
                else:
                    # Append to label cell
                    txt = (r27.cells[1].text or "").strip()
                    if "{{ areas_for_improvement }}" not in txt:
                        _set_cell_text(r27.cells[1], txt + " {{ areas_for_improvement }}")
                break
    except (IndexError, AttributeError):
        pass

    # R32: Evaluator signature row → inject evaluator_signature
    try:
        r32 = t2.rows[32]
        txt = (r32.cells[0].text or "").strip()
        if "Evaluator" in txt and "{{ evaluator_signature }}" not in txt:
            _set_cell_text(r32.cells[0], "Evaluator Signature: {{ evaluator_signature }}  Date: {{ evaluator_date }}")
    except (IndexError, AttributeError):
        pass

    # Add HR/GM signature table if missing
    has_hr = any(
        "{{ hr_signature }}" in (c.text or "")
        for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    if not has_hr:
        sig_tbl = _add_table(doc, 2, 4)
        sig_tbl.rows[0].cells[0].text = "HR Signature:"
        sig_tbl.rows[0].cells[1].text = "{{ hr_signature }}"
        sig_tbl.rows[0].cells[2].text = "GM Signature:"
        sig_tbl.rows[0].cells[3].text = "{{ gm_signature }}"
        sig_tbl.rows[1].cells[0].text = "HR Remarks:"
        sig_tbl.rows[1].cells[1].text = "{{ hr_remarks }}"
        sig_tbl.rows[1].cells[2].text = "GM Remarks:"
        sig_tbl.rows[1].cells[3].text = "{{ gm_remarks }}"

    doc.save(_dest_path(filename))
    print(f"  Fixed: {filename}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(TMPL_DIR, exist_ok=True)
    print("Fixing HR DOCX templates...")
    fix_leave_application()
    fix_duty_resumption()
    fix_passport_release()
    fix_grievance()
    fix_interview_assessment()
    fix_staff_appraisal()
    fix_station_clearance()
    fix_visa_renewal()
    fix_contract_renewal()
    print("\nAll done. Run  python scripts/scan_hr_placeholders.py  to verify.")


if __name__ == "__main__":
    main()
