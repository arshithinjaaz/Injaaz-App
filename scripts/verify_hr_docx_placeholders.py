"""
Verify HR DOCX placeholder placement: ensure {{ placeholder }} is in the correct cell
(next to its label) and report any layout issues.

Run from project root: python scripts/verify_hr_docx_placeholders.py
"""
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
HR_DOCS = os.path.join(BASE, "HR Documents")

# From inject_all_hr_placeholders
FORMS = [
    ("Leave Application Form - INJAAZ.DOCX", "leave_application"),
    ("Duty Resumption Form - INJAAZ.DOCX", "duty_resumption"),
    ("Passport Release & Submission Form - INJAAZ.DOCX", "passport_release"),
    ("Employee grievance disciplinary action-form.docx", "grievance"),
    ("Interview Assessment Form - INJAAZ.DOCX", "interview_assessment"),
    ("Staff Appraisal Form - INJAAZ.DOCX", "staff_appraisal"),
    ("Station Clearance Form - INJAAZ.DOCX", "station_clearance"),
    ("Visa Renewal Form - INJAAZ.DOCX", "visa_renewal"),
    ("Employee Contract Renewal Assessment Form Word.docx", "contract_renewal"),
]

PLACEHOLDER_RE = re.compile(r'\{\{\s*(\w+)\s*\}\}')


def _is_header_table(table):
    return len(table.rows) == 1 and len(table.rows[0].cells) == 2


def verify_docx(doc_path):
    """Verify placeholder placement in a DOCX. Returns (ok, issues, report_lines)."""
    if not os.path.isfile(doc_path):
        return False, ["File not found"], []
    doc = Document(doc_path)
    issues = []
    report = []
    tables_start = 0
    if doc.tables and _is_header_table(doc.tables[0]):
        tables_start = 1
        report.append("  (Skipping header table 0)")

    for ti in range(tables_start, len(doc.tables)):
        table = doc.tables[ti]
        report.append(f"\n  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows):
            cells = row.cells
            for ci, cell in enumerate(cells):
                text = (cell.text or "").strip()
                placeholders = PLACEHOLDER_RE.findall(text)
                if placeholders:
                    # Check: placeholder should be in value cell (right of label)
                    prev_cell_text = (cells[ci - 1].text or "").strip() if ci > 0 else ""
                    report.append(f"    Row {ri}, Cell {ci}: {placeholders}")
                    report.append(f"      Cell content: {repr(text[:60])}")
                    if ci > 0:
                        report.append(f"      Prev cell (label): {repr(prev_cell_text[:50])}")
                        # Warn only if prev cell has placeholder but NO label (merged cells show same content - OK if label present)
                        if PLACEHOLDER_RE.search(prev_cell_text):
                            # Check if label is present (text before placeholder)
                            for ph in placeholders:
                                idx = prev_cell_text.find("{{ " + ph + " }}")
                                if idx > 2 and prev_cell_text[:idx].strip():  # label before placeholder
                                    break  # OK - label preserved
                            else:
                                issues.append(f"Table {ti} Row {ri} Cell {ci}: placeholder may be in wrong cell (prev also has placeholder)")
                    if ci == 0:
                        # OK if cell has label before placeholder (merged cell)
                        has_label = any(ph in text and text.find("{{ " + ph + " }}") > 2 for ph in placeholders)
                        if not has_label:
                            issues.append(f"Table {ti} Row {ri} Cell 0: placeholder in first cell (no label to left)")

    return len(issues) == 0, issues, report


def main():
    print("=" * 70)
    print("HR DOCX Placeholder Verification")
    print("=" * 70)
    all_ok = True
    for filename, form_type in FORMS:
        path = os.path.join(HR_DOCS, filename)
        print(f"\n--- {filename} ---")
        ok, issues, report = verify_docx(path)
        for line in report:
            print(line)
        if issues:
            all_ok = False
            print("  ISSUES:")
            for i in issues:
                print(f"    ! {i}")
        else:
            print("  OK")
    print("\n" + "=" * 70)
    print("PASS" if all_ok else "REVIEW NEEDED - some placeholders may be misplaced")
    return 0 if all_ok else 1


if __name__ == "__main__":
    sys.exit(main())
