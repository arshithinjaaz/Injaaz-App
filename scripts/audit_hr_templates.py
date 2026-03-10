"""
Deep audit of every HR DOCX template: print table structure with cell text,
merged cell info, and placeholder positions so we can see alignment issues.
"""
import os, sys, re
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

HR_DOCS = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "HR Documents")
PH = re.compile(r'\{\{\s*(\w+)\s*\}\}')

FORMS = [
    "Leave Application Form - INJAAZ.DOCX",
    "Commencement Form - INJAAZ.DOCX",
    "Duty Resumption Form - INJAAZ.DOCX",
    "Passport Release & Submission Form - INJAAZ.DOCX",
    "Employee grievance disciplinary action-form.docx",
    "Interview Assessment Form - INJAAZ.DOCX",
    "Staff Appraisal Form - INJAAZ.DOCX",
    "Station Clearance Form - INJAAZ.DOCX",
    "Visa Renewal Form - INJAAZ.DOCX",
    "Employee Performance Evaluation Form - INJAAZ.DOCX",
    "Employee Contract Renewal Assessment Form Word.docx",
]

TARGET = sys.argv[1] if len(sys.argv) > 1 else None

for fname in FORMS:
    if TARGET and TARGET.lower() not in fname.lower():
        continue
    path = os.path.join(HR_DOCS, fname)
    if not os.path.exists(path):
        print(f"MISSING: {fname}"); continue
    doc = Document(path)
    print(f"\n{'='*80}")
    print(f"  {fname}")
    print(f"  Tables: {len(doc.tables)}  Paragraphs: {len(doc.paragraphs)}")
    print(f"{'='*80}")

    for pi, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t:
            phs = PH.findall(t)
            marker = f"  [PH: {phs}]" if phs else ""
            print(f"  PARA {pi}: {repr(t[:100])}{marker}")

    for ti, tbl in enumerate(doc.tables):
        ncols = len(tbl.columns)
        print(f"\n  TABLE {ti}: {len(tbl.rows)}R x {ncols}C")
        for ri, row in enumerate(tbl.rows):
            cells = row.cells
            # Detect merged cells by comparing cell._tc elements
            seen_tc = {}
            parts = []
            for ci, cell in enumerate(cells):
                tc_id = id(cell._tc)
                txt = (cell.text or "").strip().replace("\n", "\\n")[:55]
                phs = PH.findall(cell.text or "")
                ph_tag = f" [{','.join(phs)}]" if phs else ""
                if tc_id in seen_tc:
                    parts.append(f"C{ci}=C{seen_tc[tc_id]}")
                else:
                    seen_tc[tc_id] = ci
                    parts.append(f"C{ci}: {repr(txt)}{ph_tag}")
            # Only print rows with content
            has_content = any((cell.text or "").strip() for cell in cells)
            if has_content:
                print(f"    R{ri}: " + " | ".join(parts))
