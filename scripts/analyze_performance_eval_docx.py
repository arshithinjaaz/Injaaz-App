"""Analyze Performance Evaluation DOCX structure for placeholder injection."""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
path = os.path.join(base, "HR Documents", "Employee Performance Evaluation Form - INJAAZ.DOCX")
doc = Document(path)
print("Paragraphs:", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs[:25]):
    print(f"  {i}: {repr(p.text.strip()[:70])}")
print("\nTables:", len(doc.tables))
for ti, table in enumerate(doc.tables):
    print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:15]):
        cells = [c.text.strip()[:40] for c in row.cells]
        print(f"    Row {ri}: {cells}")
