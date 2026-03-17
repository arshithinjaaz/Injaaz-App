"""
Apply standard HR header to all HR form DOCX files: headline left, logo right.
Uses same format as Commencement Form (headline 20pt, subheading 9pt, logo right).

Run from project root: python scripts/apply_hr_header_all_forms.py
"""
import os
import sys
import shutil

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_margins(cell, **kwargs):
    """Set table cell margins. Values in twentieths of a point (dxa)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m in ("top", "start", "bottom", "end"):
        if m in kwargs:
            node = OxmlElement(f"w:{m}")
            node.set(qn("w:w"), str(kwargs[m]))
            node.set(qn("w:type"), "dxa")
            tcMar.append(node)
    tcPr.append(tcMar)


# (filename, headline, intro_search_start, intro_default)
# intro_search_start: search doc for paragraph starting with this; if found, use as subheading and remove from body
# intro_default: if no intro found, use this as subheading (or None for headline only)
HR_FORMS_CONFIG = [
    ("Commencement Form - INJAAZ.DOCX", "Commencement Form",
     "To complete the administrative aspect",
     "To complete the administrative aspect of your Employment please complete this form within 5 days of joining and fax it back to AH or email it to joana@ajmanholding.ae"),
    ("Leave Application Form - INJAAZ.DOCX", "Leave Application Form", None, None),
    ("Duty Resumption Form - INJAAZ.DOCX", "Duty Resumption Form", None, None),
    ("Passport Release & Submission Form - INJAAZ.DOCX", "Passport Release & Submission Form", None, None),
    ("Employee grievance disciplinary action-form.docx", "Employee Grievance / Disciplinary Action Form", None, None),
    ("Employee Performance Evaluation Form - INJAAZ.DOCX", "Employee Performance Evaluation Form", None, None),
    ("Interview Assessment Form - INJAAZ.DOCX", "Interview Assessment Form", None, None),
    ("Staff Appraisal Form - INJAAZ.DOCX", "Staff Appraisal Form", None, None),
    ("Station Clearance Form - INJAAZ.DOCX", "Station Clearance Form", None, None),
    ("Visa Renewal Form - INJAAZ.DOCX", "Visa Renewal Form", None, None),
    ("Employee Contract Renewal Assessment Form Word.docx", "Employee Contract Renewal Assessment Form", None, None),
]


def add_header_table_to_doc(doc_path, logo_path, headline, intro_search_start=None, intro_default=None,
                             backup=True, backup_subdir="templates"):
    """
    Add or replace header table: left = headline (+ optional intro), right = logo.
    intro_search_start: if set, find first paragraph starting with this and use as subheading (and remove from body).
    intro_default: if no intro found from search, use this as subheading; if None and no search, headline only.
    """
    if not os.path.isfile(doc_path):
        raise FileNotFoundError(f"Template not found: {doc_path}")
    if not os.path.isfile(logo_path):
        raise FileNotFoundError(f"Logo not found: {logo_path}")

    base_name = os.path.basename(doc_path)
    if backup:
        backup_dir = os.path.join(os.path.dirname(doc_path), backup_subdir)
        os.makedirs(backup_dir, exist_ok=True)
        backup_path = os.path.join(backup_dir, base_name.replace(".docx", " (before header).docx").replace(".DOCX", " (before header).docx"))
        shutil.copy2(doc_path, backup_path)
        print(f"  Backed up to: {backup_path}")

    doc = Document(doc_path)
    body = doc._element.body

    intro_text = None
    intro_para_element = None
    if intro_search_start:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith(intro_search_start):
                intro_text = text
                intro_para_element = para._element
                break
    if intro_text is None and intro_default:
        intro_text = intro_default

    had_existing_table = False
    first = body[0] if len(body) else None
    if first is not None:
        tag = first.tag.split("}")[-1] if "}" in first.tag else first.tag
        if tag == "tbl":
            body.remove(first)
            had_existing_table = True

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    try:
        table.rows[0].cells[0].width = Cm(13.5)
        table.rows[0].cells[1].width = Cm(3.5)
    except Exception:
        pass
    FORM_FONT = "Calibri"
    left_cell = table.rows[0].cells[0]
    left_cell.text = ""

    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.paragraph_format.keep_together = True
    p_left.paragraph_format.space_after = Pt(0)
    run = p_left.add_run(headline)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = FORM_FONT

    if intro_text:
        p_intro = left_cell.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_intro.paragraph_format.space_before = Pt(0)
        p_intro.paragraph_format.space_after = Pt(12)
        run_intro = p_intro.add_run(intro_text)
        run_intro.font.name = FORM_FONT
        run_intro.font.size = Pt(9)

    right_cell = table.rows[0].cells[1]
    for p in right_cell.paragraphs:
        p.clear()
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    run_right = p_right.add_run()
    run_right.add_picture(logo_path, width=Cm(1.8))

    set_cell_margins(left_cell, top=60, bottom=0)
    set_cell_margins(right_cell, top=60, bottom=0)
    try:
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "240")
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)
    except Exception:
        pass

    tbl_element = table._tbl
    body.remove(tbl_element)
    body.insert(0, tbl_element)

    if not had_existing_table:
        to_remove = []
        for i in range(1, min(8, len(body))):
            child = body[i]
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                to_remove.append(child)
        if intro_para_element is not None and intro_para_element in body:
            to_remove.append(intro_para_element)
        for el in to_remove:
            if el in body:
                body.remove(el)

    if doc.paragraphs:
        doc.paragraphs[0].paragraph_format.space_before = Pt(0)

    doc.save(doc_path)
    print(f"  Updated: {doc_path}")


def process_one(doc_path, logo_path, headline, intro_search_start, intro_default, backup):
    """Apply header to one document."""
    add_header_table_to_doc(
        doc_path, logo_path,
        headline=headline,
        intro_search_start=intro_search_start,
        intro_default=intro_default,
        backup=backup,
    )


def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    hr_docs = os.path.join(base, "HR Documents")
    logo_path = os.path.join(base, "static", "logo.png")
    if not os.path.isfile(logo_path):
        print("Logo not found at static/logo.png")
        return

    backup = "--no-backup" not in sys.argv
    for filename, headline, intro_search_start, intro_default in HR_FORMS_CONFIG:
        doc_path = os.path.join(hr_docs, filename)
        if not os.path.isfile(doc_path):
            print(f"Skip (not found): {filename}")
            continue
        print(f"Processing: {filename}")
        try:
            process_one(doc_path, logo_path, headline, intro_search_start, intro_default, backup)
        except Exception as e:
            print(f"  Error: {e}")

    print("Done.")


if __name__ == "__main__":
    main()
