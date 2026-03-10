"""Audit the 'before placeholders' backup templates to see original label positions."""
import os, sys, re
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document

TMPL = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "HR Documents", "templates")

TARGET = sys.argv[1] if len(sys.argv) > 1 else None

for fname in sorted(os.listdir(TMPL)):
    if "before placeholders" not in fname:
        continue
    if TARGET and TARGET.lower() not in fname.lower():
        continue
    path = os.path.join(TMPL, fname)
    doc = Document(path)
    print(f"\n{'='*80}\n  {fname}\n{'='*80}")
    for ti, tbl in enumerate(doc.tables):
        ncols = len(tbl.columns)
        print(f"\n  TABLE {ti}: {len(tbl.rows)}R x {ncols}C")
        for ri, row in enumerate(tbl.rows):
            cells = row.cells
            seen_tc = {}
            parts = []
            for ci, cell in enumerate(cells):
                tc_id = id(cell._tc)
                txt = (cell.text or "").strip().replace("\n", "\\n")[:55]
                if tc_id in seen_tc:
                    parts.append(f"C{ci}=C{seen_tc[tc_id]}")
                else:
                    seen_tc[tc_id] = ci
                    parts.append(f"C{ci}: {repr(txt)}")
            has_content = any((cell.text or "").strip() for cell in cells)
            if has_content:
                print(f"    R{ri}: " + " | ".join(parts))
    for pi, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t:
            print(f"  PARA {pi}: {repr(t[:100])}")
