"""
Inject docxtpl placeholders into the shared Commencement Form DOCX
so that UI form data is filled into HR Documents/Commencement Form - INJAAZ.DOCX.

The underline ("box") is on the LABEL paragraph in Word, not the next one.
We put placeholders in the same paragraph as the label so the value appears ON the line.

Run from project root: python scripts/inject_commencement_placeholders.py
"""
import os
import sys
import shutil

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document


def inject_commencement_placeholders(doc_path, backup=True):
    """
    Put placeholders in the LABEL paragraph (the one that has the bottom border/underline).
    Clear the next paragraph so the value does not appear again below the line.
    """
    if not os.path.isfile(doc_path):
        raise FileNotFoundError(f"Template not found: {doc_path}")

    if backup:
        backup_dir = os.path.join(os.path.dirname(doc_path), "templates")
        os.makedirs(backup_dir, exist_ok=True)
        backup_path = os.path.join(backup_dir, os.path.basename(doc_path))
        if os.path.isfile(backup_path):
            os.remove(backup_path)
        shutil.copy2(doc_path, backup_path)
        print(f"Backed up original to: {backup_path}")

    doc = Document(doc_path)

    # (label_paragraph_index, placeholders, clear_next_n) - indices after header table (heading left, logo right)
    inserts = [
        (3, ["{{ employee_name }}"], 1),
        (5, ["{{ position }}"], 1),
        (7, ["{{ contacts }}"], 1),
        (9, ["{{ department }}"], 1),
        (11, ["{{ organization }}"], 1),
        (13, ["{{ date_of_joining }}"], 1),
        (16, ["{{ bank_name }}"], 1),
        (18, ["{{ bank_branch }}"], 1),
        (20, ["{{ account_number }}"], 1),
        (22, ["{{ employee_signature }}", "Date: {{ employee_sign_date }}"], 2),
        (26, ["{{ reporting_to_name }}"], 1),
        (28, ["{{ reporting_to_designation }}"], 1),
        (30, ["{{ reporting_to_contact }}"], 1),
        (32, ["{{ reporting_to_signature }}", "Date: {{ reporting_sign_date }}"], 2),
    ]

    for idx, placeholder_texts, clear_next_n in inserts:
        if idx >= len(doc.paragraphs):
            print(f"Warning: paragraph index {idx} out of range, skipping")
            continue
        label_para = doc.paragraphs[idx]
        # Add placeholder on same line as label so signature and date appear next to label, not below
        placeholder_str = " " + " ".join(placeholder_texts)
        label_para.add_run(placeholder_str)
        # Clear the next paragraph(s) so value does not appear again below
        for k in range(1, clear_next_n + 1):
            next_idx = idx + k
            if next_idx < len(doc.paragraphs):
                doc.paragraphs[next_idx].clear()

    doc.save(doc_path)
    print(f"Inserted placeholders on label lines (inside box) and saved: {doc_path}")


def main():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    hr_docs = os.path.join(base, "HR Documents")
    doc_path = os.path.join(hr_docs, "Commencement Form - INJAAZ.DOCX")
    inject_commencement_placeholders(doc_path, backup=True)


if __name__ == "__main__":
    main()
