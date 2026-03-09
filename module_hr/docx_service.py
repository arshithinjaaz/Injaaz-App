"""
Populate HR DOCX templates with form data from the UI.
Uses docxtpl to fill placeholders in HR Documents (e.g. Commencement Form - INJAAZ.DOCX).
"""
import os
import base64
import logging
import tempfile
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)


def fit_docx_to_one_page(source_stream):
    """
    Make document compact: reduce margins, paragraph spacing, table padding
    so it fits on one sheet with minimal white space.
    """
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    source_stream.seek(0)
    doc = Document(source_stream)

    # Compact margins
    for section in doc.sections:
        section.top_margin = Cm(0.6)
        section.bottom_margin = Cm(0.6)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    FORM_FONT = "Calibri"
    BODY_SIZE = Pt(9)
    for para in doc.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.line_spacing = Pt(11)
        for run in para.runs:
            run.font.name = FORM_FONT
            if run.font.size is None or run.font.size.pt < 14:
                run.font.size = BODY_SIZE

    # Tables: compact padding, font, and proper alignment (skip first table = header)
    for ti, table in enumerate(doc.tables):
        if ti == 0:
            continue  # Preserve header: logo right alignment, headline size
        for row in table.rows:
            for cell in row.cells:
                # Reduce cell padding (values in dxa: 50 = ~1.8pt)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                existing = tcPr.find(qn("w:tcMar"))
                if existing is not None:
                    tcPr.remove(existing)
                tcMar = OxmlElement("w:tcMar")
                for side in ("top", "start", "bottom", "end"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:w"), "50")
                    el.set(qn("w:type"), "dxa")
                    tcMar.append(el)
                tcPr.append(tcMar)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    pf = para.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    for run in para.runs:
                        run.font.name = FORM_FONT
                        if run.font.size and run.font.size.pt > 10:
                            run.font.size = Pt(9)
                # Remove any coloured background so cells are plain (no tint)
                existing_shd = tcPr.find(qn("w:shd"))
                if existing_shd is not None:
                    tcPr.remove(existing_shd)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _get_docxtpl():
    """Lazy import so HR module loads even if docxtpl is not installed (DOCX download will fail with clear error)."""
    try:
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Mm
        return DocxTemplate, InlineImage, Mm
    except ImportError as e:
        raise ImportError(
            "docxtpl is required for HR DOCX download. Install with: pip install docxtpl"
        ) from e


def _get_hr_documents_path():
    """Path to HR Documents folder (project root)"""
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, 'HR Documents')


def _get_template_path(form_type):
    """Return template path for form type. Prefer shared file in HR Documents/ root, then templates/."""
    folder = _get_hr_documents_path()
    templates_sub = os.path.join(folder, 'templates')
    templates = {
        'commencement': 'Commencement Form - INJAAZ.DOCX',
        'leave_application': 'Leave Application Form - INJAAZ.DOCX',
        'leave': 'Leave Application Form - INJAAZ.DOCX',
        'duty_resumption': 'Duty Resumption Form - INJAAZ.DOCX',
        'passport_release': 'Passport Release & Submission Form - INJAAZ.DOCX',
        'grievance': 'Employee grievance disciplinary action-form.docx',
        'performance_evaluation': 'Employee Performance Evaluation Form - INJAAZ.DOCX',
        'interview_assessment': 'Interview Assessment Form - INJAAZ.DOCX',
        'staff_appraisal': 'Staff Appraisal Form - INJAAZ.DOCX',
        'station_clearance': 'Station Clearance Form - INJAAZ.DOCX',
        'visa_renewal': 'Visa Renewal Form - INJAAZ.DOCX',
        'contract_renewal': 'Employee Contract Renewal Assessment Form Word.docx',
    }
    name = templates.get(form_type)
    if not name:
        return None
    for base in (folder, templates_sub):
        path = os.path.join(base, name)
        if os.path.isfile(path):
            return path
    return None


def _fmt_date(v):
    if not v:
        return '-'
    try:
        d = datetime.fromisoformat(str(v).replace('Z', '+00:00'))
        return d.strftime('%d/%m/%Y')
    except (ValueError, TypeError):
        return str(v)


def _normalize_form_data_for_docx(form_data, form_type):
    """
    Apply field aliases so UI form field names map to DOCX template placeholders.
    Ensures all HR form data maps correctly to Word template placeholders.
    """
    fd = dict(form_data or {})

    # --- Leave Application ---
    if form_type in ('leave', 'leave_application'):
        if fd.get('start_date') and not fd.get('first_day_of_leave'):
            fd['first_day_of_leave'] = fd['start_date']
        if fd.get('end_date') and not fd.get('last_day_of_leave'):
            fd['last_day_of_leave'] = fd['end_date']
        if fd.get('total_days') is not None and fd.get('total_days') != '' and not fd.get('total_days_requested'):
            fd['total_days_requested'] = fd['total_days']
        if fd.get('contact_number') and not fd.get('telephone_reachable'):
            fd['telephone_reachable'] = fd['contact_number']
        if not fd.get('today_date') and (fd.get('submitted_at') or fd.get('first_day_of_leave')):
            try:
                fd['today_date'] = datetime.now().strftime('%Y-%m-%d')
            except Exception:
                pass
        # Leave type: map value to display label for template
        leave_labels = {
            'annual': 'Annual Leave', 'sick': 'Sick Leave', 'ot_compensatory': 'OT Compensatory Off',
            'compassionate': 'Compassionate Leave', 'study': 'Study Leave', 'unpaid': 'Unpaid Leave',
            'examination': 'Examination Leave (UAE Nationals)', 'hajj': 'Hajj Leave', 'other': 'Other',
        }
        if fd.get('leave_type') in leave_labels:
            fd['leave_type_display'] = leave_labels[fd['leave_type']]
        elif fd.get('leave_type') == 'other' and fd.get('leave_type_other'):
            fd['leave_type_display'] = fd['leave_type_other']
        else:
            fd['leave_type_display'] = fd.get('leave_type') or '-'

    # --- Checkbox values: "on" -> "Completed" for station clearance and similar ---
    checkbox_keys = {
        'tasks_handed_over', 'documents_handed_over', 'files_handed_over', 'keys_returned',
        'toolbox_returned', 'access_card_returned', 'email_cancelled', 'software_hardware_returned',
        'laptop_returned', 'file_shifted', 'dues_paid', 'medical_card_returned', 'eos_transfer',
    }
    for k in checkbox_keys:
        if k in fd and fd[k] == 'on':
            fd[k] = 'Completed'

    # --- Grievance: complaint / complaint_description (template may use either) ---
    if form_type == 'grievance':
        complaint_text = fd.get('complaint_description') or fd.get('complaint') or '-'
        fd['complaint'] = complaint_text
        fd['complaint_description'] = complaint_text

    # --- Duty Resumption: ensure all template fields present ---
    if form_type == 'duty_resumption':
        if fd.get('company') and not fd.get('organization'):
            fd['organization'] = fd['company']

    # --- Passport Release: purpose_of_release, release_date ---
    if form_type == 'passport_release':
        if not fd.get('purpose_of_release'):
            fd['purpose_of_release'] = '-'

    # --- Contract renewal: map sub-ratings to section ratings for DOCX ---
    if form_type == 'contract_renewal':
        for sn, keys in [
            ('01', ['rating_01a', 'rating_01b', 'rating_01c', 'rating_01d', 'rating_01e']),
            ('02', ['rating_02a', 'rating_02b', 'rating_02c', 'rating_02d', 'rating_02e']),
            ('03', ['rating_03a', 'rating_03b', 'rating_03c', 'rating_03d', 'rating_03e']),
            ('04', ['rating_04a', 'rating_04b', 'rating_04c']),
        ]:
            vals = [float(fd.get(k) or 0) for k in keys if fd.get(k) and str(fd.get(k)).strip()]
            if vals:
                avg = round(sum(vals) / len(vals), 1)
                fd[f'rating_{sn}'] = int(avg) if avg == int(avg) else avg
            elif not fd.get(f'rating_{sn}'):
                fd[f'rating_{sn}'] = '-'
        if fd.get('areas_for_improvement') and not fd.get('areas_for_improvement_display'):
            fd['areas_for_improvement_display'] = fd['areas_for_improvement']

    # --- Interview Assessment: map rating values to display + field aliases ---
    if form_type == 'interview_assessment':
        if fd.get('position_applied') and not fd.get('position_title'):
            fd['position_title'] = fd['position_applied']
        rating_map = {'outstanding': 'Outstanding', 'v_good': 'V. Good', 'good': 'Good', 'fair': 'Fair', 'low': 'Low'}
        for k in ['rating_turnout', 'rating_confidence', 'rating_mental_alertness', 'rating_maturity',
                  'rating_communication', 'rating_technical', 'rating_training', 'rating_experience', 'rating_overall']:
            if fd.get(k) in rating_map:
                fd[f'{k}_display'] = rating_map[fd[k]]

    # --- Staff Appraisal: template uses appraisal_period, reviewer ---
    if form_type == 'staff_appraisal':
        if fd.get('review_period') and not fd.get('appraisal_period'):
            fd['appraisal_period'] = fd['review_period']
        if fd.get('evaluator_name') and not fd.get('reviewer'):
            fd['reviewer'] = fd['evaluator_name']

    # --- Station Clearance: type_of_departure display + field aliases ---
    if form_type == 'station_clearance':
        if fd.get('last_working_day') and not fd.get('last_working_date'):
            fd['last_working_date'] = fd['last_working_day']
        if fd.get('departure_reason') and not fd.get('type_of_departure'):
            fd['type_of_departure'] = fd['departure_reason']
        if fd.get('designation') and not fd.get('position'):
            fd['position'] = fd['designation']
        dep_map = {'transfer': 'Transfer', 'resignation': 'Resignation', 'termination': 'Termination', 'end_of_contract': 'End of Contract'}
        if fd.get('type_of_departure') in dep_map:
            fd['type_of_departure_display'] = dep_map[fd['type_of_departure']]
        else:
            fd['type_of_departure_display'] = fd.get('type_of_departure') or '-'

    # --- Visa Renewal: decision display ---
    if form_type == 'visa_renewal':
        dec_map = {'continue': 'Continue employment for next 2 years and willing to have visa renewed',
                   'discontinue': 'Discontinue service and require visa cancellation'}
        if fd.get('decision') in dec_map:
            fd['decision_display'] = dec_map[fd['decision']]
        else:
            fd['decision_display'] = fd.get('decision') or '-'

    return fd


def _build_generic_context(form_data, date_keys=None):
    """Build docxtpl context from form_data: all keys, dates formatted. date_keys = set of keys to format as date."""
    fd = form_data or {}
    date_keys = date_keys or set()
    if not date_keys:
        for k, v in fd.items():
            if v and ('date' in k.lower() or 'day' in k.lower() or 'start' in k.lower() or 'end' in k.lower() or 'joining' in k.lower() or 'returning' in k.lower() or 'release' in k.lower() or 'incident' in k.lower() or 'employment' in k.lower() or 'last_working' in k.lower() or 'form_date' in k.lower() or ('evaluation' in k.lower() and 'date' in k)):
                date_keys.add(k)
    skip = {'form_type', 'employee_signature', 'evaluator_signature', 'gm_signature', 'hr_signature', 'complainant_signature', 'interviewer_signature', 'replacement_signature', 'reporting_to_signature'}
    ctx = {}
    for k, v in fd.items():
        if k in skip:
            continue
        if k in date_keys or (v and isinstance(v, str) and len(v) >= 8 and v[:4].isdigit() and '-' in v):
            ctx[k] = _fmt_date(v)
        else:
            ctx[k] = v if v not in (None, '') else '-'
    # Alias workflow HR/GM fields so both placeholder names work in templates
    ctx['hr_remarks'] = ctx.get('hr_remarks') or ctx.get('hr_comments') or '-'
    ctx['gm_remarks'] = ctx.get('gm_remarks') or ctx.get('gm_comments') or '-'
    ctx['hr_comments'] = ctx.get('hr_comments') or ctx.get('hr_remarks') or '-'
    ctx['gm_comments'] = ctx.get('gm_comments') or ctx.get('gm_remarks') or '-'
    return ctx


def _add_signatures_to_context(tpl, ctx, form_data, signature_pairs):
    """signature_pairs = [(ctx_key, form_data_key), ...]. Returns tmp_paths."""
    tmp_paths = []
    for ctx_key, data_key in signature_pairs:
        img, path = _signature_to_inline_image(tpl, (form_data or {}).get(data_key))
        if path:
            tmp_paths.append(path)
        ctx[ctx_key] = img if img else '(Signed in original)'
    return tmp_paths


def _get_context_extra_for_form(form_type, fd):
    """
    Return extra context: aliases and derived fields for template placeholders.
    Only adds keys not already in form_data (generic context handles those).
    Values are not date-formatted here; generic context does that for fd keys.
    """
    def _val(v):
        return v if v not in (None, '') else '-'

    extra = {}
    if form_type in ('leave', 'leave_application'):
        extra['leave_type'] = _val(fd.get('leave_type_display') or fd.get('leave_type'))
        if not fd.get('telephone_reachable') and fd.get('mobile_no'):
            extra['telephone_reachable'] = fd['mobile_no']
    elif form_type == 'duty_resumption':
        pass  # employee_id, company, leave_ended are in fd, generic handles
    elif form_type == 'passport_release':
        pass  # purpose_of_release, release_date in fd
    elif form_type == 'grievance':
        extra['complaint'] = _val(fd.get('complaint_description') or fd.get('complaint'))
    elif form_type == 'interview_assessment':
        extra['rating_overall'] = _val(fd.get('rating_overall_display') or fd.get('rating_overall'))
    elif form_type == 'station_clearance':
        extra['type_of_departure'] = _val(fd.get('type_of_departure_display') or fd.get('type_of_departure'))
    elif form_type == 'contract_renewal':
        pass  # areas_for_improvement, comments_01-04 in fd
    elif form_type == 'staff_appraisal':
        pass  # all in fd
    elif form_type == 'visa_renewal':
        extra['decision'] = _val(fd.get('decision_display') or fd.get('decision'))
    return extra


def _generate_filled_docx_generic(form_type, form_data, output_path_or_stream, submission_id=None,
                                   context_extra=None, signature_pairs=None):
    """Load template, build context (generic + extra), add signatures, render, save."""
    template_path = _get_template_path(form_type)
    if not template_path:
        raise FileNotFoundError(f'Template not found for {form_type}')
    DocxTemplate, _, _ = _get_docxtpl()
    tpl = DocxTemplate(template_path)
    fd = _normalize_form_data_for_docx(form_data, form_type)
    ctx = _build_generic_context(fd)
    form_extra = _get_context_extra_for_form(form_type, fd)
    ctx.update(form_extra)
    if context_extra:
        ctx.update(context_extra)
    if submission_id:
        ctx['submission_id'] = submission_id
    tmp_paths = []
    if signature_pairs:
        tmp_paths = _add_signatures_to_context(tpl, ctx, form_data, signature_pairs)  # use original for signatures
    tpl.render(ctx)
    try:
        if hasattr(output_path_or_stream, 'write'):
            tpl.save(output_path_or_stream)
        else:
            tpl.save(output_path_or_stream)
    finally:
        for p in tmp_paths:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except OSError:
                pass


def _signature_to_transparent_png(data_url):
    """Convert base64 signature to PNG with transparent background (no white). Returns temp path or None."""
    if not data_url or not isinstance(data_url, str) or not data_url.startswith('data:image'):
        return None
    try:
        _, b64 = data_url.split(',', 1)
        raw = base64.b64decode(b64)
        try:
            from PIL import Image as PILImage
            pil = PILImage.open(BytesIO(raw)).convert('RGBA')
            data = pil.load()
            w, h = pil.size
            for y in range(h):
                for x in range(w):
                    r, g, b, a = data[x, y]
                    if r >= 250 and g >= 250 and b >= 250:
                        data[x, y] = (r, g, b, 0)
            f = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            pil.save(f, 'PNG')
            f.close()
            return f.name
        except Exception:
            f = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            f.write(raw)
            f.close()
            return f.name
    except Exception:
        return None


def _signature_to_inline_image(tpl, data_url, width_mm=22, height_mm=10):
    """Convert base64 data URL signature to InlineImage for docxtpl.
    Uses transparent PNG (no background colour). Returns (InlineImage, tmp_path).
    """
    if not data_url or not isinstance(data_url, str) or not data_url.startswith('data:image'):
        return None, None
    try:
        _, InlineImage, Mm = _get_docxtpl()
        tmp_path = _signature_to_transparent_png(data_url)
        if not tmp_path:
            return None, None
        img = InlineImage(tpl, tmp_path, width=Mm(width_mm), height=Mm(height_mm))
        return img, tmp_path
    except Exception:
        return None, None


def _build_commencement_context(form_data):
    """Build docxtpl context from Commencement form UI data."""
    fd = form_data or {}
    tpl = None  # Will be set when we have the template

    ctx = {
        'employee_name': fd.get('employee_name') or '-',
        'position': fd.get('position') or '-',
        'contacts': fd.get('contacts') or '-',
        'department': fd.get('department') or '-',
        'organization': fd.get('organization') or 'INJAAZ',
        'date_of_joining': _fmt_date(fd.get('date_of_joining')),
        'bank_name': fd.get('bank_name') or '-',
        'bank_branch': fd.get('bank_branch') or '-',
        'account_number': fd.get('account_number') or '-',
        'employee_sign_date': _fmt_date(fd.get('employee_sign_date')),
        'reporting_to_name': fd.get('reporting_to_name') or '-',
        'reporting_to_designation': fd.get('reporting_to_designation') or '-',
        'reporting_to_contact': fd.get('reporting_to_contact') or '-',
        'reporting_sign_date': _fmt_date(fd.get('reporting_sign_date')),
    }
    return ctx


def _add_commencement_signatures(tpl, ctx, form_data):
    """Add signature InlineImages to context. Returns list of temp file paths to delete after save."""
    fd = form_data or {}
    tmp_paths = []
    # Decent space for signature pad: 50mm x 22mm
    emp_sig, emp_path = _signature_to_inline_image(tpl, fd.get('employee_signature'), width_mm=50, height_mm=22)
    rep_sig, rep_path = _signature_to_inline_image(tpl, fd.get('reporting_to_signature'), width_mm=50, height_mm=22)
    if emp_path:
        tmp_paths.append(emp_path)
    if rep_path:
        tmp_paths.append(rep_path)
    if emp_sig:
        ctx['employee_signature'] = emp_sig
    else:
        ctx['employee_signature'] = '(Signed in original)'
    if rep_sig:
        ctx['reporting_to_signature'] = rep_sig
    else:
        ctx['reporting_to_signature'] = '(Signed in original)'
    return tmp_paths


def _build_performance_evaluation_context(form_data):
    """Build docxtpl context from Performance Evaluation form UI data."""
    fd = form_data or {}
    ctx = {
        'employee_name': fd.get('employee_name') or '-',
        'employee_id': fd.get('employee_id') or '-',
        'department': fd.get('department') or '-',
        'designation': fd.get('designation') or '-',
        'date_of_evaluation': _fmt_date(fd.get('date_of_evaluation')),
        'date_of_joining': _fmt_date(fd.get('date_of_joining')),
        'evaluation_done_by': fd.get('evaluation_done_by') or '-',
        'score_01': fd.get('score_01') or '-',
        'score_02': fd.get('score_02') or '-',
        'score_03': fd.get('score_03') or '-',
        'score_04': fd.get('score_04') or '-',
        'score_05': fd.get('score_05') or '-',
        'score_06': fd.get('score_06') or '-',
        'score_07': fd.get('score_07') or '-',
        'score_08': fd.get('score_08') or '-',
        'score_09': fd.get('score_09') or '-',
        'score_10': fd.get('score_10') or '-',
        'overall_score': fd.get('overall_score') or '-',
        'evaluator_name': fd.get('evaluator_name') or '-',
        'evaluator_designation': fd.get('evaluator_designation') or '-',
        'evaluator_observation': fd.get('evaluator_observation') or '-',
        'area_of_concern': fd.get('area_of_concern') or '-',
        'training_required': fd.get('training_required') or '-',
        'employee_comments': fd.get('employee_comments') or '-',
        'employee_sign_date': _fmt_date(fd.get('employee_sign_date')),
        'evaluator_sign_date': _fmt_date(fd.get('evaluator_sign_date')),
        'concern_incharge_name': fd.get('concern_incharge_name') or '-',
        'incharge_comments': fd.get('incharge_comments') or '-',
        'gm_remarks': fd.get('gm_remarks') or fd.get('gm_comments') or '-',
        'hr_remarks': fd.get('hr_remarks') or fd.get('hr_comments') or '-',
    }
    return ctx


def _add_performance_evaluation_signatures(tpl, ctx, form_data):
    """Add signature InlineImages to context. Returns list of temp file paths to delete after save."""
    fd = form_data or {}
    tmp_paths = []
    for key, data_key in [
        ('employee_signature', 'employee_signature'),
        ('evaluator_signature', 'evaluator_signature'),
        ('hr_signature', 'hr_signature'),
        ('gm_signature', 'gm_signature'),
    ]:
        img, path = _signature_to_inline_image(tpl, fd.get(data_key))
        if path:
            tmp_paths.append(path)
        ctx[key] = img if img else '(Signed in original)'
    return tmp_paths


def generate_performance_evaluation_docx(form_data, output_path_or_stream, submission_id=None):
    """Generate filled Performance Evaluation DOCX from UI form data."""
    template_path = _get_template_path('performance_evaluation')
    if not template_path:
        raise FileNotFoundError('Performance Evaluation template not found')

    DocxTemplate, _, _ = _get_docxtpl()
    tpl = DocxTemplate(template_path)
    ctx = _build_performance_evaluation_context(form_data)
    tmp_paths = _add_performance_evaluation_signatures(tpl, ctx, form_data)
    if submission_id:
        ctx['submission_id'] = submission_id

    tpl.render(ctx)
    try:
        if hasattr(output_path_or_stream, 'write'):
            tpl.save(output_path_or_stream)
        else:
            tpl.save(output_path_or_stream)
    finally:
        for p in tmp_paths:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except OSError:
                pass


def generate_commencement_docx(form_data, output_path_or_stream, submission_id=None):
    """
    Generate filled Commencement Form DOCX from UI form data.
    Template: HR Documents/Commencement Form - INJAAZ.DOCX
    Placeholders required in template: {{ employee_name }}, {{ position }}, {{ contacts }},
    {{ department }}, {{ organization }}, {{ date_of_joining }}, {{ bank_name }}, {{ bank_branch }},
    {{ account_number }}, {{ employee_signature }}, {{ employee_sign_date }},
    {{ reporting_to_name }}, {{ reporting_to_designation }}, {{ reporting_to_contact }},
    {{ reporting_to_signature }}, {{ reporting_sign_date }}
    """
    template_path = _get_template_path('commencement')
    if not template_path:
        raise FileNotFoundError('Commencement Form template not found in HR Documents folder')

    DocxTemplate, _, _ = _get_docxtpl()
    tpl = DocxTemplate(template_path)
    ctx = _build_commencement_context(form_data)
    tmp_paths = _add_commencement_signatures(tpl, ctx, form_data)
    if submission_id:
        ctx['submission_id'] = submission_id

    tpl.render(ctx)
    try:
        if hasattr(output_path_or_stream, 'write'):
            tpl.save(output_path_or_stream)
        else:
            tpl.save(output_path_or_stream)
    finally:
        for p in tmp_paths:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except OSError:
                pass


# Signature placeholders per form: (template_placeholder_key, form_data_key)
_SIGNATURE_PAIRS = {
    'commencement': [
        ('employee_signature', 'employee_signature'),
        ('reporting_to_signature', 'reporting_to_signature'),
    ],
    'leave_application': [
        ('employee_signature', 'employee_signature'),
        ('replacement_signature', 'replacement_signature'),
        ('gm_signature', 'gm_signature'),
        ('hr_signature', 'hr_signature'),
    ],
    'leave': [
        ('employee_signature', 'employee_signature'),
        ('replacement_signature', 'replacement_signature'),
        ('gm_signature', 'gm_signature'),
        ('hr_signature', 'hr_signature'),
    ],
    'duty_resumption': [
        ('employee_signature', 'employee_signature'),
        ('gm_signature', 'gm_signature'),
        ('hr_signature', 'hr_signature'),
    ],
    'passport_release': [
        ('employee_signature', 'employee_signature'),
        ('gm_signature', 'gm_signature'),
        ('hr_signature', 'hr_signature'),
    ],
    'grievance': [
        ('complainant_signature', 'complainant_signature'),
        ('hr_signature', 'hr_signature'),
        ('gm_signature', 'gm_signature'),
    ],
    'interview_assessment': [
        ('interviewer_signature', 'interviewer_signature'),
        ('hr_signature', 'hr_signature'),
        ('gm_signature', 'gm_signature'),
    ],
    'staff_appraisal': [
        ('employee_signature', 'employee_signature'),
        ('evaluator_signature', 'evaluator_signature'),
        ('hr_signature', 'hr_signature'),
        ('gm_signature', 'gm_signature'),
    ],
    'station_clearance': [
        ('employee_signature', 'employee_signature'),
        ('hr_signature', 'hr_signature'),
    ],
    'visa_renewal': [
        ('employee_signature', 'employee_signature'),
        ('hr_signature', 'hr_signature'),
        ('gm_signature', 'gm_signature'),
    ],
    'contract_renewal': [
        ('evaluator_signature', 'evaluator_signature'),
        ('hr_signature', 'hr_signature'),
        ('gm_signature', 'gm_signature'),
    ],
    'performance_evaluation': [
        ('employee_signature', 'employee_signature'),
        ('evaluator_signature', 'evaluator_signature'),
        ('hr_signature', 'hr_signature'),
        ('gm_signature', 'gm_signature'),
    ],
}


def generate_hr_docx(submission, output_path_or_stream):
    """
    Generate HR DOCX from submission.
    Uses actual Word templates from HR Documents — exact template layout, sections, styling.
    Word and PDF outputs follow the HR Documents templates for consistent document format.
    Returns (True, filled=True) if generated, (False, False) if not supported.
    """
    form_type = (submission.module_type or '').replace('hr_', '')
    form_data = submission.form_data or {}
    submission_id = getattr(submission, 'submission_id', None)

    template_path = _get_template_path(form_type)
    if not template_path:
        return False, False

    try:
        _generate_filled_docx_generic(
            form_type,
            form_data,
            output_path_or_stream,
            submission_id=submission_id,
            signature_pairs=_SIGNATURE_PAIRS.get(form_type),
        )
        return True, True
    except Exception as e:
        logger.exception(
            "Failed to generate DOCX for %s (submission %s): %s",
            form_type, submission_id, e
        )
        raise


def get_supported_docx_forms():
    """Return list of form types that support DOCX download (via HR Documents templates)."""
    return list(_SIGNATURE_PAIRS.keys())
