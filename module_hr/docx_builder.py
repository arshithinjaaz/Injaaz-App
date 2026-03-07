"""
Professional DOCX builder for HR forms.
Generates branded, print-ready documents from HR form submission data
using python-docx — no Word templates required.
Matches PDF template exactly for Commencement Form (B&W, logo right, same layout).
"""
import os
import base64
import tempfile
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ────────────────────────────────────────────────────────────
C_PRI   = "2D6A4F"   # deep green  (header / section bars)
C_SEC   = "52B788"   # mid green   (accents)
C_LIGHT = "D8F3DC"   # mint        (label cells)
C_LROW  = "F0FAF3"   # pale mint   (alternate value rows)
C_WHITE = "FFFFFF"
C_DARK  = "1B2E28"   # near-black
C_MUTED = "5A6E64"   # grey-green
C_LINE  = "B7DFC5"   # border

# ── PDF-matching B&W palette ─────────────────────────────────────────────────
C_BW_BLACK = "000000"
C_BW_GRAY  = "666666"
C_BW_LIGHT = "CCCCCC"

FONT    = "Calibri"

_LOGO = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "static", "logo.png"
)


# ── XML / low-level helpers ───────────────────────────────────────────────────
def _rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_borders(cell, top=None, bottom=None, left=None, right=None,
                  inside_h=None, inside_v=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    sides = {"top": top, "bottom": bottom, "left": left, "right": right,
              "insideH": inside_h, "insideV": inside_v}
    for side, color in sides.items():
        el = OxmlElement(f"w:{side}")
        if color:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color.lstrip("#"))
        else:
            el.set(qn("w:val"), "none")
        borders.append(el)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _cell_text(cell, text, bold=False, italic=False, size=9,
               color=C_DARK, align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=1, space_after=1):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(str(text) if text not in (None, "", "-", "—") else "—")
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = _rgb(color)
    return p


def _set_row_height(row, cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(Cm(cm).emu / 914)))  # EMU → twips
    trHeight.set(qn("w:hRule"), "atLeast")
    existing = trPr.find(qn("w:trHeight"))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


# ── Document setup ────────────────────────────────────────────────────────────
def _new_doc():
    doc = Document()
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return doc


def _gap(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pt)


# ── Page header ───────────────────────────────────────────────────────────────
def _add_header(doc, title, doc_no=None):
    """Full-width header: logo (left) + title & doc info (right) on green bar."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_borders(table)
    table.columns[0].width = Cm(3.8)
    table.columns[1].width = Cm(13.2)

    row = table.rows[0]
    _set_row_height(row, 2.2)

    # Left: logo on green
    lc = row.cells[0]
    _set_bg(lc, C_PRI)
    _set_borders(lc, right=C_SEC)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)
    if os.path.isfile(_LOGO):
        try:
            run = lp.add_run()
            run.add_picture(_LOGO, width=Cm(3.2), height=Cm(1.8))
        except Exception:
            _fallback_logo(lp)
    else:
        _fallback_logo(lp)

    # Right: company name + form title
    rc = row.cells[1]
    _set_bg(rc, C_PRI)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = rc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.4)

    r1 = p.add_run("INJAAZ TECHNICAL SERVICES LLC\n")
    r1.font.name = FONT
    r1.font.size = Pt(7.5)
    r1.font.color.rgb = _rgb(C_SEC)

    r2 = p.add_run(title.upper())
    r2.font.name = FONT
    r2.font.size = Pt(15)
    r2.font.bold = True
    r2.font.color.rgb = _rgb(C_WHITE)

    if doc_no:
        r3 = p.add_run(f"\n{doc_no}")
        r3.font.name = FONT
        r3.font.size = Pt(7.5)
        r3.font.color.rgb = _rgb(C_LIGHT)

    _gap(doc, 5)


def _fallback_logo(p):
    r = p.add_run("INJAAZ")
    r.font.name = FONT
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = _rgb(C_WHITE)


# ── Section header bar ────────────────────────────────────────────────────────
def _section(doc, title):
    table = doc.add_table(rows=1, cols=1)
    _no_table_borders(table)
    table.columns[0].width = Cm(17.0)
    cell = table.rows[0].cells[0]
    _set_bg(cell, C_PRI)
    _set_row_height(table.rows[0], 0.6)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"▌  {title.upper()}")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = _rgb(C_WHITE)
    _gap(doc, 0)


# ── Data grid (label | value pairs) ──────────────────────────────────────────
def _data_table(doc, pairs, cols=2):
    """
    cols=2 → one pair per row  [label | value]
    cols=4 → two pairs per row [label | value | label | value]
    Skip pairs where value is falsy (empty/None/dash).
    """
    pairs = [(lbl, val) for lbl, val in pairs]  # keep all, show — for empty
    if not pairs:
        return

    if cols == 4:
        t = doc.add_table(rows=0, cols=4)
        _no_table_borders(t)
        widths = [Cm(3.8), Cm(4.7), Cm(3.8), Cm(4.7)]
        for i, w in enumerate(widths):
            t.columns[i].width = w
        for i in range(0, len(pairs), 2):
            row = t.add_row()
            _set_row_height(row, 0.62)
            bg = C_LROW if (i // 2) % 2 == 0 else C_WHITE
            lbl1, val1 = pairs[i]
            _set_bg(row.cells[0], C_LIGHT)
            _set_borders(row.cells[0], bottom=C_LINE)
            _cell_text(row.cells[0], lbl1, bold=True, size=7.5, color=C_PRI)
            _set_bg(row.cells[1], bg)
            _set_borders(row.cells[1], bottom=C_LINE, right=C_LINE)
            _cell_text(row.cells[1], val1 or "—", size=8.5, color=C_DARK)
            if i + 1 < len(pairs):
                lbl2, val2 = pairs[i + 1]
                _set_bg(row.cells[2], C_LIGHT)
                _set_borders(row.cells[2], bottom=C_LINE)
                _cell_text(row.cells[2], lbl2, bold=True, size=7.5, color=C_PRI)
                _set_bg(row.cells[3], bg)
                _set_borders(row.cells[3], bottom=C_LINE)
                _cell_text(row.cells[3], val2 or "—", size=8.5, color=C_DARK)
            else:
                for ci in (2, 3):
                    _set_bg(row.cells[ci], bg)
                    row.cells[ci].paragraphs[0].clear()
    else:
        t = doc.add_table(rows=0, cols=2)
        _no_table_borders(t)
        t.columns[0].width = Cm(5.2)
        t.columns[1].width = Cm(11.8)
        for i, (lbl, val) in enumerate(pairs):
            row = t.add_row()
            _set_row_height(row, 0.62)
            bg = C_LROW if i % 2 == 0 else C_WHITE
            _set_bg(row.cells[0], C_LIGHT)
            _set_borders(row.cells[0], bottom=C_LINE)
            _cell_text(row.cells[0], lbl, bold=True, size=7.5, color=C_PRI)
            _set_bg(row.cells[1], bg)
            _set_borders(row.cells[1], bottom=C_LINE)
            _cell_text(row.cells[1], val or "—", size=8.5, color=C_DARK)
    _gap(doc, 5)


# ── Long text field (full width) ──────────────────────────────────────────────
def _long_field(doc, label, value):
    t = doc.add_table(rows=2, cols=1)
    _no_table_borders(t)
    t.columns[0].width = Cm(17.0)
    lc = t.rows[0].cells[0]
    _set_bg(lc, C_LIGHT)
    _set_borders(lc, bottom=C_LINE)
    _cell_text(lc, label, bold=True, size=7.5, color=C_PRI)
    vc = t.rows[1].cells[0]
    _set_bg(vc, C_LROW)
    _set_borders(vc, bottom=C_LINE)
    p = vc.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.2)
    run = p.add_run(str(value) if value not in (None, "", "-") else "—")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = _rgb(C_DARK)
    _gap(doc, 5)


# ── Rating / score table ──────────────────────────────────────────────────────
def _rating_table(doc, rows_data):
    """
    rows_data = list of (criterion_label, score_value, max_val, weight_label)
    Renders a score table with a coloured score badge.
    """
    if not rows_data:
        return
    t = doc.add_table(rows=0, cols=3)
    _no_table_borders(t)
    t.columns[0].width = Cm(9.5)
    t.columns[1].width = Cm(3.5)
    t.columns[2].width = Cm(4.0)

    # Header row
    hrow = t.add_row()
    _set_row_height(hrow, 0.55)
    for ci, (txt, c) in enumerate([("Criterion", C_PRI), ("Score", C_PRI), ("Max / Weight", C_PRI)]):
        _set_bg(hrow.cells[ci], C_PRI)
        _cell_text(hrow.cells[ci], txt, bold=True, size=7.5, color=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (criterion, score, max_val, weight) in enumerate(rows_data):
        row = t.add_row()
        _set_row_height(row, 0.62)
        bg = C_LROW if i % 2 == 0 else C_WHITE
        _set_bg(row.cells[0], bg)
        _set_borders(row.cells[0], bottom=C_LINE)
        _cell_text(row.cells[0], criterion, size=8.5, color=C_DARK)

        # Score cell: highlighted if filled
        score_str = str(score) if score not in (None, "", "-", "—") else "—"
        _set_bg(row.cells[1], C_SEC if score_str != "—" else bg)
        _set_borders(row.cells[1], bottom=C_LINE)
        _cell_text(row.cells[1], score_str, bold=(score_str != "—"),
                   size=9, color=C_WHITE if score_str != "—" else C_MUTED,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        _set_bg(row.cells[2], bg)
        _set_borders(row.cells[2], bottom=C_LINE)
        _cell_text(row.cells[2], str(weight) if weight else str(max_val),
                   size=8, color=C_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    _gap(doc, 5)


# ── Checkbox / clearance table ────────────────────────────────────────────────
def _checklist_table(doc, items):
    """
    items = list of (label, status_value, date_value)
    Renders a checklist with ✔ / ✗ indicators.
    """
    if not items:
        return
    t = doc.add_table(rows=0, cols=3)
    _no_table_borders(t)
    t.columns[0].width = Cm(10.5)
    t.columns[1].width = Cm(3.0)
    t.columns[2].width = Cm(3.5)

    hrow = t.add_row()
    _set_row_height(hrow, 0.55)
    for ci, txt in enumerate(["Item", "Status", "Date"]):
        _set_bg(hrow.cells[ci], C_PRI)
        _cell_text(hrow.cells[ci], txt, bold=True, size=7.5, color=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (label, status, date_val) in enumerate(items):
        row = t.add_row()
        _set_row_height(row, 0.62)
        bg = C_LROW if i % 2 == 0 else C_WHITE
        _set_bg(row.cells[0], bg)
        _set_borders(row.cells[0], bottom=C_LINE)
        _cell_text(row.cells[0], label, size=8.5, color=C_DARK)

        done = str(status).lower() in ("on", "completed", "yes", "true", "1")
        _set_bg(row.cells[1], "D8F3DC" if done else "FEE2E2")
        _set_borders(row.cells[1], bottom=C_LINE)
        _cell_text(row.cells[1], "✔  Done" if done else "—",
                   bold=done, size=8.5,
                   color="2D6A4F" if done else "B91C1C",
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        _set_bg(row.cells[2], bg)
        _set_borders(row.cells[2], bottom=C_LINE)
        _cell_text(row.cells[2], _fmt(date_val) if date_val else "—",
                   size=8, color=C_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    _gap(doc, 5)


# ── Signature block ───────────────────────────────────────────────────────────
def _sig_block(doc, signatures):
    """
    signatures = list of (label, base64_data_url_or_None, date_value)
    """
    if not signatures:
        return
    n = len(signatures)
    col_w = Cm(17.0 / n)

    t = doc.add_table(rows=3, cols=n)
    _no_table_borders(t)
    for i in range(n):
        t.columns[i].width = col_w

    # Row 0 – label
    for i, (label, _, _) in enumerate(signatures):
        c = t.rows[0].cells[i]
        _set_bg(c, C_PRI)
        _set_borders(c, right=C_SEC if i < n - 1 else None)
        _cell_text(c, label, bold=True, size=8, color=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=3)

    # Row 1 – signature image or blank box
    _set_row_height(t.rows[1], 1.6)
    for i, (_, img_data, _) in enumerate(signatures):
        c = t.rows[1].cells[i]
        _set_bg(c, C_WHITE)
        _set_borders(c, bottom=C_LINE, left=C_LINE, right=C_LINE)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if img_data and isinstance(img_data, str) and img_data.startswith("data:image"):
            try:
                _, b64 = img_data.split(",", 1)
                img_bytes = base64.b64decode(b64)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
                    f.write(img_bytes)
                    tmp = f.name
                run = p.add_run()
                run.add_picture(tmp, width=Mm(40), height=Mm(14))
                try:
                    os.unlink(tmp)
                except OSError:
                    pass
            except Exception:
                r = p.add_run("(Signed)")
                r.font.name = FONT
                r.font.size = Pt(8)
                r.font.color.rgb = _rgb(C_MUTED)
        else:
            r = p.add_run(" ")
            r.font.size = Pt(30)

    # Row 2 – date
    for i, (_, _, date_val) in enumerate(signatures):
        c = t.rows[2].cells[i]
        _set_bg(c, C_LROW)
        _set_borders(c, top=C_LINE)
        _cell_text(c, f"Date: {_fmt(date_val)}" if date_val else "Date: ___________",
                   size=8, color=C_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    _gap(doc, 5)


# ── Footer ────────────────────────────────────────────────────────────────────
def _footer(doc, doc_no=None, form_date=None):
    # Thin green bar
    t = doc.add_table(rows=1, cols=1)
    _no_table_borders(t)
    t.columns[0].width = Cm(17.0)
    c = t.rows[0].cells[0]
    _set_bg(c, C_PRI)
    _set_row_height(t.rows[0], 0.15)
    c.paragraphs[0].clear()

    _gap(doc, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    parts = ["INJAAZ TECHNICAL SERVICES LLC"]
    if doc_no:
        parts.append(f"Doc: {doc_no}")
    if form_date:
        parts.append(f"Date: {_fmt(form_date)}")
    r = p.add_run("   |   ".join(parts))
    r.font.name = FONT
    r.font.size = Pt(7)
    r.font.color.rgb = _rgb(C_MUTED)


# ── Date formatter ────────────────────────────────────────────────────────────
def _fmt(v):
    if not v:
        return "—"
    try:
        d = datetime.fromisoformat(str(v).replace("Z", "+00:00"))
        return d.strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        s = str(v)
        return s if s not in ("", "-", "—") else "—"


def _fd(form_data, key, default="—"):
    v = (form_data or {}).get(key)
    if v in (None, "", "-"):
        return default
    return str(v)


# ── PDF-matching layout (B&W, same structure as hr_pdf_builder) ───────────────
def _sig_to_transparent_png(data_url):
    """Convert base64 signature to PNG with transparent background. Returns temp path or None."""
    if not data_url or not isinstance(data_url, str) or not data_url.startswith("data:image"):
        return None
    try:
        _, b64 = data_url.split(",", 1)
        raw = base64.b64decode(b64)
        try:
            from PIL import Image as PILImage
            pil = PILImage.open(BytesIO(raw)).convert("RGBA")
            data = pil.load()
            w, h = pil.size
            for y in range(h):
                for x in range(w):
                    r, g, b, a = data[x, y]
                    if r >= 250 and g >= 250 and b >= 250:
                        data[x, y] = (r, g, b, 0)
            f = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            pil.save(f, "PNG")
            f.close()
            return f.name
        except Exception:
            f = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            f.write(raw)
            f.close()
            return f.name
    except Exception:
        return None


def _add_header_pdf_style(doc, form_name):
    """Logo right, headline left: Injaaz Facility Management's + form name. B&W."""
    t = doc.add_table(rows=1, cols=2)
    _no_table_borders(t)
    t.columns[0].width = Cm(12.5)
    t.columns[1].width = Cm(4.5)
    row = t.rows[0]
    # Left: Injaaz Facility Management's + form name
    lc = row.cells[0]
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = lc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("Injaaz Facility Management's\n")
    r1.font.name = FONT
    r1.font.size = Pt(7)
    r1.font.color.rgb = _rgb(C_BW_GRAY)
    r2 = p.add_run(form_name)
    r2.font.name = FONT
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = _rgb(C_BW_BLACK)
    # Right: logo
    rc = row.cells[1]
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(0)
    if os.path.isfile(_LOGO):
        try:
            run = rp.add_run()
            run.add_picture(_LOGO, width=Cm(1.7), height=Cm(1.7))
        except Exception:
            _fallback_logo_bw(rp)
    else:
        _fallback_logo_bw(rp)
    # Bottom rule
    _add_rule(doc, C_BW_BLACK, 1)
    _gap(doc, 4)


def _fallback_logo_bw(p):
    r = p.add_run("INJAAZ")
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = _rgb(C_BW_BLACK)


def _add_rule(doc, color_hex, pt=0.5):
    t = doc.add_table(rows=1, cols=1)
    _no_table_borders(t)
    t.columns[0].width = Cm(17.0)
    c = t.rows[0].cells[0]
    _set_borders(c, bottom=color_hex)
    c.paragraphs[0].clear()
    _set_row_height(t.rows[0], 0.12)


def _section_bar_numbered(doc, num, title):
    """01. Title with underline. B&W."""
    t = doc.add_table(rows=1, cols=1)
    _no_table_borders(t)
    t.columns[0].width = Cm(17.0)
    c = t.rows[0].cells[0]
    p = c.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{num}. {title}")
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = _rgb(C_BW_BLACK)
    _set_borders(c, bottom=C_BW_BLACK)
    _gap(doc, 2)


def _instruction_line(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(6)
        run.font.color.rgb = _rgb(C_BW_GRAY)
    _gap(doc, 4)


def _data_table_pdf_style(doc, pairs, cols=2):
    """B&W: label | value, horizontal lines only."""
    if not pairs:
        return
    if cols == 4:
        rows_data = []
        for i in range(0, len(pairs), 2):
            r = [pairs[i]]
            if i + 1 < len(pairs):
                r.append(pairs[i + 1])
            rows_data.append(r)
        t = doc.add_table(rows=0, cols=4)
        widths = [Cm(3.7), Cm(4.6), Cm(3.7), Cm(4.6)]
    else:
        rows_data = [[p] for p in pairs]
        t = doc.add_table(rows=0, cols=2)
        widths = [Cm(4.8), Cm(12.2)]
    _no_table_borders(t)
    for i, w in enumerate(widths):
        t.columns[i].width = w
    for row_pairs in rows_data:
        row = t.add_row()
        _set_row_height(row, 0.55)
        idx = 0
        for lbl, val in row_pairs:
            _cell_text(row.cells[idx], lbl.upper(), bold=True, size=7, color=C_BW_GRAY)
            _set_borders(row.cells[idx], bottom=C_BW_LIGHT)
            _cell_text(row.cells[idx + 1], val or "—", size=8, color=C_BW_BLACK)
            _set_borders(row.cells[idx + 1], bottom=C_BW_LIGHT)
            idx += 2
        while idx < len(widths):
            row.cells[idx].paragraphs[0].clear()
            idx += 1
    _gap(doc, 5)


def _sig_block_pdf_style(doc, signatures):
    """Transparent signature, no box. Left-aligned. B&W."""
    if not signatures:
        return
    n = len(signatures)
    col_w = Cm(min(4.0, 17.0 / n))
    t = doc.add_table(rows=3, cols=n)
    _no_table_borders(t)
    for i in range(n):
        t.columns[i].width = col_w
    # Label row
    for i, (label, _, _) in enumerate(signatures):
        c = t.rows[0].cells[i]
        _cell_text(c, label, bold=True, size=7, color=C_BW_BLACK, space_before=0, space_after=2)
    # Signature row - no borders, no background, decent space
    _set_row_height(t.rows[1], 1.2)
    tmp_paths = []
    for i, (_, img_data, _) in enumerate(signatures):
        c = t.rows[1].cells[i]
        p = c.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        img_data = (img_data or "").strip() if isinstance(img_data, str) else img_data
        if img_data and isinstance(img_data, str) and img_data.startswith("data:image"):
            path = _sig_to_transparent_png(img_data)
            if path:
                tmp_paths.append(path)
                try:
                    run = p.add_run()
                    run.add_picture(path, width=Mm(45), height=Mm(18))
                except Exception:
                    p.add_run("Sign").font.size = Pt(7)
            else:
                p.add_run("Sign").font.size = Pt(7)
        else:
            r = p.add_run("Sign")
            r.font.name = FONT
            r.font.size = Pt(7)
            r.font.color.rgb = _rgb(C_BW_GRAY)
    for path in tmp_paths:
        try:
            os.unlink(path)
        except OSError:
            pass
    # Date row
    for i, (_, _, date_val) in enumerate(signatures):
        c = t.rows[2].cells[i]
        _cell_text(c, _fmt(date_val) if date_val else "Date", size=6, color=C_BW_GRAY, space_before=0, space_after=0)
    _gap(doc, 6)


def _info_box_pdf_style(doc, text):
    """Note box with light border. B&W."""
    t = doc.add_table(rows=1, cols=1)
    _no_table_borders(t)
    t.columns[0].width = Cm(17.0)
    c = t.rows[0].cells[0]
    _set_borders(c, top=C_BW_LIGHT, bottom=C_BW_LIGHT, left=C_BW_LIGHT, right=C_BW_LIGHT)
    p = c.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    r1 = p.add_run("Note ")
    r1.font.name = FONT
    r1.font.size = Pt(7)
    r1.font.bold = True
    r1.font.color.rgb = _rgb(C_BW_BLACK)
    r2 = p.add_run(text)
    r2.font.name = FONT
    r2.font.size = Pt(7)
    r2.font.color.rgb = _rgb(C_BW_BLACK)
    _gap(doc, 5)


def _footer_pdf_style(doc):
    """Generated timestamp. B&W."""
    _gap(doc, 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Generated {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    r.font.name = FONT
    r.font.size = Pt(6)
    r.font.color.rgb = _rgb(C_BW_GRAY)


# ═════════════════════════════════════════════════════════════════════════════
# Form-specific builders
# ═════════════════════════════════════════════════════════════════════════════

def _build_leave_application(fd):
    doc = _new_doc()
    _add_header(doc, "Leave Application Form")
    _section(doc, "Employee Information")
    _data_table(doc, [
        ("Employee Name",   _fd(fd, "employee_name")),
        ("Employee ID",     _fd(fd, "employee_id")),
        ("Job Title",       _fd(fd, "job_title")),
        ("Department",      _fd(fd, "department")),
        ("Date of Joining", _fmt(fd.get("date_of_joining"))),
        ("Mobile No.",      _fd(fd, "mobile_no")),
        ("Last Leave Date", _fmt(fd.get("last_leave_date"))),
        ("Today's Date",    _fmt(fd.get("today_date"))),
    ], cols=4)

    _section(doc, "Leave Details")
    leave_type = _fd(fd, "leave_type")
    if leave_type == "other":
        leave_type = _fd(fd, "leave_type_other", "Other")
    _data_table(doc, [
        ("Leave Type",            leave_type),
        ("Salary Advance",        _fd(fd, "salary_advance", "No").upper()),
        ("First Day of Leave",    _fmt(fd.get("first_day_of_leave"))),
        ("Last Day of Leave",     _fmt(fd.get("last_day_of_leave"))),
        ("Total Days",            _fd(fd, "total_days_requested")),
        ("Date Returning",        _fmt(fd.get("date_returning_to_work"))),
        ("Reachable Telephone",   _fd(fd, "telephone_reachable")),
        ("Replacement Name",      _fd(fd, "replacement_name")),
    ], cols=4)

    _section(doc, "Signatures")
    _sig_block(doc, [
        ("Employee",      fd.get("employee_signature"),    fd.get("today_date")),
        ("Replacement",   fd.get("replacement_signature"), None),
        ("GM Approval",   fd.get("gm_signature"),          None),
    ])

    _section(doc, "HR Use Only")
    _data_table(doc, [
        ("HR Checked",      _fd(fd, "hr_checked")),
        ("Balance C/F",     _fd(fd, "hr_balance_cf")),
        ("Contract Year",   _fd(fd, "hr_contract_year")),
        ("Paid",            _fd(fd, "hr_paid")),
        ("Unpaid",          _fd(fd, "hr_unpaid")),
    ], cols=4)
    if fd.get("hr_comments"):
        _long_field(doc, "HR Comments", fd.get("hr_comments"))
    _sig_block(doc, [("HR Signature", fd.get("hr_signature"), fd.get("hr_date"))])
    _footer(doc)
    return doc


def _build_commencement(fd):
    """Commencement Form — exact same layout as PDF (B&W, logo right, same sections)."""
    doc = _new_doc()
    for section in doc.sections:
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)
        section.top_margin = Cm(0.7)
        section.bottom_margin = Cm(1.4)
    _add_header_pdf_style(doc, "Commencement Form")
    _instruction_line(doc,
        "To complete the administrative aspect of your Employment please complete this form within 5 days of joining "
        "and fax it back to AH or email it to joana@ajmanholding.ae")
    _section_bar_numbered(doc, "01", "Personal Details")
    _data_table_pdf_style(doc, [
        ("Full Name", _fd(fd, "employee_name")),
        ("Position / Title", _fd(fd, "position")),
        ("Department", _fd(fd, "department")),
        ("Organization", _fd(fd, "organization", "INJAAZ")),
        ("Contact Number", _fd(fd, "contacts")),
        ("Date of Joining (DD/MM/YYYY)", _fmt(fd.get("date_of_joining"))),
    ], cols=2)
    _section_bar_numbered(doc, "02", "Bank Account Details")
    _data_table_pdf_style(doc, [
        ("Bank Name", _fd(fd, "bank_name")),
        ("Branch", _fd(fd, "bank_branch")),
        ("Account Number", _fd(fd, "account_number")),
    ], cols=2)
    _section_bar_numbered(doc, "03", "Employee Declaration")
    _sig_block_pdf_style(doc, [
        ("Employee Signature", fd.get("employee_signature"), fd.get("employee_sign_date")),
    ])
    _section_bar_numbered(doc, "04", "Reporting Manager")
    _data_table_pdf_style(doc, [
        ("Manager Name", _fd(fd, "reporting_to_name")),
        ("Designation / Title", _fd(fd, "reporting_to_designation")),
        ("Contact Number", _fd(fd, "reporting_to_contact")),
    ], cols=2)
    _sig_block_pdf_style(doc, [
        ("Reporting Officer Signature", fd.get("reporting_to_signature"), fd.get("reporting_sign_date")),
    ])
    _info_box_pdf_style(doc,
        "Need a Salary Letter? If you require assistance with a salary letter to open a new bank account, "
        "please forward your request to the HR Department.")
    _footer_pdf_style(doc)
    return doc


def _build_duty_resumption(fd):
    doc = _new_doc()
    _add_header(doc, "Duty Resumption Form")
    _section(doc, "Employee Details")
    _data_table(doc, [
        ("Requester",       _fd(fd, "requester")),
        ("Employee Name",   _fd(fd, "employee_name")),
        ("Employee ID",     _fd(fd, "employee_id")),
        ("Job Title",       _fd(fd, "job_title")),
        ("Company",         _fd(fd, "company", "INJAAZ LLC")),
    ], cols=4)

    _section(doc, "Leave & Resumption Dates")
    _data_table(doc, [
        ("Leave Started",            _fmt(fd.get("leave_started"))),
        ("Leave Ended",              _fmt(fd.get("leave_ended"))),
        ("Planned Resumption Date",  _fmt(fd.get("planned_resumption_date"))),
        ("Actual Resumption Date",   _fmt(fd.get("actual_resumption_date"))),
        ("Note",                     _fd(fd, "note")),
    ], cols=4)

    if fd.get("line_manager_remarks"):
        _section(doc, "Line Manager Remarks")
        _long_field(doc, "Remarks", fd.get("line_manager_remarks"))

    _section(doc, "Signatures")
    sigs = [("Employee", fd.get("employee_signature"), fd.get("sign_date"))]
    if fd.get("gm_signature"):
        sigs.append(("GM Approval", fd.get("gm_signature"), None))
    if fd.get("hr_signature"):
        sigs.append(("HR", fd.get("hr_signature"), None))
    _sig_block(doc, sigs)
    _footer(doc)
    return doc


def _build_passport_release(fd):
    doc = _new_doc()
    form_type_val = _fd(fd, "passport_form_type", "release").replace("_", " ").title()
    _add_header(doc, f"Passport {form_type_val} Form")
    _section(doc, "Request Details")
    _data_table(doc, [
        ("Form Type",    form_type_val),
        ("Date",         _fmt(fd.get("form_date"))),
        ("Requester",    _fd(fd, "requester")),
        ("Employee Name",_fd(fd, "employee_name")),
        ("Employee ID",  _fd(fd, "employee_id")),
        ("Job Title",    _fd(fd, "job_title")),
        ("Project",      _fd(fd, "project")),
        ("Release Date", _fmt(fd.get("release_date"))),
    ], cols=4)

    if fd.get("purpose_of_release"):
        _section(doc, "Purpose of Release")
        _long_field(doc, "Purpose", fd.get("purpose_of_release"))

    _section(doc, "Signatures")
    sigs = [("Employee", fd.get("employee_signature"), fd.get("form_date"))]
    if fd.get("gm_signature"):
        sigs.append(("GM Approval", fd.get("gm_signature"), None))
    if fd.get("hr_signature"):
        sigs.append(("HR", fd.get("hr_signature"), None))
    _sig_block(doc, sigs)
    _footer(doc)
    return doc


def _build_grievance(fd):
    doc = _new_doc()
    _add_header(doc, "Employee Grievance / Disciplinary Action Form")
    _section(doc, "First Party (Complainant)")
    _data_table(doc, [
        ("Employee Name",      _fd(fd, "complainant_name")),
        ("Employee ID",        _fd(fd, "complainant_id")),
        ("Designation",        _fd(fd, "complainant_designation")),
        ("Contact No.",        _fd(fd, "complainant_contact")),
        ("Date of Incident",   _fmt(fd.get("date_of_incident"))),
        ("Shift / Time",       _fd(fd, "shift_time")),
        ("Location",           _fd(fd, "issue_location", "").replace("_", " ").title()),
    ], cols=4)

    _section(doc, "Second Party")
    _data_table(doc, [
        ("Employee Name",   _fd(fd, "second_party_name")),
        ("Staff ID",        _fd(fd, "second_party_id")),
        ("Department",      _fd(fd, "second_party_department")),
        ("Place of Incident",_fd(fd, "place_of_incident")),
        ("Shift / Time",    _fd(fd, "second_party_shift")),
        ("Contact No.",     _fd(fd, "second_party_contact")),
    ], cols=4)

    _section(doc, "Complaint Details")
    _long_field(doc, "Description of Complaint", fd.get("complaint_description"))
    _data_table(doc, [
        ("Witnesses",     _fd(fd, "witnesses")),
        ("Who Informed",  _fd(fd, "who_informed")),
        ("Attachment",    _fd(fd, "attachment")),
    ], cols=2)

    _section(doc, "HR Review")
    _data_table(doc, [
        ("Statement 2nd Party", _fd(fd, "statement_2nd_party")),
        ("Statement Verified",  _fd(fd, "hr_statement_verified")),
        ("1st / Recurring",     _fd(fd, "hr_first_recurring")),
    ], cols=4)
    if fd.get("hr_remarks"):
        _long_field(doc, "HR Remarks", fd.get("hr_remarks"))
    if fd.get("gm_remarks"):
        _long_field(doc, "GM Remarks", fd.get("gm_remarks"))

    _section(doc, "Signatures")
    sigs = [("Complainant", fd.get("complainant_signature"), fd.get("date_of_incident"))]
    if fd.get("hr_signature"):
        sigs.append(("HR", fd.get("hr_signature"), None))
    if fd.get("gm_signature"):
        sigs.append(("GM", fd.get("gm_signature"), None))
    _sig_block(doc, sigs)
    _footer(doc)
    return doc


def _build_visa_renewal(fd):
    doc = _new_doc()
    _add_header(doc, "Visa Renewal Form")
    _section(doc, "Employee Information")
    _data_table(doc, [
        ("Date",             _fmt(fd.get("form_date"))),
        ("Employee Name",    _fd(fd, "employee_name")),
        ("Employee ID",      _fd(fd, "employee_id")),
        ("Employer",         _fd(fd, "employer", "INJAAZ")),
        ("Position",         _fd(fd, "position")),
        ("Years Completed",  _fd(fd, "years_completed")),
        ("Decision",         _fd(fd, "decision", "").replace("_", " ").title()),
    ], cols=4)

    _section(doc, "Signature")
    _sig_block(doc, [
        ("Employee", fd.get("employee_signature"), fd.get("form_date")),
    ])
    _footer(doc)
    return doc


def _build_interview_assessment(fd):
    doc = _new_doc()
    _add_header(doc, "Interview Assessment Form")
    _section(doc, "Candidate Information")
    _data_table(doc, [
        ("Candidate Name",         _fd(fd, "candidate_name")),
        ("Position Title",         _fd(fd, "position_title")),
        ("Academic Qualification", _fd(fd, "academic_qualification")),
        ("Age",                    _fd(fd, "age")),
        ("Gender",                 _fd(fd, "gender")),
        ("Marital Status",         _fd(fd, "marital_status")),
        ("No. of Dependents",      _fd(fd, "dependents")),
        ("Nationality",            _fd(fd, "nationality")),
        ("Current Job Title",      _fd(fd, "current_job_title")),
        ("Years of Experience",    _fd(fd, "years_experience")),
        ("Current Salary",         _fd(fd, "current_salary")),
        ("Expected Salary",        _fd(fd, "expected_salary")),
        ("Interview Date",         _fmt(fd.get("interview_date"))),
        ("Interview By",           _fd(fd, "interview_by")),
    ], cols=4)

    _section(doc, "Assessment Ratings")
    _rating_table(doc, [
        ("Turn-out & Appearance",         _fd(fd, "rating_turnout"),          "Outstanding→Low", ""),
        ("Confidence",                    _fd(fd, "rating_confidence"),        "Outstanding→Low", ""),
        ("Mental Alertness",              _fd(fd, "rating_mental_alertness"),  "Outstanding→Low", ""),
        ("Maturity & Emotional Stability",_fd(fd, "rating_maturity"),          "Outstanding→Low", ""),
        ("Communication Skills",          _fd(fd, "rating_communication"),     "Outstanding→Low", ""),
        ("Technical Knowledge",           _fd(fd, "rating_technical"),         "Outstanding→Low", ""),
        ("Relevant Training",             _fd(fd, "rating_training"),          "Outstanding→Low", ""),
        ("Relevant Experience",           _fd(fd, "rating_experience"),        "Outstanding→Low", ""),
        ("Overall Rating",                _fd(fd, "rating_overall"),           "Outstanding→Low", ""),
    ])

    if fd.get("overall_assessment"):
        _section(doc, "Overall Assessment")
        _long_field(doc, "Assessment", fd.get("overall_assessment"))

    _data_table(doc, [("Eligible for Employment", _fd(fd, "eligibility", "").upper())], cols=2)

    _section(doc, "Signatures")
    sigs = []
    if fd.get("hr_signature"):
        sigs.append(("HR", fd.get("hr_signature"), None))
    if fd.get("gm_signature"):
        sigs.append(("GM", fd.get("gm_signature"), None))
    if sigs:
        _sig_block(doc, sigs)
    _footer(doc)
    return doc


def _build_staff_appraisal(fd):
    doc = _new_doc()
    _add_header(doc, "Staff Appraisal Form")
    _section(doc, "Employee Information")
    _data_table(doc, [
        ("Employee Name",   _fd(fd, "employee_name")),
        ("Employee ID",     _fd(fd, "employee_id")),
        ("Department",      _fd(fd, "department")),
        ("Position",        _fd(fd, "position")),
        ("Appraisal Period",_fd(fd, "appraisal_period")),
        ("Reviewer",        _fd(fd, "reviewer")),
    ], cols=4)

    _section(doc, "Performance Ratings  (Scale 1 – 5)")
    _rating_table(doc, [
        ("Punctuality",    _fd(fd, "rating_punctuality"),     "1-5", "15%"),
        ("Job Knowledge",  _fd(fd, "rating_job_knowledge"),   "1-5", "15%"),
        ("Quality of Work",_fd(fd, "rating_quality"),         "1-5", "15%"),
        ("Productivity",   _fd(fd, "rating_productivity"),    "1-5", "15%"),
        ("Communication",  _fd(fd, "rating_communication"),   "1-5", "10%"),
        ("Teamwork",       _fd(fd, "rating_teamwork"),        "1-5", "10%"),
        ("Problem-Solving",_fd(fd, "rating_problem_solving"), "1-5", "10%"),
        ("Adaptability",   _fd(fd, "rating_adaptability"),    "1-5",  "5%"),
        ("Leadership",     _fd(fd, "rating_leadership"),      "1-5",  "5%"),
        ("TOTAL SCORE",    _fd(fd, "total_score"),            "5",   "100%"),
    ])

    # Comments
    comment_fields = [
        ("comments_punctuality",    "Punctuality"),
        ("comments_job_knowledge",  "Job Knowledge"),
        ("comments_quality",        "Quality"),
        ("comments_productivity",   "Productivity"),
        ("comments_communication",  "Communication"),
        ("comments_teamwork",       "Teamwork"),
        ("comments_problem_solving","Problem-Solving"),
        ("comments_adaptability",   "Adaptability"),
        ("comments_leadership",     "Leadership"),
    ]
    comment_pairs = [(lbl, _fd(fd, key)) for key, lbl in comment_fields
                     if fd.get(key) and fd[key] not in ("", "-")]
    if comment_pairs:
        _section(doc, "Evaluator Comments")
        _data_table(doc, comment_pairs, cols=2)

    if fd.get("employee_strengths"):
        _section(doc, "Employee Strengths")
        _long_field(doc, "Strengths & Achievements", fd.get("employee_strengths"))

    _section(doc, "Signatures")
    sigs = [("Employee", fd.get("employee_signature"), None)]
    if fd.get("hr_signature"):
        sigs.append(("HR", fd.get("hr_signature"), None))
    if fd.get("gm_signature"):
        sigs.append(("GM", fd.get("gm_signature"), None))
    _sig_block(doc, sigs)
    _footer(doc)
    return doc


def _build_station_clearance(fd):
    doc = _new_doc()
    _add_header(doc, "Station Clearance Form")
    _section(doc, "Employee Information")
    _data_table(doc, [
        ("Employee Name",     _fd(fd, "employee_name")),
        ("Employee ID",       _fd(fd, "employee_id")),
        ("Position",          _fd(fd, "position")),
        ("Department",        _fd(fd, "department")),
        ("Section",           _fd(fd, "section")),
        ("Employment Date",   _fmt(fd.get("employment_date"))),
        ("Type of Departure", _fd(fd, "type_of_departure", "").replace("_", " ").title()),
        ("Last Working Date", _fmt(fd.get("last_working_date"))),
    ], cols=4)

    _section(doc, "Department Clearance")
    dept_items = [
        ("Has completed / handed over all tasks on hand",     fd.get("tasks_handed_over"),     fd.get("dept_date_1")),
        ("Has handed over all original working documents",    fd.get("documents_handed_over"), fd.get("dept_date_2")),
        ("Has handed over all normal & electronic files",     fd.get("files_handed_over"),     fd.get("dept_date_3")),
        ("Keys Returned",                                     fd.get("keys_returned"),         fd.get("dept_date_4")),
        ("Toolbox Returned",                                  fd.get("toolbox_returned"),      fd.get("dept_date_5")),
        ("Access Card",                                       fd.get("access_card_returned"),  fd.get("dept_date_6")),
    ]
    _checklist_table(doc, dept_items)
    if fd.get("dept_others"):
        _long_field(doc, "Department – Others", fd.get("dept_others"))

    _section(doc, "IT Clearance")
    it_items = [
        ("E-mail Account Cancelled",                   fd.get("email_cancelled"),         None),
        ("Has returned all software / hardware",       fd.get("software_hardware_returned"),None),
        ("Laptop Returned",                            fd.get("laptop_returned"),           None),
    ]
    _checklist_table(doc, it_items)
    if fd.get("it_others"):
        _long_field(doc, "IT – Others", fd.get("it_others"))

    _section(doc, "HR Clearance")
    hr_items = [
        ("Employee file shifted to Exit folder",       fd.get("file_shifted"),  None),
        ("Payment of outstanding dues (Salary)",       fd.get("dues_paid"),     None),
        ("Medical Card Returned",                      fd.get("medical_card_returned"), None),
    ]
    _checklist_table(doc, hr_items)
    if fd.get("hr_others"):
        _long_field(doc, "HR – Others", fd.get("hr_others"))

    _section(doc, "Finance Clearance")
    fin_items = [("EOS Benefits Transfer", fd.get("eos_transfer"), None)]
    _checklist_table(doc, fin_items)
    if fd.get("finance_others"):
        _long_field(doc, "Finance – Others", fd.get("finance_others"))

    if fd.get("remarks"):
        _section(doc, "Remarks")
        _long_field(doc, "Remarks", fd.get("remarks"))

    _section(doc, "Signatures")
    sigs = [("Employee", fd.get("employee_signature"), None)]
    if fd.get("hr_signature"):
        sigs.append(("HR", fd.get("hr_signature"), None))
    _sig_block(doc, sigs)
    _footer(doc)
    return doc


def _build_performance_evaluation(fd):
    doc = _new_doc()
    _add_header(doc, "Performance Evaluation Form")
    _section(doc, "Employee Information")
    _data_table(doc, [
        ("Employee Name",    _fd(fd, "employee_name")),
        ("Employee ID",      _fd(fd, "employee_id")),
        ("Department",       _fd(fd, "department")),
        ("Designation",      _fd(fd, "designation")),
        ("Date of Evaluation",_fmt(fd.get("date_of_evaluation"))),
        ("Date of Joining",  _fmt(fd.get("date_of_joining"))),
        ("Evaluation By",    _fd(fd, "evaluation_done_by")),
    ], cols=4)

    _section(doc, "Performance Scores  (Scale 1 – 10)")
    _rating_table(doc, [
        ("Score 01", _fd(fd, "score_01"), "10", ""),
        ("Score 02", _fd(fd, "score_02"), "10", ""),
        ("Score 03", _fd(fd, "score_03"), "10", ""),
        ("Score 04", _fd(fd, "score_04"), "10", ""),
        ("Score 05", _fd(fd, "score_05"), "10", ""),
        ("Score 06", _fd(fd, "score_06"), "10", ""),
        ("Score 07", _fd(fd, "score_07"), "10", ""),
        ("Score 08", _fd(fd, "score_08"), "10", ""),
        ("Score 09", _fd(fd, "score_09"), "10", ""),
        ("Score 10", _fd(fd, "score_10"), "10", ""),
        ("OVERALL SCORE", _fd(fd, "overall_score"), "100", ""),
    ])

    _section(doc, "Evaluator Details")
    _data_table(doc, [
        ("Evaluator Name",        _fd(fd, "evaluator_name")),
        ("Evaluator Designation", _fd(fd, "evaluator_designation")),
    ], cols=4)
    for lbl, key in [
        ("Evaluator Observation",  "evaluator_observation"),
        ("Area of Concern",        "area_of_concern"),
        ("Training Required",      "training_required"),
        ("Employee Comments",      "employee_comments"),
        ("Concern In-charge",      "concern_incharge_name"),
        ("In-charge Comments",     "incharge_comments"),
        ("GM Remarks",             "gm_remarks"),
        ("HR Remarks",             "hr_remarks"),
    ]:
        if fd.get(key) and fd[key] not in ("", "-"):
            _long_field(doc, lbl, fd.get(key))

    _section(doc, "Signatures")
    sigs = [("Employee", fd.get("employee_signature"), fd.get("employee_sign_date"))]
    sigs.append(("Evaluator", fd.get("evaluator_signature"), fd.get("evaluator_sign_date")))
    if fd.get("hr_signature"):
        sigs.append(("HR", fd.get("hr_signature"), None))
    if fd.get("gm_signature"):
        sigs.append(("GM", fd.get("gm_signature"), None))
    _sig_block(doc, sigs)
    _footer(doc)
    return doc


def _build_contract_renewal(fd):
    doc = _new_doc()
    _add_header(doc, "Contract Renewal Assessment Form")
    _section(doc, "Employee Information")
    _data_table(doc, [
        ("Employee Name",     _fd(fd, "employee_name")),
        ("Employee ID",       _fd(fd, "employee_id")),
        ("Department",        _fd(fd, "department")),
        ("Designation",       _fd(fd, "designation")),
        ("Date of Joining",   _fmt(fd.get("date_of_joining"))),
        ("Contract End Date", _fmt(fd.get("contract_end_date"))),
        ("Date of Evaluation",_fmt(fd.get("date_of_evaluation"))),
        ("Evaluation By",     _fd(fd, "evaluation_by")),
    ], cols=4)

    # Sub-rating labels
    sub_labels = {
        "01": [
            ("01A", "Completes assigned tasks efficiently and accurately"),
            ("01B", "Demonstrates a strong understanding of job responsibilities"),
            ("01C", "Meets deadlines consistently"),
            ("01D", "Produces work of high quality"),
            ("01E", "Takes initiative to improve work processes"),
        ],
        "02": [
            ("02A", "Shows a positive attitude towards work and colleagues"),
            ("02B", "Accepts feedback constructively and strives for improvement"),
            ("02C", "Maintains professionalism in all interactions"),
            ("02D", "Demonstrates adaptability to changing situations"),
            ("02E", "Upholds company policies and values"),
        ],
        "03": [
            ("03A", "Communicates clearly and effectively with team members"),
            ("03B", "Works collaboratively and contributes to team goals"),
            ("03C", "Resolves conflicts or issues amicably"),
            ("03D", "Demonstrates good listening skills"),
            ("03E", "Keeps supervisors and colleagues informed of progress"),
        ],
        "04": [
            ("04A", "Arrives on time and prepared for work"),
            ("04B", "Maintains consistent attendance"),
            ("04C", "Provides notice and valid reasons for absences"),
        ],
    }
    section_titles = {
        "01": "SN 01 – Job Performance",
        "02": "SN 02 – Attitude & Work Ethics",
        "03": "SN 03 – Communication & Teamwork",
        "04": "SN 04 – Punctuality & Attendance",
    }

    for sn, sub_list in sub_labels.items():
        _section(doc, section_titles[sn])
        rows = []
        for suffix, label in sub_list:
            key = f"rating_{sn}{suffix[-1].lower()}"
            rows.append((label, _fd(fd, key), "1–5", ""))
        # Average
        avg_key = f"rating_{sn}"
        rows.append((f"Section {sn} Average", _fd(fd, avg_key), "1–5", "Avg"))
        _rating_table(doc, rows)
        comment_key = f"comments_{sn}"
        if fd.get(comment_key):
            _long_field(doc, "Comments", fd.get(comment_key))

    _section(doc, "Summary")
    _data_table(doc, [
        ("Total Score",          _fd(fd, "overall_score")),
        ("Recommendation",       _fd(fd, "recommendation", "").replace("_", " ").title()),
        ("Strength",             _fd(fd, "strength")),
        ("Areas for Improvement",_fd(fd, "areas_for_improvement")),
    ], cols=2)

    _section(doc, "Evaluator Signature")
    _sig_block(doc, [
        ("Evaluator", fd.get("evaluator_signature"), fd.get("evaluator_date")),
    ])
    _footer(doc)
    return doc


# ═════════════════════════════════════════════════════════════════════════════
# Public dispatch
# ═════════════════════════════════════════════════════════════════════════════
_BUILDERS = {
    "leave_application": _build_leave_application,
    "leave":             _build_leave_application,
    "commencement":      _build_commencement,
    "duty_resumption":   _build_duty_resumption,
    "passport_release":  _build_passport_release,
    "grievance":         _build_grievance,
    "visa_renewal":      _build_visa_renewal,
    "interview_assessment":  _build_interview_assessment,
    "staff_appraisal":       _build_staff_appraisal,
    "station_clearance":     _build_station_clearance,
    "performance_evaluation":_build_performance_evaluation,
    "contract_renewal":      _build_contract_renewal,
}


def build_professional_docx(form_type, form_data, output_stream, submission_id=None):
    """
    Generate a professional branded DOCX for the given form_type.
    Writes to output_stream (BytesIO or file path).
    Returns True on success, False if form_type is not supported.
    """
    builder = _BUILDERS.get(form_type)
    if not builder:
        return False
    doc = builder(form_data or {})
    if hasattr(output_stream, "write"):
        doc.save(output_stream)
    else:
        doc.save(output_stream)
    return True


def supports_builder(form_type):
    return form_type in _BUILDERS
